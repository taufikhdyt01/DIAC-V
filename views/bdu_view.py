import os
import sys
import io
import pandas as pd
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QLabel, QPushButton, QFrame, QGridLayout, QSpacerItem,
                             QSizePolicy, QScrollArea, QApplication, QMenu, QAction,
                             QTabWidget, QLineEdit, QComboBox, QTableWidget, QTableWidgetItem,
                             QHeaderView, QMessageBox, QFileDialog, QDateEdit, QCheckBox)
from PyQt5.QtGui import QPixmap, QIcon, QFont, QColor, QPalette, QCursor, QImage
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QPoint, QDate, QThread
from views.loading_screen import LoadingScreen, QuickLoadingDialog, show_loading_dialog
import tempfile
import shutil
import subprocess

# Import local modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import APP_NAME, SECONDARY_COLOR, PRIMARY_COLOR, BG_COLOR, DEPARTMENTS

try:
    from modules.formula_helper import SimpleFormulaEvaluator, FORMULA_CELLS, evaluate_formulas_background
    HAS_FORMULA_HELPER = True
except ImportError:
    HAS_FORMULA_HELPER = False
    print("Formula helper not available")
    
INDUSTRY_SUBTYPE_MAPPING = {
    "Business A - Palm Oil": [
        "Palm Oil - CPO", 
        "Palm Oil - CPKO", 
        "Palm Oil - EFB", 
        "Palm Oil - Plantation"
    ],
    "Business B - Mining, Oil & Gas, NFI A": [
        "Mining - Coal", 
        "Mining - Gold", 
        "Mining - Nickel", 
        "Mining - Tin", 
        "Mining - Bauxite", 
        "O&G - Upstream", 
        "O&G - Midstream", 
        "O&G - Downstream", 
        "Non Food - Textile", 
        "Non Food - Manufacturing / Heavy", 
        "Non Food - Tech & Telecom", 
        "Non Food - Transport & Log"
    ],
    "Business C - Food Beverage & Dairy, Agroindustry": [
        "F&B - Processed Food", 
        "F&B - Beverages", 
        "F&B - Dairy Product", 
        "F&B - Confectionery", 
        "F&B - Meat Processing", 
        "F&B - Seasoning", 
        "Agroindustry - Fishery & Aquaculture", 
        "Agroindustry - Food Crops", 
        "Agroindustry - Tobacco", 
        "Agroindustry - Sugar", 
        "Agroindustry - Livestock & Poultry"
    ],
    "Business D - OM, BOO, BOT": [],
    "Business E - NFI B": [
        "Non Food - Tourism & Hospitality", 
        "Non Food - Residential", 
        "Non Food - Construction & Real Estate", 
        "Non Food - PDAM/SPAM"
    ]
}

SEISMIC_ZONE_DESCRIPTIONS = {
    "ZONE-1": "2.5 or less. Usually not felt, but can be recorded by seismographs",
    "ZONE-2": "2.5 - 5.4. Often felt, but causes only minor damage",
    "ZONE-3": "5.5 - 6.0. Can cause slight damage to buildings and other structures",
    "ZONE-4": "6.6 - 6.9. Can cause significant damage in populated areas",
    "Zone-5": "7.0 - 7.9. Major earthquake with serious damage.",
    "Zone-6": "8.0 or larger. Great earthquake. Can destroy communities near the epicenter."
}

WIND_SPEED_DESCRIPTIONS = {
    "LEVEL-1": "0 - 2,0 m/s",
    "LEVEL-2": "2,0 - 4,0 m/s",
    "LEVEL-3": "4,0 - 6,0 m/s",
    "LEVEL-4": "6,0 - 8,0 m/s",
    "LEVEL-5": "8,0 - 10,0 m/s",
    "LEVEL-6": "10,0 m/s"
}

# Daftar provinsi di Indonesia
INDONESIA_PROVINCES = [
    "Aceh", "Sumatera Utara", "Sumatera Barat", "Riau", "Jambi", "Sumatera Selatan", 
    "Bengkulu", "Lampung", "Kepulauan Bangka Belitung", "Kepulauan Riau", 
    "DKI Jakarta", "Jawa Barat", "Jawa Tengah", "DI Yogyakarta", "Jawa Timur", "Banten", 
    "Bali", "Nusa Tenggara Barat", "Nusa Tenggara Timur", 
    "Kalimantan Barat", "Kalimantan Tengah", "Kalimantan Selatan", "Kalimantan Timur", "Kalimantan Utara", 
    "Sulawesi Utara", "Sulawesi Tengah", "Sulawesi Selatan", "Sulawesi Tenggara", "Gorontalo", "Sulawesi Barat", 
    "Maluku", "Maluku Utara", "Papua", "Papua Barat"
]

# Kota-kota dan kabupaten di Indonesia berdasarkan provinsi
INDONESIA_CITIES = {
    "Aceh": [
        "Kota Banda Aceh", "Kota Langsa", "Kota Lhokseumawe", "Kota Sabang", "Kota Subulussalam",
        "Kabupaten Aceh Barat", "Kabupaten Aceh Barat Daya", "Kabupaten Aceh Besar", 
        "Kabupaten Aceh Jaya", "Kabupaten Aceh Selatan", "Kabupaten Aceh Singkil", 
        "Kabupaten Aceh Tamiang", "Kabupaten Aceh Tengah", "Kabupaten Aceh Tenggara", 
        "Kabupaten Aceh Timur", "Kabupaten Aceh Utara", "Kabupaten Bener Meriah", 
        "Kabupaten Bireuen", "Kabupaten Gayo Lues", "Kabupaten Nagan Raya", 
        "Kabupaten Pidie", "Kabupaten Pidie Jaya", "Kabupaten Simeulue"
    ],
    
    "Sumatera Utara": [
        "Kota Medan", "Kota Binjai", "Kota Gunungsitoli", "Kota Padang Sidempuan", 
        "Kota Pematangsiantar", "Kota Sibolga", "Kota Tanjungbalai", "Kota Tebing Tinggi",
        "Kabupaten Asahan", "Kabupaten Batu Bara", "Kabupaten Dairi", "Kabupaten Deli Serdang", 
        "Kabupaten Humbang Hasundutan", "Kabupaten Karo", "Kabupaten Labuhanbatu", 
        "Kabupaten Labuhanbatu Selatan", "Kabupaten Labuhanbatu Utara", "Kabupaten Langkat", 
        "Kabupaten Mandailing Natal", "Kabupaten Nias", "Kabupaten Nias Barat", 
        "Kabupaten Nias Selatan", "Kabupaten Nias Utara", "Kabupaten Padang Lawas", 
        "Kabupaten Padang Lawas Utara", "Kabupaten Pakpak Bharat", "Kabupaten Samosir", 
        "Kabupaten Serdang Bedagai", "Kabupaten Simalungun", "Kabupaten Tapanuli Selatan", 
        "Kabupaten Tapanuli Tengah", "Kabupaten Tapanuli Utara", "Kabupaten Toba Samosir"
    ],
    
    "Sumatera Barat": [
        "Kota Bukittinggi", "Kota Padang", "Kota Padang Panjang", "Kota Pariaman", 
        "Kota Payakumbuh", "Kota Sawahlunto", "Kota Solok",
        "Kabupaten Agam", "Kabupaten Dharmasraya", "Kabupaten Kepulauan Mentawai", 
        "Kabupaten Lima Puluh Kota", "Kabupaten Padang Pariaman", "Kabupaten Pasaman", 
        "Kabupaten Pasaman Barat", "Kabupaten Pesisir Selatan", "Kabupaten Sijunjung", 
        "Kabupaten Solok", "Kabupaten Solok Selatan", "Kabupaten Tanah Datar"
    ],
    
    "Riau": [
        "Kota Dumai", "Kota Pekanbaru",
        "Kabupaten Bengkalis", "Kabupaten Indragiri Hilir", "Kabupaten Indragiri Hulu", 
        "Kabupaten Kampar", "Kabupaten Kepulauan Meranti", "Kabupaten Kuantan Singingi", 
        "Kabupaten Pelalawan", "Kabupaten Rokan Hilir", "Kabupaten Rokan Hulu", 
        "Kabupaten Siak"
    ],
    
    "Jambi": [
        "Kota Jambi", "Kota Sungai Penuh",
        "Kabupaten Batanghari", "Kabupaten Bungo", "Kabupaten Kerinci", 
        "Kabupaten Merangin", "Kabupaten Muaro Jambi", "Kabupaten Sarolangun", 
        "Kabupaten Tanjung Jabung Barat", "Kabupaten Tanjung Jabung Timur", 
        "Kabupaten Tebo"
    ],
    
    "Sumatera Selatan": [
        "Kota Lubuklinggau", "Kota Pagar Alam", "Kota Palembang", "Kota Prabumulih",
        "Kabupaten Banyuasin", "Kabupaten Empat Lawang", "Kabupaten Lahat", 
        "Kabupaten Muara Enim", "Kabupaten Musi Banyuasin", "Kabupaten Musi Rawas", 
        "Kabupaten Musi Rawas Utara", "Kabupaten Ogan Ilir", "Kabupaten Ogan Komering Ilir", 
        "Kabupaten Ogan Komering Ulu", "Kabupaten Ogan Komering Ulu Selatan", 
        "Kabupaten Ogan Komering Ulu Timur", "Kabupaten Penukal Abab Lematang Ilir"
    ],
    
    "Bengkulu": [
        "Kota Bengkulu",
        "Kabupaten Bengkulu Selatan", "Kabupaten Bengkulu Tengah", "Kabupaten Bengkulu Utara", 
        "Kabupaten Kaur", "Kabupaten Kepahiang", "Kabupaten Lebong", 
        "Kabupaten Mukomuko", "Kabupaten Rejang Lebong", "Kabupaten Seluma"
    ],
    
    "Lampung": [
        "Kota Bandar Lampung", "Kota Metro",
        "Kabupaten Lampung Barat", "Kabupaten Lampung Selatan", "Kabupaten Lampung Tengah", 
        "Kabupaten Lampung Timur", "Kabupaten Lampung Utara", "Kabupaten Mesuji", 
        "Kabupaten Pesawaran", "Kabupaten Pesisir Barat", "Kabupaten Pringsewu", 
        "Kabupaten Tanggamus", "Kabupaten Tulang Bawang", "Kabupaten Tulang Bawang Barat", 
        "Kabupaten Way Kanan"
    ],
    
    "Kepulauan Bangka Belitung": [
        "Kota Pangkalpinang",
        "Kabupaten Bangka", "Kabupaten Bangka Barat", "Kabupaten Bangka Selatan", 
        "Kabupaten Bangka Tengah", "Kabupaten Belitung", "Kabupaten Belitung Timur"
    ],
    
    "Kepulauan Riau": [
        "Kota Batam", "Kota Tanjungpinang",
        "Kabupaten Bintan", "Kabupaten Karimun", "Kabupaten Kepulauan Anambas", 
        "Kabupaten Lingga", "Kabupaten Natuna"
    ],
    
    "DKI Jakarta": [
        "Kota Jakarta Barat", "Kota Jakarta Pusat", "Kota Jakarta Selatan", 
        "Kota Jakarta Timur", "Kota Jakarta Utara", "Kabupaten Kepulauan Seribu"
    ],
    
    "Jawa Barat": [
        "Kota Bandung", "Kota Banjar", "Kota Bekasi", "Kota Bogor", "Kota Cimahi", 
        "Kota Cirebon", "Kota Depok", "Kota Sukabumi", "Kota Tasikmalaya",
        "Kabupaten Bandung", "Kabupaten Bandung Barat", "Kabupaten Bekasi", "Kabupaten Bogor", 
        "Kabupaten Ciamis", "Kabupaten Cianjur", "Kabupaten Cirebon", "Kabupaten Garut", 
        "Kabupaten Indramayu", "Kabupaten Karawang", "Kabupaten Kuningan", "Kabupaten Majalengka", 
        "Kabupaten Pangandaran", "Kabupaten Purwakarta", "Kabupaten Subang", "Kabupaten Sukabumi", 
        "Kabupaten Sumedang", "Kabupaten Tasikmalaya"
    ],
    
    "Jawa Tengah": [
        "Kota Magelang", "Kota Pekalongan", "Kota Salatiga", "Kota Semarang", "Kota Surakarta", "Kota Tegal",
        "Kabupaten Banjarnegara", "Kabupaten Banyumas", "Kabupaten Batang", "Kabupaten Blora", 
        "Kabupaten Boyolali", "Kabupaten Brebes", "Kabupaten Cilacap", "Kabupaten Demak", 
        "Kabupaten Grobogan", "Kabupaten Jepara", "Kabupaten Karanganyar", "Kabupaten Kebumen", 
        "Kabupaten Kendal", "Kabupaten Klaten", "Kabupaten Kudus", "Kabupaten Magelang", 
        "Kabupaten Pati", "Kabupaten Pekalongan", "Kabupaten Pemalang", "Kabupaten Purbalingga", 
        "Kabupaten Purworejo", "Kabupaten Rembang", "Kabupaten Semarang", "Kabupaten Sragen", 
        "Kabupaten Sukoharjo", "Kabupaten Tegal", "Kabupaten Temanggung", "Kabupaten Wonogiri", 
        "Kabupaten Wonosobo"
    ],
    
    "DI Yogyakarta": [
        "Kota Yogyakarta",
        "Kabupaten Bantul", "Kabupaten Gunungkidul", "Kabupaten Kulon Progo", 
        "Kabupaten Sleman"
    ],
    
    "Jawa Timur": [
        "Kota Batu", "Kota Blitar", "Kota Kediri", "Kota Madiun", "Kota Malang", 
        "Kota Mojokerto", "Kota Pasuruan", "Kota Probolinggo", "Kota Surabaya",
        "Kabupaten Bangkalan", "Kabupaten Banyuwangi", "Kabupaten Blitar", "Kabupaten Bojonegoro", 
        "Kabupaten Bondowoso", "Kabupaten Gresik", "Kabupaten Jember", "Kabupaten Jombang", 
        "Kabupaten Kediri", "Kabupaten Lamongan", "Kabupaten Lumajang", "Kabupaten Madiun", 
        "Kabupaten Magetan", "Kabupaten Malang", "Kabupaten Mojokerto", "Kabupaten Nganjuk", 
        "Kabupaten Ngawi", "Kabupaten Pacitan", "Kabupaten Pamekasan", "Kabupaten Pasuruan", 
        "Kabupaten Ponorogo", "Kabupaten Probolinggo", "Kabupaten Sampang", "Kabupaten Sidoarjo", 
        "Kabupaten Situbondo", "Kabupaten Sumenep", "Kabupaten Trenggalek", "Kabupaten Tuban", 
        "Kabupaten Tulungagung"
    ],
    
    "Banten": [
        "Kota Cilegon", "Kota Serang", "Kota Tangerang", "Kota Tangerang Selatan",
        "Kabupaten Lebak", "Kabupaten Pandeglang", "Kabupaten Serang", "Kabupaten Tangerang"
    ],
    
    "Bali": [
        "Kota Denpasar",
        "Kabupaten Badung", "Kabupaten Bangli", "Kabupaten Buleleng", "Kabupaten Gianyar", 
        "Kabupaten Jembrana", "Kabupaten Karangasem", "Kabupaten Klungkung", "Kabupaten Tabanan"
    ],
    
    "Nusa Tenggara Barat": [
        "Kota Bima", "Kota Mataram",
        "Kabupaten Bima", "Kabupaten Dompu", "Kabupaten Lombok Barat", "Kabupaten Lombok Tengah", 
        "Kabupaten Lombok Timur", "Kabupaten Lombok Utara", "Kabupaten Sumbawa", 
        "Kabupaten Sumbawa Barat"
    ],
    
    "Nusa Tenggara Timur": [
        "Kota Kupang",
        "Kabupaten Alor", "Kabupaten Belu", "Kabupaten Ende", "Kabupaten Flores Timur", 
        "Kabupaten Kupang", "Kabupaten Lembata", "Kabupaten Malaka", "Kabupaten Manggarai", 
        "Kabupaten Manggarai Barat", "Kabupaten Manggarai Timur", "Kabupaten Nagekeo", 
        "Kabupaten Ngada", "Kabupaten Rote Ndao", "Kabupaten Sabu Raijua", "Kabupaten Sikka", 
        "Kabupaten Sumba Barat", "Kabupaten Sumba Barat Daya", "Kabupaten Sumba Tengah", 
        "Kabupaten Sumba Timur", "Kabupaten Timor Tengah Selatan", "Kabupaten Timor Tengah Utara"
    ],
    
    "Kalimantan Barat": [
        "Kota Pontianak", "Kota Singkawang",
        "Kabupaten Bengkayang", "Kabupaten Kapuas Hulu", "Kabupaten Kayong Utara", 
        "Kabupaten Ketapang", "Kabupaten Kubu Raya", "Kabupaten Landak", "Kabupaten Melawi", 
        "Kabupaten Mempawah", "Kabupaten Sambas", "Kabupaten Sanggau", "Kabupaten Sekadau", 
        "Kabupaten Sintang"
    ],
    
    "Kalimantan Tengah": [
        "Kota Palangka Raya",
        "Kabupaten Barito Selatan", "Kabupaten Barito Timur", "Kabupaten Barito Utara", 
        "Kabupaten Gunung Mas", "Kabupaten Kapuas", "Kabupaten Katingan", 
        "Kabupaten Kotawaringin Barat", "Kabupaten Kotawaringin Timur", "Kabupaten Lamandau", 
        "Kabupaten Murung Raya", "Kabupaten Pulang Pisau", "Kabupaten Sukamara", 
        "Kabupaten Seruyan"
    ],
    
    "Kalimantan Selatan": [
        "Kota Banjarbaru", "Kota Banjarmasin",
        "Kabupaten Balangan", "Kabupaten Banjar", "Kabupaten Barito Kuala", "Kabupaten Hulu Sungai Selatan", 
        "Kabupaten Hulu Sungai Tengah", "Kabupaten Hulu Sungai Utara", "Kabupaten Kotabaru", 
        "Kabupaten Tabalong", "Kabupaten Tanah Bumbu", "Kabupaten Tanah Laut", "Kabupaten Tapin"
    ],
    
    "Kalimantan Timur": [
        "Kota Balikpapan", "Kota Bontang", "Kota Samarinda",
        "Kabupaten Berau", "Kabupaten Kutai Barat", "Kabupaten Kutai Kartanegara", 
        "Kabupaten Kutai Timur", "Kabupaten Mahakam Ulu", "Kabupaten Paser", 
        "Kabupaten Penajam Paser Utara"
    ],
    
    "Kalimantan Utara": [
        "Kota Tarakan",
        "Kabupaten Bulungan", "Kabupaten Malinau", "Kabupaten Nunukan", "Kabupaten Tana Tidung"
    ],
    
    "Sulawesi Utara": [
        "Kota Bitung", "Kota Kotamobagu", "Kota Manado", "Kota Tomohon",
        "Kabupaten Bolaang Mongondow", "Kabupaten Bolaang Mongondow Selatan", 
        "Kabupaten Bolaang Mongondow Timur", "Kabupaten Bolaang Mongondow Utara", 
        "Kabupaten Kepulauan Sangihe", "Kabupaten Kepulauan Siau Tagulandang Biaro", 
        "Kabupaten Kepulauan Talaud", "Kabupaten Minahasa", "Kabupaten Minahasa Selatan", 
        "Kabupaten Minahasa Tenggara", "Kabupaten Minahasa Utara"
    ],
    
    "Sulawesi Tengah": [
        "Kota Palu",
        "Kabupaten Banggai", "Kabupaten Banggai Kepulauan", "Kabupaten Banggai Laut", 
        "Kabupaten Buol", "Kabupaten Donggala", "Kabupaten Morowali", "Kabupaten Morowali Utara", 
        "Kabupaten Parigi Moutong", "Kabupaten Poso", "Kabupaten Sigi", "Kabupaten Tojo Una-Una", 
        "Kabupaten Tolitoli"
    ],
    
    "Sulawesi Selatan": [
        "Kota Makassar", "Kota Palopo", "Kota Parepare",
        "Kabupaten Bantaeng", "Kabupaten Barru", "Kabupaten Bone", "Kabupaten Bulukumba", 
        "Kabupaten Enrekang", "Kabupaten Gowa", "Kabupaten Jeneponto", "Kabupaten Kepulauan Selayar", 
        "Kabupaten Luwu", "Kabupaten Luwu Timur", "Kabupaten Luwu Utara", "Kabupaten Maros", 
        "Kabupaten Pangkajene dan Kepulauan", "Kabupaten Pinrang", "Kabupaten Sidenreng Rappang", 
        "Kabupaten Sinjai", "Kabupaten Soppeng", "Kabupaten Takalar", "Kabupaten Tana Toraja", 
        "Kabupaten Toraja Utara", "Kabupaten Wajo"
    ],
    
    "Sulawesi Tenggara": [
        "Kota Baubau", "Kota Kendari",
        "Kabupaten Bombana", "Kabupaten Buton", "Kabupaten Buton Selatan", "Kabupaten Buton Tengah", 
        "Kabupaten Buton Utara", "Kabupaten Kolaka", "Kabupaten Kolaka Timur", "Kabupaten Kolaka Utara", 
        "Kabupaten Konawe", "Kabupaten Konawe Kepulauan", "Kabupaten Konawe Selatan", 
        "Kabupaten Konawe Utara", "Kabupaten Muna", "Kabupaten Muna Barat", "Kabupaten Wakatobi"
    ],
    
    "Gorontalo": [
        "Kota Gorontalo",
        "Kabupaten Boalemo", "Kabupaten Bone Bolango", "Kabupaten Gorontalo", 
        "Kabupaten Gorontalo Utara", "Kabupaten Pohuwato"
    ],
    
    "Sulawesi Barat": [
        "Kabupaten Majene", "Kabupaten Mamasa", "Kabupaten Mamuju", "Kabupaten Mamuju Tengah", 
        "Kabupaten Pasangkayu", "Kabupaten Polewali Mandar"
    ],
    
    "Maluku": [
        "Kota Ambon", "Kota Tual",
        "Kabupaten Buru", "Kabupaten Buru Selatan", "Kabupaten Kepulauan Aru", 
        "Kabupaten Maluku Barat Daya", "Kabupaten Maluku Tengah", "Kabupaten Maluku Tenggara", 
        "Kabupaten Maluku Tenggara Barat", "Kabupaten Seram Bagian Barat", "Kabupaten Seram Bagian Timur"
    ],
    
    "Maluku Utara": [
        "Kota Ternate", "Kota Tidore Kepulauan",
        "Kabupaten Halmahera Barat", "Kabupaten Halmahera Tengah", "Kabupaten Halmahera Timur", 
        "Kabupaten Halmahera Selatan", "Kabupaten Halmahera Utara", "Kabupaten Kepulauan Sula", 
        "Kabupaten Pulau Morotai", "Kabupaten Pulau Taliabu"
    ],
    
    "Papua": [
        "Kota Jayapura",
        "Kabupaten Asmat", "Kabupaten Biak Numfor", "Kabupaten Boven Digoel", "Kabupaten Deiyai", 
        "Kabupaten Dogiyai", "Kabupaten Intan Jaya", "Kabupaten Jayapura", "Kabupaten Jayawijaya", 
        "Kabupaten Keerom", "Kabupaten Kepulauan Yapen", "Kabupaten Lanny Jaya", "Kabupaten Mamberamo Raya", 
        "Kabupaten Mamberamo Tengah", "Kabupaten Mappi", "Kabupaten Merauke", "Kabupaten Mimika", 
        "Kabupaten Nabire", "Kabupaten Nduga", "Kabupaten Paniai", "Kabupaten Pegunungan Bintang", 
        "Kabupaten Puncak", "Kabupaten Puncak Jaya", "Kabupaten Sarmi", "Kabupaten Supiori", 
        "Kabupaten Tolikara", "Kabupaten Waropen", "Kabupaten Yahukimo", "Kabupaten Yalimo"
    ],
    
    "Papua Barat": [
        "Kota Sorong",
        "Kabupaten Fakfak", "Kabupaten Kaimana", "Kabupaten Manokwari", "Kabupaten Manokwari Selatan", 
        "Kabupaten Maybrat", "Kabupaten Pegunungan Arfak", "Kabupaten Raja Ampat", "Kabupaten Sorong", 
        "Kabupaten Sorong Selatan", "Kabupaten Tambrauw", "Kabupaten Teluk Bintuni", "Kabupaten Teluk Wondama"
    ]
}

EFFLUENT_WARRANTY_OPTIONS = [
    "PERMENKES No. 2 Tahun 2023 (Parameter Wajib Air Minum)",
    "PERMENKES No. 2 Tahun 2023 (Parameter Air untuk Keperluan Higiene dan Sanitasi)",
    "PERMENLHK RI No. P.68 Tahun 2016 (Baku Mutu Air Limbah Domestik)",
    "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 1 dan Sejenisnya)",
    "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 2 dan Sejenisnya)",
    "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 3 dan Sejenisnya)",
    "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 4 dan Sejenisnya)"
]

PUMP_BRAND_TYPE_MAPPING = {
    "GRUNDFOS": [
        "Vertical Multistage Centrifugal Pump",
        "End Suction Centrifugal Pump"
    ],
    "EBARA": [
        "Vertical Multistage Centrifugal Pump", 
        "End Suction Centrifugal Pump"
    ],
    "CNP": [
        "Vertical Multistage Centrifugal Pump",
        "Horizontal Multistage Centrifugal Pump",
        "Submersible Pump"
    ],
    "LEO": [
        "Vertical Multistage Centrifugal Pump",
        "End Suction Centrifugal Pump", 
        "Submersible Pump"
    ]
}

PUMP_BRAND_TYPE_MODEL_MAPPING = {
    ("GRUNDFOS", "Vertical Multistage Centrifugal Pump"): ["CR", "CRN"],
    ("GRUNDFOS", "End Suction Centrifugal Pump"): ["NKG"],
    ("EBARA", "Vertical Multistage Centrifugal Pump"): ["3S"],
    ("EBARA", "End Suction Centrifugal Pump"): ["FSSC"],
    ("CNP", "Vertical Multistage Centrifugal Pump"): ["CDMF"],
    ("CNP", "Horizontal Multistage Centrifugal Pump"): ["CHL"],
    ("CNP", "Submersible Pump"): ["WQ"],
    ("LEO", "Vertical Multistage Centrifugal Pump"): ["LVRS"],
    ("LEO", "End Suction Centrifugal Pump"): ["LEP"],
    ("LEO", "Submersible Pump"): ["SWE", "XSP"]
}

def get_effluent_warranty_parameters_for_tooltip(warranty_type):
    """
    Get effluent warranty parameters in simple format for tooltip
    
    Parameters:
    - warranty_type: Selected warranty type
    
    Returns:
    - str: Formatted parameter string for tooltip (vertical format)
    """
    
    # Data parameter berdasarkan tabel yang Anda berikan
    parameter_data = {
        "PERMENKES No. 2 Tahun 2023 (Parameter Wajib Air Minum)": [
            "pH: 6.5-8.5",
            "TDS: <300 mg/L",
            "Turbidity: <3 NTU", 
            "Color: 10 TCU",
            "Nitrate: 20 mg/L",
            "Nitrite: 3 mg/L",
            "Cr6+: 0.01 mg/L",
            "Fe: 0.2 mg/L",
            "Mn: 0.1 mg/L",
            "Sisa Cl: 0.2-0.5 mg/L",
            "As: 0.01 mg/L",
            "Cd: 0.003 mg/L",
            "Pb: 0.01 mg/L",
            "F: 1.5 mg/L",
            "Al: 0.2 mg/L",
            "E.Coli: 0 CFU/100 mL",
            "Total Coliform: 0 CFU/100 mL",
            "Odor: odorless"
        ],
        
        "PERMENKES No. 2 Tahun 2023 (Parameter Air untuk Keperluan Higiene dan Sanitasi)": [
            "pH: 6.5-8.5",
            "TDS: <300 mg/L",
            "Turbidity: <3 NTU",
            "Color: 10 TCU", 
            "Nitrate: 20 mg/L",
            "Nitrite: 3 mg/L",
            "Cr6+: 0.01 mg/L",
            "Fe: 0.2 mg/L",
            "Mn: 0.1 mg/L",
            "E.Coli: 0 CFU/100 mL",
            "Total Coliform: 0 CFU/100 mL",
            "Odor: odorless"
        ],
        
        "PERMENLHK RI No. P.68 Tahun 2016 (Baku Mutu Air Limbah Domestik)": [
            "pH: 6-9",
            "BOD: 30 mg/L",
            "COD: 100 mg/L",
            "TSS: 30 mg/L",
            "FOG: 5 mg/L",
            "Ammonia: 10 mg/L",
            "Total Coliform: 3000 jumlah/100 mL"
        ],
        
        "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 1 dan Sejenisnya)": [
            "pH: 6-9",
            "TDS: 1000 mg/L",
            "TSS: 40 mg/L",
            "Color: 15 Pt-Co Unit",
            "BOD: 2 mg/L",
            "COD: 10 mg/L",
            "DO: 6 mg/L",
            "Sulfate: 300 mg/L",
            "Chloride: 300 mg/L",
            "Nitrate: 10 mg/L",
            "Nitrite: 0.06 mg/L",
            "Ammonia: 0.1 mg/L",
            "Total Nitrogen: 15 mg/L",
            "Total Phospate: 0.2 mg/L",
            "Fluoride: 1 mg/L",
            "Sulphure as H2S: 0.002 mg/L",
            "Cyanide: 0.02 mg/L",
            "Free Chlorine: 0.03 mg/L",
            "Oil and Grease: 1 mg/L",
            "Total Detergent: 0.2 mg/L",
            "Fecal Coliform: 100 MPN/100 mL",
            "Total Coliform: 1000 MPN/100 mL"
        ],
        
        "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 2 dan Sejenisnya)": [
            "pH: 6-9",
            "TDS: 1000 mg/L",
            "TSS: 50 mg/L",
            "Color: 50 Pt-Co Unit",
            "BOD: 3 mg/L",
            "COD: 25 mg/L",
            "DO: 4 mg/L",
            "Sulfate: 300 mg/L",
            "Chloride: 300 mg/L",
            "Nitrate: 10 mg/L",
            "Nitrite: 0.06 mg/L",
            "Ammonia: 0.2 mg/L",
            "Total Nitrogen: 15 mg/L",
            "Total Phospate: 0.2 mg/L",
            "Fluoride: 1.5 mg/L",
            "Sulphure as H2S: 0.002 mg/L",
            "Cyanide: 0.02 mg/L",
            "Free Chlorine: 0.03 mg/L",
            "Oil and Grease: 1 mg/L",
            "Total Detergent: 0.2 mg/L",
            "Fecal Coliform: 1000 MPN/100 mL",
            "Total Coliform: 5000 MPN/100 mL"
        ],
        
        "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 3 dan Sejenisnya)": [
            "pH: 6-9",
            "TDS: 1000 mg/L",
            "TSS: 100 mg/L",
            "Color: 100 Pt-Co Unit",
            "BOD: 6 mg/L",
            "COD: 40 mg/L",
            "DO: 3 mg/L",
            "Sulfate: 300 mg/L",
            "Chloride: 300 mg/L",
            "Nitrate: 20 mg/L",
            "Nitrite: 0.06 mg/L",
            "Ammonia: 0.5 mg/L",
            "Total Nitrogen: 25 mg/L",
            "Total Phospate: 1 mg/L",
            "Fluoride: 1.5 mg/L",
            "Sulphure as H2S: 0.002 mg/L",
            "Cyanide: 0.02 mg/L",
            "Free Chlorine: 0.03 mg/L",
            "Oil and Grease: 1 mg/L",
            "Total Detergent: 0.2 mg/L",
            "Fecal Coliform: 2000 MPN/100 mL",
            "Total Coliform: 10000 MPN/100 mL"
        ],
        
        "PP RI No. 22 Tahun 2021 (Baku Mutu Air Sungai Kelas 4 dan Sejenisnya)": [
            "pH: 6-9",
            "TDS: 2000 mg/L",
            "TSS: 400 mg/L",
            "BOD: 12 mg/L",
            "COD: 80 mg/L",
            "DO: 1 mg/L",
            "Sulfate: 400 mg/L",
            "Chloride: 600 mg/L",
            "Nitrate: 20 mg/L",
            "Oil and Grease: 10 mg/L",
            "Fecal Coliform: 2000 MPN/100 mL",
            "Total Coliform: 10000 MPN/100 mL"
        ]
    }
    
    # Ambil parameter untuk warranty type yang dipilih
    parameters = parameter_data.get(warranty_type, [])
    
    if not parameters:
        return "No parameters available"
    
    formatted_params = []
    
    for param in parameters: 
        formatted_params.append(f"• {param}")
    
    # JOIN DENGAN \n UNTUK FORMAT VERTIKAL (KE BAWAH)
    return "\n".join(formatted_params)

def setup_effluent_warranty_item_tooltips(dropdown_widget):
    """
    Setup tooltip pada setiap item di dropdown Effluent Warranty
    
    Parameters:
    - dropdown_widget: QComboBox widget
    """
    try:
        # Loop through all items in dropdown
        for i in range(dropdown_widget.count()):
            item_text = dropdown_widget.itemText(i)
            
            # Skip placeholder item
            if item_text == "-- Select Value --":
                continue
            
            # Get parameters for this warranty type
            tooltip_text = get_effluent_warranty_parameters_for_tooltip(item_text)
            
            # Set tooltip untuk item ini menggunakan setItemData dengan Qt.ToolTipRole
            dropdown_widget.setItemData(i, tooltip_text, Qt.ToolTipRole)
        
    except Exception as e:
        print(f"Error setting up item tooltips: {str(e)}")
        
# Pengecekan opsional untuk modul docx2pdf
HAS_DOCX2PDF = False
try:
    import docx2pdf
    HAS_DOCX2PDF = True
except ImportError:
    pass

# Pengecekan opsional untuk modul PyMuPDF (fitz)
HAS_PYMUPDF = False
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    pass

# Kelas thread untuk mengkonversi dokumen Word ke gambar thumbnail
class ThumbnailGeneratorThread(QThread):
    # Signal saat thumbnail siap
    thumbnail_ready = pyqtSignal(int, QPixmap)
    finished = pyqtSignal()
    error = pyqtSignal(str)
    
    def __init__(self, file_path, max_pages=100):  # Ubah max_pages menjadi 100
        super().__init__()
        self.file_path = file_path
        self.max_pages = max_pages
        self.temp_dir = None
        
    def run(self):
        try:
             # Inisialisasi COM library untuk thread ini
            import pythoncom
            pythoncom.CoInitialize()
            
            # Buat direktori sementara
            self.temp_dir = tempfile.mkdtemp()
            
            # Konversi Word ke PDF jika itu docx dan modul docx2pdf tersedia
            _, ext = os.path.splitext(self.file_path)
            if ext.lower() == '.docx' and HAS_DOCX2PDF:
                pdf_path = os.path.join(self.temp_dir, "temp.pdf")
                docx2pdf.convert(self.file_path, pdf_path)
                
                # Gunakan PDF untuk menghasilkan thumbnail jika PyMuPDF tersedia
                if HAS_PYMUPDF:
                    self.generate_thumbnails_from_pdf(pdf_path)
                else:
                    self.generate_thumbnails_alternative()
            else:
                # Gunakan pendekatan alternatif untuk file non-PDF dan jika docx2pdf tidak tersedia
                self.generate_thumbnails_alternative()
                
            self.finished.emit()
        except Exception as e:
            self.error.emit(str(e))
        finally:
            # Bersihkan direktori sementara saat selesai
            if self.temp_dir and os.path.exists(self.temp_dir):
                try:
                    shutil.rmtree(self.temp_dir)
                except:
                    pass
                
            # Uninisialisasi COM library
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass

    def generate_thumbnails_from_pdf(self, pdf_path):
        """Menghasilkan thumbnail dari file PDF menggunakan PyMuPDF jika tersedia"""
        try:
            # Pastikan PyMuPDF tersedia
            if not HAS_PYMUPDF:
                self.error.emit("PyMuPDF (modul fitz) tidak tersedia. Menggunakan metode alternatif.")
                self.generate_thumbnails_alternative()
                return
                
            doc = fitz.open(pdf_path)
            # Proses semua halaman tanpa batasan 10
            for page_num in range(min(self.max_pages, doc.page_count)):
                page = doc.load_page(page_num)
                
                # Render halaman sebagai gambar
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                img_data = pix.samples
                
                # Buat QImage dari data piksel
                img = QImage(img_data, pix.width, pix.height, pix.stride, QImage.Format_RGB888)
                pixmap = QPixmap.fromImage(img)
                
                # Kirim thumbnail ke UI
                self.thumbnail_ready.emit(page_num, pixmap)
            
            doc.close()
        except Exception as e:
            self.error.emit(f"Error PyMuPDF: {str(e)}. Menggunakan metode alternatif.")
            self.generate_thumbnails_alternative()
    
    def generate_thumbnails_alternative(self):
        """Metode alternatif untuk menghasilkan thumbnail sederhana"""
        try:
            import docx
            
            # Buat thumbnail sederhana yang menunjukkan judul dokumen dan beberapa baris teks
            doc = docx.Document(self.file_path)
            
            # Hitung jumlah halaman (perkiraan)
            total_paragraphs = len(doc.paragraphs)
            estimated_pages = max(1, total_paragraphs // 20)  # Asumsikan ~20 paragraf per halaman
            
            # Tampilkan semua halaman yang diperkirakan, tanpa dibatasi 10
            for page_num in range(min(self.max_pages, estimated_pages)):
                # Buat gambar kosong
                width, height = 800, 1000
                img = QImage(width, height, QImage.Format_RGB888)
                img.fill(Qt.white)
                
                # Buat objek painter
                painter = QtGui.QPainter(img)
                painter.setFont(QFont("Arial", 12))
                
                # Gambar judul dokumen
                painter.setFont(QFont("Arial", 16, QFont.Bold))
                painter.drawText(40, 40, "Page {}".format(page_num + 1))
                
                # Gambar beberapa teks dari dokumen
                painter.setFont(QFont("Arial", 11))
                
                y_pos = 80
                start_para = page_num * 20
                end_para = min(start_para + 20, total_paragraphs)
                
                for i in range(start_para, end_para):
                    text = doc.paragraphs[i].text
                    if text:
                        wrapped_text = self.wrap_text(painter, text, width - 80)
                        for line in wrapped_text:
                            painter.drawText(40, y_pos, line)
                            y_pos += 20
                            if y_pos > height - 40:
                                break
                    y_pos += 10
                    if y_pos > height - 40:
                        break
                
                painter.end()
                
                # Konversi ke QPixmap
                pixmap = QPixmap.fromImage(img)
                
                # Kirim thumbnail ke UI
                self.thumbnail_ready.emit(page_num, pixmap)
        except Exception as e:
            self.error.emit(f"Gagal membuat thumbnail alternatif: {str(e)}")
            
            # Buat thumbnail sangat sederhana sebagai fallback terakhir
            self.generate_simple_fallback(estimated_pages=min(100, estimated_pages))  # Tingkatkan jumlah halaman
    
    def generate_simple_fallback(self, estimated_pages=100):
        """Membuat thumbnail sederhana sebagai fallback terakhir"""
        try:
            for page_num in range(min(self.max_pages, estimated_pages)):
                # Buat gambar kosong
                width, height = 800, 1000
                img = QImage(width, height, QImage.Format_RGB888)
                img.fill(Qt.white)
                
                # Buat objek painter
                painter = QtGui.QPainter(img)
                
                # Gambar judul dokumen
                painter.setFont(QFont("Arial", 16, QFont.Bold))
                painter.drawText(40, 40, "Page {}".format(page_num + 1))
                
                # Gambar pesan kesalahan atau info
                painter.setFont(QFont("Arial", 12))
                painter.drawText(40, 80, "Pratinjau dokumen tidak tersedia.")
                painter.drawText(40, 110, "Silakan gunakan tombol 'Open in Word'")
                painter.drawText(40, 140, "untuk membuka dokumen langsung di Word.")
                
                # Gambar ikon dokumen besar di tengah
                painter.setFont(QFont("Arial", 100))
                painter.drawText(width/2 - 80, height/2, "📄")
                
                painter.end()
                
                # Konversi ke QPixmap
                pixmap = QPixmap.fromImage(img)
                
                # Kirim thumbnail ke UI
                self.thumbnail_ready.emit(page_num, pixmap)
        except Exception as e:
            self.error.emit(f"Gagal membuat thumbnail fallback: {str(e)}")
    
    def wrap_text(self, painter, text, max_width):
        """Membantu memotong teks untuk muat dalam lebar tertentu"""
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            test_line = current_line + " " + word if current_line else word
            width = painter.fontMetrics().horizontalAdvance(test_line)
            
            if width <= max_width:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
            
        return lines
class BDUGroupView(QMainWindow):
    """View untuk BDU Group"""
    back_to_dashboard = pyqtSignal()
    
    def __init__(self, auth_manager):
        super().__init__()
        self.auth_manager = auth_manager
        self.current_user = auth_manager.get_current_user()
        self.excel_data = None
        self.sheet_tabs = {}
        self.data_fields = {}
        self.linked_dropdowns = {}
        self.user_codes = []
        self.formula_evaluator = None
        self.formula_widgets = {}
        
        # Excel path
        self.excel_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "SET_BDU.xlsx")
        
        self.initUI()
        self.load_excel_data()
    
    def initUI(self):
        """Initialize the UI"""
        # Set window properties
        self.setWindowTitle(f"BDU Group - {APP_NAME}")
        self.setMinimumSize(1000, 700)
        
        # Set central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Header
        self.setup_header()
        main_layout.addWidget(self.header_widget)
        
        # Content area
        content_widget = QScrollArea()
        content_widget.setWidgetResizable(True)
        content_widget.setStyleSheet(f"background-color: {BG_COLOR}; border: none;")
        
        # Main content
        self.main_content = QWidget()
        self.main_content.setStyleSheet(f"background-color: {BG_COLOR};")
        
        content_layout = QVBoxLayout(self.main_content)
        content_layout.setContentsMargins(20, 20, 20, 20)
        content_layout.setSpacing(15)
        
        # Page title
        title_layout = QHBoxLayout()
        
        # Back button
        back_btn = QPushButton()
        back_btn.setIcon(self.style().standardIcon(QtWidgets.QStyle.SP_ArrowBack))
        back_btn.setFixedSize(36, 36)
        back_btn.setCursor(Qt.PointingHandCursor)
        back_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {PRIMARY_COLOR};
                border-radius: 18px;
                color: white;
            }}
            QPushButton:hover {{
                background-color: #1f2c39;
            }}
        """)
        back_btn.clicked.connect(self.go_back_to_dashboard)
        
        page_title = QLabel("BDU Group")
        page_title.setFont(QFont("Segoe UI", 22, QFont.Bold))
        page_title.setStyleSheet(f"color: {PRIMARY_COLOR};")
        
        # Get BDU icon from departments
        bdu_dept = next((dept for dept in DEPARTMENTS if dept["id"] == "BDU"), None)
        if bdu_dept:
            dept_icon = QLabel(bdu_dept["emoji"])
            dept_icon.setFont(QFont("Segoe UI", 22))
            dept_icon.setStyleSheet(f"color: {bdu_dept['color']};")
            title_layout.addWidget(dept_icon)
        
        title_layout.addWidget(back_btn)
        title_layout.addSpacing(10)
        title_layout.addWidget(page_title)
        title_layout.addStretch()
        
        # User Code Dropdown
        user_code_label = QLabel("User Code:")
        user_code_label.setFont(QFont("Segoe UI", 10))
        user_code_label.setStyleSheet("color: #666;")
        
        self.user_code_dropdown = QComboBox()
        self.user_code_dropdown.setFont(QFont("Segoe UI", 10))
        self.user_code_dropdown.setMinimumWidth(100)
        self.user_code_dropdown.setStyleSheet(f"""
            QComboBox {{
                padding: 5px;
                border: 1px solid #ccc;
                border-radius: 4px;
                background-color: white;
                min-height: 25px;
            }}
            QComboBox:hover {{
                border: 1px solid {SECONDARY_COLOR};
            }}
            QComboBox::drop-down {{
                subcontrol-origin: padding;
                subcontrol-position: top right;
                width: 20px;
                border-left-width: 1px;
                border-left-color: #ccc;
                border-left-style: solid;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
            }}
        """)
        self.user_code_dropdown.currentTextChanged.connect(self.on_user_code_changed)
        
        # Add User Code dropdown to title layout
        title_layout.addWidget(user_code_label)
        title_layout.addWidget(self.user_code_dropdown)
        title_layout.addSpacing(15)
        
        # Add refresh button
        refresh_btn = QPushButton("Refresh Data")
        refresh_btn.setFont(QFont("Segoe UI", 10))
        refresh_btn.setCursor(Qt.PointingHandCursor)
        refresh_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {SECONDARY_COLOR};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 6px 12px;
            }}
            QPushButton:hover {{
                background-color: #2980B9;
            }}
        """)
        refresh_btn.clicked.connect(self.refresh_with_calculation)
        
        title_layout.addWidget(refresh_btn)
        
        content_layout.addLayout(title_layout)
        
        # Tab widget for different sheets
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet(f"""
            QTabWidget::pane {{
                border: 1px solid #cccccc;
                background-color: white;
                border-radius: 5px;
            }}
            QTabBar::tab {{
                background-color: #f0f0f0;
                color: #333333;
                padding: 8px 15px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                border: 1px solid #cccccc;
                border-bottom: none;
            }}
            QTabBar::tab:selected {{
                background-color: white;
                border-bottom-color: white;
            }}
            QTabBar::tab:hover {{
                background-color: #e0e0e0;
            }}
        """)
        
        content_layout.addWidget(self.tab_widget)
        
        # Loading message
        self.loading_label = QLabel("Loading data from SET_BDU.xlsx...")
        self.loading_label.setFont(QFont("Segoe UI", 12))
        self.loading_label.setAlignment(Qt.AlignCenter)
        self.loading_label.setStyleSheet("color: #666; margin: 20px;")
        
        content_layout.addWidget(self.loading_label)
        
        # Set the main content to the scroll area
        content_widget.setWidget(self.main_content)
        main_layout.addWidget(content_widget)
        
        # Status bar
        self.statusBar().showMessage(f"BDU Group Module | User: {self.current_user['username']}")
        self.statusBar().setStyleSheet("background-color: #f0f0f0; color: #555;")
    
    def load_user_codes_from_excel(self):
        """Load user codes from 'User Code' sheet in Excel - FIXED VERSION"""
        try:
            if not os.path.exists(self.excel_path):
                print(f"Excel file not found: {self.excel_path}")
                return
            
            # SOLUSI: Specify dtype untuk kolom Code sebagai string
            # Ini akan mempertahankan leading zeros
            dtype_dict = {'Code': str}  # Force Code column to be read as string
            
            # Read the 'User Code' sheet dengan dtype specification
            df = pd.read_excel(self.excel_path, sheet_name='User Code', dtype=dtype_dict)
            
            # Check if 'Code' column exists
            if 'Code' in df.columns:
                # Get all codes, filter out NaN values - sudah dalam format string
                codes = df['Code'].dropna().tolist()
                
                # Filter out empty strings dan pastikan format string
                self.user_codes = []
                for code in codes:
                    code_str = str(code).strip()
                    if code_str and code_str.lower() != 'nan':
                        self.user_codes.append(code_str)
                
                # Update dropdown
                self.update_user_code_dropdown()
            else:
                print("'Code' column not found in 'User Code' sheet")
                print(f"Available columns: {df.columns.tolist()}")
                self.user_codes = []
                
        except Exception as e:
            print(f"Error loading user codes from Excel: {str(e)}")
            import traceback
            traceback.print_exc()
            self.user_codes = []

    def update_user_code_dropdown(self):
        """Update the user code dropdown with loaded codes"""
        self.user_code_dropdown.clear()
        
        if self.user_codes:
            # Add placeholder option
            self.user_code_dropdown.addItem("-- Select User Code --")
            
            # Add all user codes
            for code in self.user_codes:
                self.user_code_dropdown.addItem(code)
            
            # Style the placeholder item
            self.user_code_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            self.user_code_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            self.user_code_dropdown.setCurrentIndex(0)
            
        else:
            # No codes available
            self.user_code_dropdown.addItem("-- No User Codes Available --")
            self.user_code_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            self.user_code_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            self.user_code_dropdown.setEnabled(False)

    def on_user_code_changed(self, selected_code):
        """Handle user code selection change - UPDATED VERSION"""
        if selected_code and selected_code != "-- Select User Code --" and selected_code != "-- No User Codes Available --":
            
            # Simpan user code yang dipilih ke Excel
            self.save_selected_user_code_to_excel(selected_code)
            
            # Update status bar to show selected user code
            self.statusBar().showMessage(f"BDU Group Module | User: {self.current_user['username']} | Selected Code: {selected_code}")
            
            # Get user details for selected code
            user_details = self.get_user_details_by_code(selected_code)
            if user_details:
                print(f"User details: {user_details}")
                
                # Show user info in status or somewhere in UI
                user_info = f"Selected: {user_details.get('Name', 'Unknown')} ({user_details.get('Position', 'Unknown')})"
                print(f"User info: {user_info}")
            
        else:
            # Reset status bar if no valid code selected
            self.statusBar().showMessage(f"BDU Group Module | User: {self.current_user['username']}")

    def save_selected_user_code_to_excel(self, user_code):
        """Simpan user code yang dipilih ke Excel DATA_TEMP.B1 untuk digunakan saat generate proposal"""
        try:
            from openpyxl import load_workbook
            
            # Load workbook
            wb = load_workbook(self.excel_path)
            
            # Cek apakah sheet DATA_TEMP ada
            if 'DATA_TEMP' not in wb.sheetnames:
                print(f"Warning: Sheet DATA_TEMP tidak ditemukan, membuat sheet baru")
                # Buat sheet baru jika tidak ada
                sheet = wb.create_sheet('DATA_TEMP')
            else:
                sheet = wb['DATA_TEMP']
            
            # Simpan user code di cell B1
            sheet['B1'] = user_code
            
            # Simpan workbook
            wb.save(self.excel_path)
            wb.close()
            
            return True
            
        except Exception as e:
            print(f"Error saat menyimpan user code ke Excel: {str(e)}")
            return False
    
    def get_user_details_by_code(self, selected_code):
        """Get user details from Excel for the selected code"""
        try:
            if not os.path.exists(self.excel_path):
                return None
            
            # Read the 'User Code' sheet
            df = pd.read_excel(self.excel_path, sheet_name='User Code')
            
            # Find the row with matching code
            user_row = df[df['Code'].astype(str).str.strip() == selected_code]
            
            if not user_row.empty:
                # Convert the row to dictionary and return
                user_details = user_row.iloc[0].to_dict()
                
                # Clean up NaN values
                for key, value in user_details.items():
                    if pd.isna(value):
                        user_details[key] = ""
                
                return user_details
            
            return None
            
        except Exception as e:
            print(f"Error getting user details for code {selected_code}: {str(e)}")
            return None
    
    def setup_header(self):
        """Setup header widget dengan logo, judul, dan menu"""
        self.header_widget = QWidget()
        self.header_widget.setFixedHeight(60)
        self.header_widget.setStyleSheet(f"background-color: {PRIMARY_COLOR};")
        
        header_layout = QHBoxLayout(self.header_widget)
        header_layout.setContentsMargins(15, 0, 15, 0)
        
        # BDU title with icon
        bdu_dept = next((dept for dept in DEPARTMENTS if dept["id"] == "BDU"), None)
        if bdu_dept:
            dept_icon = QLabel(bdu_dept["emoji"])
            dept_icon.setFont(QFont("Segoe UI", 18))
            dept_icon.setStyleSheet("color: white;")
            header_layout.addWidget(dept_icon)
        
        # App title
        title_label = QLabel(f"{APP_NAME} - BDU Group")
        title_label.setFont(QFont("Segoe UI", 16, QFont.Bold))
        title_label.setStyleSheet("color: white;")
        
        # Spacer
        spacer = QSpacerItem(40, 20, QSizePolicy.Expanding, QSizePolicy.Minimum)
        
        # User button (menggunakan teks emoji)
        user_menu_btn = QPushButton("👤")
        user_menu_btn.setFont(QFont("Segoe UI", 14))
        user_menu_btn.setStyleSheet("""
            QPushButton {
                border: none;
                background-color: transparent;
                color: white;
            }
            QPushButton:hover {
                background-color: rgba(255, 255, 255, 0.2);
                border-radius: 5px;
            }
        """)
        user_menu_btn.setFixedSize(36, 36)
        user_menu_btn.setCursor(Qt.PointingHandCursor)
        user_menu_btn.clicked.connect(self.show_user_menu)
        
        # Add all elements to header layout
        header_layout.addWidget(title_label)
        header_layout.addItem(spacer)
        header_layout.addWidget(user_menu_btn)
    
    def show_user_menu(self):
        """Tampilkan menu pengguna"""
        sender = self.sender()
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background-color: white;
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 5px;
            }
            QMenu::item {
                padding: 5px 15px;
                border-radius: 3px;
            }
            QMenu::item:selected {
                background-color: #f0f0f0;
            }
        """)
        
        # Add user info at the top (non-clickable)
        user_info = QAction(f"{self.current_user['name']} ({self.current_user['department']})", self)
        user_info.setEnabled(False)
        menu.addAction(user_info)
        
        menu.addSeparator()
        
        # Add menu actions
        dashboard_action = QAction("Back to Dashboard", self)
        settings_action = QAction("Settings", self)
        
        menu.addAction(dashboard_action)
        menu.addAction(settings_action)
        
        # Connect actions
        dashboard_action.triggered.connect(self.go_back_to_dashboard)
        
        # Show menu at button position
        menu.exec_(sender.mapToGlobal(QPoint(0, sender.height())))
    
    def go_back_to_dashboard(self):
        """Kembali ke dashboard"""
        self.back_to_dashboard.emit()
        self.close()
    
    def update_dependent_dropdown(self, parent_value, child_dropdown):
        """Update dropdown yang bergantung berdasarkan nilai dropdown utama"""
        # Bersihkan dropdown anak
        child_dropdown.clear()
        
        # Simpan nilai saat ini agar bisa digunakan nanti
        current_value = getattr(child_dropdown, "_last_value", None)
        
        # Cek jenis dropdown berdasarkan property
        dropdown_type = child_dropdown.property("pump_dropdown_type")
        
        # Isi dengan nilai-nilai yang sesuai berdasarkan pilihan pada dropdown utama
        if parent_value in INDUSTRY_SUBTYPE_MAPPING:
            # Ini adalah dropdown industry -> sub industry
            child_dropdown.addItem("-- Select Value --")
            child_dropdown.addItems(INDUSTRY_SUBTYPE_MAPPING[parent_value])
            child_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            child_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
        elif parent_value in INDONESIA_CITIES:
            # Ini adalah dropdown province -> city
            child_dropdown.addItem("-- Select Value --")
            child_dropdown.addItems(INDONESIA_CITIES[parent_value])
            child_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            child_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
        elif dropdown_type == "pump_type" and parent_value in PUMP_BRAND_TYPE_MAPPING:
            # Ini adalah dropdown pump brand -> pump type
            child_dropdown.addItem("-- Select Value --")
            child_dropdown.addItems(PUMP_BRAND_TYPE_MAPPING[parent_value])
            child_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            child_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            child_dropdown.setEnabled(True)
        elif dropdown_type == "pump_type" and parent_value not in PUMP_BRAND_TYPE_MAPPING:
            # Brand yang tidak memiliki turunan - disable type dropdown
            child_dropdown.addItem("-- Not Required for this Brand --")
            child_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            child_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            child_dropdown.setEnabled(False)
        else:
            # Default jika tidak ada mapping
            child_dropdown.addItem("-- Select Value --")
            child_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            child_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                    
    def update_pump_model_dropdown(self, pump_brand, pump_type, model_dropdown):
        """Update pump model dropdown berdasarkan brand dan type yang dipilih"""
        model_dropdown.clear()
        
        # Cek apakah brand memiliki turunan
        if pump_brand not in PUMP_BRAND_TYPE_MAPPING:
            # Brand tidak memiliki turunan - disable model dropdown
            model_dropdown.addItem("-- Not Required for this Brand --")
            model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            model_dropdown.setEnabled(False)
            return
        
        # Skip jika nilai masih placeholder
        if (pump_brand == "-- Select Value --" or pump_type == "-- Select Value --" or
            pump_type == "-- Not Required for this Brand --" or
            not pump_brand or not pump_type):
            model_dropdown.addItem("-- Select Pump Brand and Type First --")
            model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            model_dropdown.setEnabled(True)
            return
        
        model_key = (pump_brand, pump_type)
        
        if model_key in PUMP_BRAND_TYPE_MODEL_MAPPING:
            model_dropdown.addItem("-- Select Value --")
            model_dropdown.addItems(PUMP_BRAND_TYPE_MODEL_MAPPING[model_key])
            model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            model_dropdown.setEnabled(True)
        else:
            # Jika tidak ada mapping spesifik untuk brand+type ini
            model_dropdown.addItem("-- No Models Available --")
            model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
            model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
            model_dropdown.setEnabled(False)
    
    def get_absolute_path(self, relative_path):
        """Konversi path relatif menjadi absolut relatif terhadap root project"""
        if os.path.isabs(relative_path):
            return relative_path
        
        project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        return os.path.join(project_root, relative_path)
    
    def refresh_with_calculation(self):
        """Refresh data and process formulas in background"""
        
        def refresh_process(progress_callback=None):
            try:
                if progress_callback:
                    progress_callback(20, "Refreshing Excel data...")
                
                # Reload Excel data
                self.load_excel_data()
                
                if progress_callback:
                    progress_callback(60, "Processing formulas...")
                
                # Process formulas in background
                if hasattr(self, 'formula_evaluator') and self.formula_evaluator and HAS_FORMULA_HELPER:
                    evaluate_formulas_background(self.formula_evaluator)
                
                if progress_callback:
                    progress_callback(100, "Refresh completed!")
                
                return "Data refreshed successfully with background formula processing"
                    
            except Exception as e:
                if progress_callback:
                    progress_callback(100, f"Error: {str(e)}")
                return f"Error during refresh: {str(e)}"
        
        # Show loading screen
        loading_screen = LoadingScreen(
            parent=self,
            title="Refreshing Data",
            message="Reloading data and processing formulas..."
        )
        loading_screen.show()
        loading_screen.start_loading(refresh_process)
        
        # Connect completion handler
        def on_refresh_complete(success, message):
            if success:
                self.statusBar().showMessage("Data refreshed successfully", 3000)
            else:
                QMessageBox.critical(self, "Error", f"Error during refresh: {message}")
                self.statusBar().clearMessage()
        
        loading_screen.worker.task_completed.connect(on_refresh_complete)    
    
    def load_excel_data(self):
        """Load data from SET_BDU.xlsx"""
        
        def load_data_process(progress_callback=None):
            try:
                if progress_callback:
                    progress_callback(5, "Checking Excel file...")
                
                if not os.path.exists(self.excel_path):
                    return f"Error: File SET_BDU.xlsx not found in the data directory."
                
                if progress_callback:
                    progress_callback(10, "Loading user codes...")

                # Load user codes first
                self.load_user_codes_from_excel()

                if progress_callback:
                    progress_callback(15, "Preparing data loading...")
                
                # Initialize formula evaluator
                if HAS_FORMULA_HELPER:
                    self.formula_evaluator = SimpleFormulaEvaluator(self.excel_path)
                    if self.formula_evaluator.load_workbook():
                        # TAMBAHKAN INI - Process formulas di background
                        if progress_callback:
                            progress_callback(20, "Processing formulas in background...")
                        evaluate_formulas_background(self.formula_evaluator)
                    else:
                        print("Warning: Could not initialize formula evaluator")
                        self.formula_evaluator = None
                 
                if progress_callback:
                    progress_callback(35, "Reading Excel file structure...")
                
                # Clear existing tabs
                self.tab_widget.clear()
                self.sheet_tabs = {}
                self.data_fields = {}
                
                # Hide loading message when tabs exist
                self.loading_label.setVisible(True)
                
                # Read Excel file
                xl = pd.ExcelFile(self.excel_path)
                sheet_names = xl.sheet_names
                
                if progress_callback:
                    progress_callback(45, "Filtering relevant sheets...")
                
                # Filter sheets to only include those starting with DATA_ or DIP_
                filtered_sheets = [sheet for sheet in sheet_names if sheet.startswith("DATA_") or sheet.startswith("DIP_")]
                
                if len(filtered_sheets) == 0:
                    return "No DATA_ or DIP_ sheets found in SET_BDU.xlsx."
                
                # Hide loading label as we have data
                self.loading_label.setVisible(False)
                
                if progress_callback:
                    progress_callback(55, f"Processing {len(filtered_sheets)} sheets...")
                
                # Create a tab for each filtered sheet
                sheet_progress_step = 40 / len(filtered_sheets)  # Distribute 40% progress among sheets
                
                for i, sheet_name in enumerate(filtered_sheets):
                    current_progress = 55 + (i * sheet_progress_step)
                    
                    if progress_callback:
                        progress_callback(int(current_progress), f"Processing sheet: {sheet_name}")
                    
                    # Get display name
                    display_name = sheet_name
                    if sheet_name.startswith("DIP_"):
                        display_name = sheet_name[4:]
                    elif sheet_name.startswith("DATA_"):
                        display_name = sheet_name[5:]
                    
                    try:
                        df = pd.read_excel(self.excel_path, sheet_name=sheet_name, header=None)
                        
                        if sheet_name == "DATA_PROPOSAL":
                            # Special handling for proposal sheet
                            scroll_area = QScrollArea()
                            scroll_area.setWidgetResizable(True)
                            
                            proposal_widget = QWidget()
                            proposal_layout = QVBoxLayout(proposal_widget)
                            proposal_layout.setContentsMargins(10, 10, 10, 0)
                            proposal_layout.setSpacing(0)
                            
                            # Check for existing proposal file
                            file_exists = False
                            relative_word_file_path = ""
                            
                            if not df.empty and not pd.isna(df.iloc[0, 0]):
                                relative_word_file_path = str(df.iloc[0, 0]).strip()
                                absolute_word_file_path = self.get_absolute_path(relative_word_file_path)
                                file_exists = os.path.exists(absolute_word_file_path)
                            
                            if file_exists:
                                self.proposal_relative_path = relative_word_file_path
                                if self.process_proposal_document(absolute_word_file_path):
                                    if hasattr(self, 'proposal_document_widget'):
                                        proposal_layout.addWidget(self.proposal_document_widget)
                            else:
                                # Create proposal interface for when file doesn't exist
                                self.create_proposal_interface(proposal_layout)
                            
                            scroll_area.setWidget(proposal_widget)
                            self.tab_widget.addTab(scroll_area, display_name)
                            self.sheet_tabs[sheet_name] = proposal_widget
                        else:
                            # Regular processing for other sheets
                            scroll_area = QScrollArea()
                            scroll_area.setWidgetResizable(True)
                            
                            sheet_widget = QWidget()
                            sheet_layout = QVBoxLayout(sheet_widget)
                            sheet_layout.setContentsMargins(15, 15, 15, 15)
                            
                            # Process the sheet data
                            self.process_sheet_data(df, sheet_name, sheet_layout)
                            
                            scroll_area.setWidget(sheet_widget)
                            self.tab_widget.addTab(scroll_area, display_name)
                            self.sheet_tabs[sheet_name] = sheet_widget
                            
                    except Exception as e:
                        print(f"Error processing sheet {sheet_name}: {str(e)}")
                        # Create an error tab for this sheet
                        error_widget = QWidget()
                        error_layout = QVBoxLayout(error_widget)
                        
                        error_label = QLabel(f"Error loading {sheet_name}: {str(e)}")
                        error_label.setStyleSheet("color: #E74C3C;")
                        error_layout.addWidget(error_label)
                        
                        self.tab_widget.addTab(error_widget, display_name)
                
                if progress_callback:
                    progress_callback(100, "Excel data loaded successfully!")
                
                return "Excel data loaded successfully"
                
            except Exception as e:
                if progress_callback:
                    progress_callback(100, f"Error: {str(e)}")
                return f"Error loading data: {str(e)}"
        
        # For large files, show loading screen. For small files, load directly.
        try:
            file_size = os.path.getsize(self.excel_path) / (1024 * 1024)  # Size in MB
            
            if file_size > 5:  # Show loading screen for files larger than 5MB
                loading_screen = LoadingScreen(
                    parent=self,
                    title="Loading Excel Data",
                    message="Reading and processing Excel sheets..."
                )
                loading_screen.show()
                loading_screen.start_loading(load_data_process)
                
                def on_load_complete(success, message):
                    if not success:
                        self.loading_label.setText(message)
                        self.loading_label.setStyleSheet("color: #E74C3C; margin: 20px;")
                        self.loading_label.setVisible(True)
                        print(f"Error loading Excel data: {message}")
                
                loading_screen.worker.task_completed.connect(on_load_complete)
            else:
                # Load directly for smaller files
                result = load_data_process()
                if "Error" in result:
                    self.loading_label.setText(result)
                    self.loading_label.setStyleSheet("color: #E74C3C; margin: 20px;")
                    self.loading_label.setVisible(True)
                    print(f"Error loading Excel data: {result}")
        except Exception as e:
            self.loading_label.setText(f"Error loading data: {str(e)}")
            self.loading_label.setStyleSheet("color: #E74C3C; margin: 20px;")
            self.loading_label.setVisible(True)
            print(f"Error loading Excel data: {str(e)}") 
            
    def create_proposal_interface(self, layout):
        """Create interface for when proposal file doesn't exist"""
        # Container untuk header file yang tidak ada
        top_container = QWidget()
        top_layout = QHBoxLayout(top_container)
        top_layout.setContentsMargins(0, 0, 0, 0)
        
        # Panel kiri untuk ikon dan label
        left_panel = QWidget()
        left_panel.setMaximumWidth(300)
        left_layout = QHBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(5)
        
        # Ikon file
        file_icon = QLabel("📄")
        file_icon.setFont(QFont("Segoe UI", 14))
        file_icon.setStyleSheet("color: #3498DB; background-color: transparent;")
        left_layout.addWidget(file_icon)
        
        # Label status proposal
        file_label = QLabel("Proposal Not Generated")
        file_label.setFont(QFont("Segoe UI", 11))
        file_label.setStyleSheet("color: #333; background-color: transparent;")
        left_layout.addWidget(file_label)
        left_layout.addStretch()
        
        # Panel kanan untuk tombol
        right_panel = QWidget()
        right_layout = QHBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setAlignment(Qt.AlignRight | Qt.AlignTop)
        
        # Tombol Run Projection
        run_projection_btn = QPushButton("Run Projection")
        run_projection_btn.setFont(QFont("Segoe UI", 10))
        run_projection_btn.setCursor(Qt.PointingHandCursor)
        run_projection_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 5px 10px;
                margin-right: 5px;
            }}
            QPushButton:hover {{
                background-color: #2980b9;
            }}
        """)
        run_projection_btn.clicked.connect(self.run_projection)
        right_layout.addWidget(run_projection_btn)
        
        # Tombol Generate Proposal
        generate_btn = QPushButton("Generate Proposal")
        generate_btn.setFont(QFont("Segoe UI", 10))
        generate_btn.setCursor(Qt.PointingHandCursor)
        generate_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: #27ae60;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 5px 10px;
            }}
            QPushButton:hover {{
                background-color: #2ecc71;
            }}
        """)
        generate_btn.clicked.connect(self.run_generate_proposal)
        right_layout.addWidget(generate_btn)
        
        # Menyusun panel kiri dan kanan
        top_layout.addWidget(left_panel)
        top_layout.addWidget(right_panel)
        
        # Tambahkan container utama ke layout proposal
        layout.addWidget(top_container)
        
        # Tambahkan instruksi yang singkat
        info_label = QLabel("Click 'Generate Proposal' button to create a new proposal document.")
        info_label.setAlignment(Qt.AlignCenter)
        info_label.setFont(QFont("Segoe UI", 10))
        info_label.setStyleSheet("color: #666;")
        info_label.setMaximumHeight(20)
        layout.addWidget(info_label)
        
        # Tambahkan spacer
        layout.addStretch(1)   
                
    def force_excel_calculation(self, excel_path, progress_callback=None):
        """Force Excel to recalculate all formulas before reading data with progress updates"""
        try:
            import xlwings as xw
            import time
            
            if progress_callback:
                progress_callback(10, "Starting Excel application...")
            
            # Buka Excel dengan xlwings (tidak visible)
            app = xw.App(visible=False)
            app.display_alerts = False
            
            if progress_callback:
                progress_callback(30, "Opening workbook...")
            
            # Buka workbook
            wb = xw.Book(excel_path)
            
            if progress_callback:
                progress_callback(50, "Forcing formula calculation...")
            
            # Paksa kalkulasi semua formula
            wb.api.Application.CalculateFullRebuild()
            wb.api.Application.Calculate()
            
            if progress_callback:
                progress_callback(70, "Waiting for calculation to complete...")
            
            # Tunggu sebentar untuk memastikan kalkulasi selesai
            time.sleep(2)
            
            if progress_callback:
                progress_callback(90, "Saving workbook...")
            
            # Simpan workbook
            wb.save()
            
            if progress_callback:
                progress_callback(100, "Calculation completed!")
            
            # Tutup
            wb.close()
            app.quit()
            
            return True
            
        except ImportError:
            # Jika xlwings tidak tersedia, gunakan metode alternatif dengan VBScript
            return self.force_calculation_vbs(excel_path, progress_callback)
        except Exception as e:
            print(f"Error forcing Excel calculation: {str(e)}")
            if progress_callback:
                progress_callback(100, f"Error: {str(e)}")
            return False
    
    def run_projection(self):
        """Fungsi untuk menjalankan projection dengan loading screen - Complete Version with Process 4 & 5"""
        
        def projection_process(progress_callback=None):
            try:
                import os
                import pandas as pd
                from openpyxl import load_workbook
                import time
                import sys
                import subprocess
                from PyQt5.QtWidgets import QApplication, QMessageBox
                
                print("=" * 80)
                print("🚀 STARTING COMPLETE PROJECTION PROCESS")
                print("=" * 80)
                
                if progress_callback:
                    progress_callback(2, "Initializing projection process...")
                
                # Use the customer-specific Excel file if available
                if hasattr(self, 'excel_path'):
                    set_bdu_path = self.excel_path
                    print(f"📁 Using customer-specific Excel: {set_bdu_path}")
                else:
                    # Fall back to default path if customer file isn't set
                    data_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
                    set_bdu_path = os.path.join(data_folder, "SET_BDU.xlsx")
                    print(f"📁 Using default Excel: {set_bdu_path}")
                
                # Base data folder path
                data_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
                print(f"📂 Data folder: {data_folder}")
                
                # Path ke file-file Excel (UPDATED PATHS)
                sbt_anapak_path = os.path.join(data_folder, "SBT_PROCESS", "SBT_ANAPAK.xlsx")
                sbt_instrument_path = os.path.join(data_folder, "SBT_EQUIPMENT AND TOOLS", "SBT-INSTRUMENT versi 1.0.xlsm")
                sbt_dosingpump_path = os.path.join(data_folder,  "SBT_EQUIPMENT AND TOOLS", "SBT-DOSINGPUMP versi 1.0.xlsm")
                all_udf_path = os.path.join(data_folder, "ALL_UDF.py")
                
                print(f"📄 SBT_ANAPAK path: {sbt_anapak_path}")
                print(f"📄 SBT_INSTRUMENT path: {sbt_instrument_path}")
                print(f"📄 SBT_DOSINGPUMP path: {sbt_dosingpump_path}")
                
                if progress_callback:
                    progress_callback(5, "Checking required files...")
                
                print("\n🔍 CHECKING REQUIRED FILES:")
                # Pastikan file yang dibutuhkan ada
                missing_files = []
                required_files = [set_bdu_path, sbt_anapak_path, sbt_instrument_path, all_udf_path]
                for file_path in required_files:
                    if not os.path.exists(file_path):
                        missing_files.append(os.path.basename(file_path))
                        print(f"❌ Missing: {file_path}")
                    else:
                        print(f"✅ Found: {file_path}")
                
                if missing_files:
                    error_msg = f"Required files not found: {', '.join(missing_files)}"
                    print(f"🚨 ERROR: {error_msg}")
                    raise Exception(error_msg)
                
                print("✅ All required files found!")
                
                if progress_callback:
                    progress_callback(8, "Opening SET_BDU workbook...")
                
                print("\n📖 PROSES 1: READING SET_BDU DATA")
                print("-" * 50)
                
                # PROSES 1: Transfer data dari SET_BDU ke SBT_ANAPAK
                print(f"🔓 Opening SET_BDU: {set_bdu_path}")
                wb_bdu = load_workbook(set_bdu_path, data_only=True)
                print(f"📋 Available sheets: {wb_bdu.sheetnames}")
                
                if progress_callback:
                    progress_callback(10, "Reading data from DIP sheets...")
                
                # Ambil data dari DIP_Customer Information dan DIP_Project Information
                print("📄 Reading DIP_Customer Information...")
                if "DIP_Customer Information" not in wb_bdu.sheetnames:
                    raise Exception("Sheet 'DIP_Customer Information' not found!")
                sheet_customer = wb_bdu["DIP_Customer Information"]
                
                print("📄 Reading DIP_Project Information...")
                if "DIP_Project Information" not in wb_bdu.sheetnames:
                    raise Exception("Sheet 'DIP_Project Information' not found!")
                sheet_project = wb_bdu["DIP_Project Information"]
                
                # Data untuk SBT_ANAPAK (UPDATED MAPPING - combining old and new)
                print("📊 Extracting data for ANAPAK...")
                anapak_values = {
                    'C1': sheet_customer['B4'].value,   # DIP_Customer Information.B4
                    'C4': sheet_project['B70'].value,   # DIP_Project Information.B70
                    'C5': sheet_project['B71'].value,   # DIP_Project Information.B71
                    'C6': sheet_project['B72'].value,   # DIP_Project Information.B72
                    'C7': sheet_project['B73'].value,   # DIP_Project Information.B73
                    'C8': sheet_project['B74'].value,   # DIP_Project Information.B74
                    'C9': sheet_project['B77'].value,   # DIP_Project Information.B77
                    'C10': sheet_project['B78'].value,  # DIP_Project Information.B78
                    'C11': sheet_project['B82'].value,  # DIP_Project Information.B82
                    'C16': sheet_project['B42'].value,  # DIP_Project Information.B42
                    'C17': sheet_project['B43'].value,  # DIP_Project Information.B43
                    'C18': sheet_project['B45'].value,  # DIP_Project Information.B45
                    'C19': sheet_project['B51'].value,  # DIP_Project Information.B51
                    'C20': sheet_project['B59'].value   # DIP_Project Information.B59
                }
                
                print("📋 ANAPAK Data extracted:")
                for key, value in anapak_values.items():
                    print(f"   {key}: {value}")
                
                # Ambil data untuk pemilihan pump file
                pump_brand = sheet_project['B42'].value
                pump_type = sheet_project['B43'].value
                pump_model = sheet_project['B44'].value
                
                print(f"🔧 Pump Configuration:")
                print(f"   Brand: {pump_brand}")
                print(f"   Type: {pump_type}")
                print(f"   Model: {pump_model}")
                
                # Ambil data untuk Proses 4 & 5
                project_b59_value = sheet_project['B59'].value  # Untuk SBT_INSTRUMENT
                project_b42_value = sheet_project['B42'].value  # Untuk Proses 3.4
                
                print(f"📊 Additional data for Process 4 & 5:")
                print(f"   B59 (for INSTRUMENT): {project_b59_value}")
                print(f"   B42 (for Process 3.4): {project_b42_value}")
                
                wb_bdu.close()
                print("✅ SET_BDU closed successfully")
                
                if progress_callback:
                    progress_callback(15, "Transferring data to SBT_ANAPAK...")
                
                print("\n📤 TRANSFERRING DATA TO SBT_ANAPAK")
                print("-" * 50)
                
                # Buka dan update workbook SBT_ANAPAK
                print(f"🔓 Opening SBT_ANAPAK: {sbt_anapak_path}")
                wb_anapak = load_workbook(sbt_anapak_path)
                print(f"📋 Available sheets: {wb_anapak.sheetnames}")
                
                if "DATA_INPUT" not in wb_anapak.sheetnames:
                    raise Exception("Sheet 'DATA_INPUT' not found in SBT_ANAPAK!")
                
                sheet_anapak_input = wb_anapak["DATA_INPUT"]
                print("📄 Found DATA_INPUT sheet")
                
                # Pindahkan data ke SBT_ANAPAK
                print("📊 Writing data to ANAPAK...")
                for cell_addr, value in anapak_values.items():
                    old_value = sheet_anapak_input[cell_addr].value
                    sheet_anapak_input[cell_addr] = value
                    print(f"   {cell_addr}: {old_value} → {value}")
                
                # Simpan SBT_ANAPAK
                print("💾 Saving SBT_ANAPAK...")
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ SBT_ANAPAK saved and closed successfully")
                
                if progress_callback:
                    progress_callback(20, "Force calculating SBT_ANAPAK formulas...")
                
                print("\n🧮 FORCE CALCULATING SBT_ANAPAK FORMULAS")
                print("-" * 50)
                
                try:
                    print("🔄 Starting force calculation for SBT_ANAPAK...")
                    self.force_excel_calculation(sbt_anapak_path, 
                        lambda pct, msg: progress_callback(20 + (pct * 0.05), f"SBT_ANAPAK: {msg}") if progress_callback else None)
                    print("✅ SBT_ANAPAK calculation completed")
                    time.sleep(3)  # Give more time for calculation to stabilize
                except Exception as e:
                    print(f"⚠️ Warning during SBT_ANAPAK calculation: {str(e)}")
                    if progress_callback:
                        progress_callback(25, f"Warning: {str(e)}")
                
                if progress_callback:
                    progress_callback(25, "Reading calculated values from SBT_ANAPAK...")
                
                print("\n📊 READING CALCULATED VALUES FROM SBT_ANAPAK")
                print("-" * 50)
                
                # CRITICAL: Use OLD METHOD - data_only=True to read calculated values
                print("🔓 Reading data from SBT_ANAPAK with data_only=True...")
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)  # KEY: data_only=True
                print(f"📋 Available sheets: {wb_anapak_read.sheetnames}")
                
                if "DATA_ENGINE ANAPAK" not in wb_anapak_read.sheetnames:
                    print("❌ Sheet 'DATA_ENGINE ANAPAK' not found!")
                    print(f"📋 Available sheets: {wb_anapak_read.sheetnames}")
                    raise Exception("Sheet 'DATA_ENGINE ANAPAK' not found!")
                
                sheet_anapak_engine = wb_anapak_read["DATA_ENGINE ANAPAK"]
                
                # Ambil nilai dari SBT_ANAPAK.DATA_ENGINE ANAPAK
                value_i66 = sheet_anapak_engine['I66'].value
                value_k67 = sheet_anapak_engine['K67'].value
                
                print(f"📊 Retrieved from ANAPAK (with data_only=True):")
                print(f"   I66: {value_i66}")
                print(f"   K67: {value_k67}")
                
                wb_anapak_read.close()
                print("✅ SBT_ANAPAK closed")
                
                # Gunakan nilai default jika masih None atau formula
                if value_i66 is None or (isinstance(value_i66, str) and value_i66.startswith('=')):
                    print("⚠️ I66 was None or formula, using default: 100")
                    value_i66 = 100
                if value_k67 is None or (isinstance(value_k67, str) and value_k67.startswith('=')):
                    print("⚠️ K67 was None or formula, using default: 24.44")
                    value_k67 = 24.44
                
                # Ensure numeric values
                try:
                    value_i66 = float(value_i66)
                    value_k67 = float(value_k67)
                except:
                    value_i66 = 100.0
                    value_k67 = 24.44
                
                print(f"✅ Final values to use:")
                print(f"   I66: {value_i66}")
                print(f"   K67: {value_k67}")
                
                if progress_callback:
                    progress_callback(30, "Determining SBT_PUMP file...")
                
                print("\n🔧 PROSES 2: DETERMINING PUMP FILE")
                print("-" * 50)
                
                # PROSES 2: Menentukan file SBT_PUMP berdasarkan brand, type, dan model
                pump_file_map = {
                    ("GRUNDFOS", "Vertical Multistage Centrifugal Pump", "CR"): "SBT-GRUNDFOS-CR Pump versi 2.0 (test version 3.51).xlsm",
                    ("GRUNDFOS", "Vertical Multistage Centrifugal Pump", "CRN"): "SBT-GRUNDFOS-CRN Pump versi 2.0 (test version 3.51).xlsm",
                    ("GRUNDFOS", "End Suction Centrifugal Pump", "NKG"): "SBT-GRUNDFOS-NKG Pump versi 2.0 (test version 3.51).xlsm",
                    ("EBARA", "Vertical Multistage Centrifugal Pump", "3S"): "SBT-EBARA-3S Pump versi 2.0 (test version 3.51).xlsm",
                    ("EBARA", "End Suction Centrifugal Pump", "FSSC"): "SBT-EBARA-FSSC Pump versi 2.0 (test version 3.51).xlsm",
                    ("CNP", "Vertical Multistage Centrifugal Pump", "CDMF"): "SBT-CNP-CDMF Pump versi 2.0 (test version 3.51).xlsm",
                    ("CNP", "Horizontal Multistage Centrifugal Pump", "CHL"): "SBT-CNP-CHL Pump versi 2.0 (test version 3.51).xlsm",
                    ("CNP", "Submersible Pump", "WQ"): "SBT-CNP-WQ Pump versi 2.0 (test version 3.51).xlsm",
                    ("LEO", "Vertical Multistage Centrifugal Pump", "LVRS"): "SBT-LEO-LVRS Pump versi 2.0 (test version 3.51).xlsm",
                    ("LEO", "End Suction Centrifugal Pump", "LEP"): "SBT-LEO-LEP Pump versi 2.0 (test version 3.51).xlsm",
                    ("LEO", "Submersible Pump", "SWE"): "SBT-LEO-SWE Pump versi 2.0 (test version 3.51).xlsm",
                    ("LEO", "Submersible Pump", "XSP"): "SBT-LEO-XSP Pump versi 2.0 (test version 3.51).xlsm"
                }
                
                # Tentukan file pump yang akan digunakan
                pump_key = (pump_brand, pump_type, pump_model)
                print(f"🔍 Looking for pump key: {pump_key}")
                
                if pump_key not in pump_file_map:
                    print(f"❌ Pump configuration not found in mapping!")
                    print("📋 Available pump configurations:")
                    for key in pump_file_map.keys():
                        print(f"   {key}")
                    raise Exception(f"Pump configuration not found: {pump_key}")
                
                pump_filename = pump_file_map[pump_key]
                sbt_pump_path = os.path.join(data_folder, "SBT_PUMP", pump_filename)
                
                print(f"✅ Selected pump file: {pump_filename}")
                print(f"📁 Full path: {sbt_pump_path}")
                
                if not os.path.exists(sbt_pump_path):
                    print(f"❌ SBT_PUMP file not found at: {sbt_pump_path}")
                    raise Exception(f"SBT_PUMP file not found: {pump_filename}")
                
                print("✅ SBT_PUMP file exists!")
                
                if progress_callback:
                    progress_callback(35, f"Using pump file: {pump_filename}")
                
                print("\n⚙️ PROSES 3.1: ANAPAK → PUMP DATA TRANSFER")
                print("-" * 50)
                
                # Transfer data ke SBT_PUMP
                print(f"🔓 Opening SBT_PUMP: {pump_filename}")
                try:
                    wb_pump = load_workbook(sbt_pump_path, keep_vba=True)
                    print(f"📋 Available sheets: {wb_pump.sheetnames}")
                    
                    # Cari sheet 'DATA INPUT'
                    target_sheet_name = "DATA INPUT"
                    found_sheet = False
                    
                    for sheet_name in wb_pump.sheetnames:
                        if sheet_name.upper() == target_sheet_name.upper():
                            target_sheet_name = sheet_name
                            found_sheet = True
                            print(f"✅ Found sheet: {sheet_name}")
                            break
                    
                    if not found_sheet:
                        print(f"❌ Sheet 'DATA INPUT' not found!")
                        print(f"📋 Available sheets: {wb_pump.sheetnames}")
                        raise Exception(f"Sheet 'DATA INPUT' not found in {pump_filename}")
                    
                    sheet_pump = wb_pump[target_sheet_name]
                    
                    # Transfer data ke SBT_PUMP - NOW WITH NUMERIC VALUES
                    old_b13 = sheet_pump['B13'].value
                    old_b14 = sheet_pump['B14'].value
                    
                    sheet_pump['B13'] = value_i66  # DATA_ENGINE ANAPAK.I66 -> DATA INPUT.B13
                    sheet_pump['B14'] = value_k67  # DATA_ENGINE ANAPAK.K67 -> DATA INPUT.B14
                    
                    print(f"📊 Updated PUMP data:")
                    print(f"   B13: {old_b13} → {value_i66}")
                    print(f"   B14: {old_b14} → {value_k67}")
                    
                    # Simpan SBT_PUMP
                    print("💾 Saving SBT_PUMP...")
                    wb_pump.save(sbt_pump_path)
                    wb_pump.close()
                    print("✅ SBT_PUMP saved and closed")
                    
                    time.sleep(1)
                    
                except Exception as e:
                    error_msg = f"Failed to transfer data to SBT_PUMP: {str(e)}"
                    print(f"🚨 ERROR: {error_msg}")
                    raise Exception(error_msg)
                
                if progress_callback:
                    progress_callback(40, "Running GENERATE_REPORT macro in SBT_PUMP...")
                
                print("-" * 50)
                print("🎯 PROSES 3.2: Menjalankan macro di SBT_PUMP")
                print("-" * 50)
                
                # PROSES 3.2: Menjalankan macro di SBT_PUMP 
                print("🎯 Attempting to run GENERATE_REPORT macro...")
                macro_success = False
                
                try:
                    print("🔄 Method: Using xlwings...")
                    import xlwings as xw
                    
                    # Inisialisasi xlwings app
                    app = xw.App(visible=False)
                    app.display_alerts = False
                    app.api.AutomationSecurity = 1
                    print("✅ xlwings app initialized")
                    
                    # Buka workbook
                    wb = xw.Book(sbt_pump_path)
                    print(f"✅ Workbook opened: {os.path.basename(sbt_pump_path)}")
                    
                    # Jalankan macro GENERATE_REPORT
                    pump_filename = os.path.basename(sbt_pump_path)
                    print(f"🚀 Running GENERATE_REPORT macro from {pump_filename}...")
                    
                    try:
                        wb.api.Application.Run("GENERATE_REPORT")
                        print("✅ Macro executed successfully")
                    except Exception as macro_error:
                        print(f"⚠️ Macro error: {str(macro_error)}")
                    
                    # Tunggu macro selesai
                    print("⏳ Waiting for macro completion...")
                    time.sleep(5)
                    
                    # Pastikan semua kalkulasi selesai
                    print("🔄 Ensuring all calculations are complete...")
                    wb.api.Application.CalculateFullRebuild()
                    wb.api.Application.Calculate()
                    time.sleep(2)
                    
                    print("💾 Saving workbook...")
                    wb.save()
                    wb.close()
                    app.quit()
                    print("✅ xlwings method completed successfully")
                    macro_success = True
                    
                except Exception as e:
                    print(f"❌ xlwings method failed: {str(e)}")
                    macro_success = False
                
                if not macro_success:
                    print("⚠️ Macro failed, but continuing with projection...")
                
                print("✅ Macro execution phase completed")
                
                if progress_callback:
                    progress_callback(45, "Processing Proses 3.3 and 3.4...")
                
                print("\n📤 PROSES 3.3: PUMP → ANAPAK TRANSFER")
                print("-" * 50)
                
                # PROSES 3.3: SBT_PUMP -> SBT_ANAPAK 
                print(f"🔓 Reading results from SBT_PUMP...")
                wb_pump_output = load_workbook(sbt_pump_path, data_only=True)  # KEY: data_only=True
                print(f"📋 Available sheets: {wb_pump_output.sheetnames}")
                
                if "DATA ENGINE" not in wb_pump_output.sheetnames:
                    print("❌ Sheet DATA ENGINE not found in SBT_PUMP!")
                    print(f"📋 Available sheets: {wb_pump_output.sheetnames}")
                    raise Exception("Sheet DATA ENGINE not found in SBT_PUMP")
                    
                sheet_pump_engine = wb_pump_output["DATA ENGINE"]
                pump_value_b19 = sheet_pump_engine['B19'].value
                
                print(f"📊 Retrieved from PUMP DATA_ENGINE.B19: {pump_value_b19}")
                
                wb_pump_output.close()
                print("✅ SBT_PUMP closed")
                
                print("\n📥 PROSES 3.4: SET_BDU → ANAPAK (B42)")
                print("-" * 50)
                
                # Update SBT_ANAPAK dengan data dari PROSES 3.3 dan 3.4
                print("🔓 Opening SBT_ANAPAK for update...")
                wb_anapak = load_workbook(sbt_anapak_path)
                
                if "DATA_OUTPUT" not in wb_anapak.sheetnames:
                    print("❌ Sheet DATA_OUTPUT not found!")
                    print(f"📋 Available sheets: {wb_anapak.sheetnames}")
                    raise Exception("Sheet DATA_OUTPUT not found!")
                
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                old_c33 = sheet_anapak_output['C33'].value
                old_c34 = sheet_anapak_output['C34'].value
                
                sheet_anapak_output['C33'] = pump_value_b19    # DATA_ENGINE.B19 -> DATA_OUTPUT.C33
                sheet_anapak_output['C34'] = project_b42_value # DIP_Project Information.B42 -> DATA_OUTPUT.C34
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT:")
                print(f"   C33: {old_c33} → {pump_value_b19}")
                print(f"   C34: {old_c34} → {project_b42_value}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ SBT_ANAPAK updated and saved")
                
                if progress_callback:
                    progress_callback(50, "Starting Process 4: SET_BDU → SBT_INSTRUMENT...")
                
                print("\n🔧 PROSES 4: SET_BDU → SBT_INSTRUMENT")
                print("-" * 50)
                
                # PROSES 4: SET_BDU -> SBT_INSTRUMENT
                print(f"🔓 Opening SBT_INSTRUMENT: {sbt_instrument_path}")
                wb_instrument = load_workbook(sbt_instrument_path, keep_vba=True)
                print(f"📋 Available sheets: {wb_instrument.sheetnames}")
                
                # Cari sheet 'DATA INPUT'
                instrument_input_sheet_name = "DATA INPUT"
                found_instrument_sheet = False
                
                for sheet_name in wb_instrument.sheetnames:
                    if sheet_name.upper() == instrument_input_sheet_name.upper():
                        instrument_input_sheet_name = sheet_name
                        found_instrument_sheet = True
                        print(f"✅ Found INSTRUMENT sheet: {sheet_name}")
                        break
                
                if not found_instrument_sheet:
                    print(f"❌ Sheet 'DATA INPUT' not found in SBT_INSTRUMENT!")
                    print(f"📋 Available sheets: {wb_instrument.sheetnames}")
                    raise Exception("Sheet 'DATA INPUT' not found in SBT_INSTRUMENT")
                
                sheet_instrument_input = wb_instrument[instrument_input_sheet_name]
                
                # Transfer data: DIP_Project Information.B59 -> DATA INPUT.B4
                old_instrument_b4 = sheet_instrument_input['B4'].value
                sheet_instrument_input['B4'] = project_b59_value
                
                print(f"📊 Updated INSTRUMENT DATA_INPUT:")
                print(f"   B4: {old_instrument_b4} → {project_b59_value}")
                
                # Simpan SBT_INSTRUMENT
                print("💾 Saving SBT_INSTRUMENT...")
                wb_instrument.save(sbt_instrument_path)
                wb_instrument.close()
                print("✅ SBT_INSTRUMENT saved and closed")
                
                if progress_callback:
                    progress_callback(55, "Starting Process 5: Complex ANAPAK ↔ INSTRUMENT iterations...")
                
                print("\n🔄 PROSES 5: COMPLEX ANAPAK ↔ INSTRUMENT ITERATIONS")
                print("=" * 60)
                
                # PROSES 5.1: SBT_ANAPAK -> SBT_INSTRUMENT
                print("\n📤 PROSES 5.1: ANAPAK → INSTRUMENT (C74 → B5)")
                print("-" * 50)
                
                # Read from ANAPAK DATA_OUTPUT.C74
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                sheet_anapak_output_read = wb_anapak_read["DATA_OUTPUT"]
                anapak_c74_value = sheet_anapak_output_read['C74'].value
                
                print(f"📊 Retrieved from ANAPAK DATA_OUTPUT.C74: {anapak_c74_value}")
                wb_anapak_read.close()
                
                # Write to INSTRUMENT DATA INPUT.B5
                wb_instrument = load_workbook(sbt_instrument_path, keep_vba=True)
                sheet_instrument_input = wb_instrument[instrument_input_sheet_name]
                
                old_instrument_b5 = sheet_instrument_input['B5'].value
                sheet_instrument_input['B5'] = anapak_c74_value
                
                print(f"📊 Updated INSTRUMENT DATA_INPUT.B5: {old_instrument_b5} → {anapak_c74_value}")
                
                wb_instrument.save(sbt_instrument_path)
                wb_instrument.close()
                print("✅ INSTRUMENT updated for Process 5.1")
                
                if progress_callback:
                    progress_callback(60, "Process 5.2: INSTRUMENT calculation and data return...")
                
                # PROSES 5.2: SBT_INSTRUMENT -> SBT_ANAPAK (Force Calculating + Data Transfer)
                print("\n📥 PROSES 5.2: INSTRUMENT → ANAPAK (Force Calc + Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada INSTRUMENT
                print("🧮 Force calculating SBT_INSTRUMENT...")
                try:
                    self.force_excel_calculation(sbt_instrument_path,
                        lambda pct, msg: progress_callback(60 + (pct * 0.03), f"INSTRUMENT Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_INSTRUMENT calculation completed")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_INSTRUMENT calculation: {str(e)}")
                
                # Read calculated values from INSTRUMENT
                wb_instrument_read = load_workbook(sbt_instrument_path, data_only=True)
                
                # Cari sheet 'DATA PROPOSAL'
                instrument_proposal_sheet_name = "DATA PROPOSAL"
                found_proposal_sheet = False
                
                for sheet_name in wb_instrument_read.sheetnames:
                    if sheet_name.upper() == instrument_proposal_sheet_name.upper():
                        instrument_proposal_sheet_name = sheet_name
                        found_proposal_sheet = True
                        print(f"✅ Found INSTRUMENT PROPOSAL sheet: {sheet_name}")
                        break
                
                if not found_proposal_sheet:
                    print(f"❌ Sheet 'DATA PROPOSAL' not found in SBT_INSTRUMENT!")
                    print(f"📋 Available sheets: {wb_instrument_read.sheetnames}")
                    raise Exception("Sheet 'DATA PROPOSAL' not found in SBT_INSTRUMENT")
                
                sheet_instrument_proposal = wb_instrument_read[instrument_proposal_sheet_name]
                
                # Read values: DATA PROPOSAL.B8 and B5
                instrument_b8_value = sheet_instrument_proposal['B8'].value
                instrument_b5_value = sheet_instrument_proposal['B5'].value
                
                print(f"📊 Retrieved from INSTRUMENT DATA_PROPOSAL:")
                print(f"   B8: {instrument_b8_value}")
                print(f"   B5: {instrument_b5_value}")
                
                wb_instrument_read.close()
                
                # Write to ANAPAK DATA_OUTPUT: B8->C75, B5->C76
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                old_c75 = sheet_anapak_output['C75'].value
                old_c76 = sheet_anapak_output['C76'].value
                
                sheet_anapak_output['C75'] = instrument_b8_value  # DATA PROPOSAL.B8 -> DATA_OUTPUT.C75
                sheet_anapak_output['C76'] = instrument_b5_value  # DATA PROPOSAL.B5 -> DATA_OUTPUT.C76
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 5.2):")
                print(f"   C75: {old_c75} → {instrument_b8_value}")
                print(f"   C76: {old_c76} → {instrument_b5_value}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 5.2")
                
                if progress_callback:
                    progress_callback(65, "Process 5.3: ANAPAK → INSTRUMENT (C80 → B5)...")
                
                # PROSES 5.3: SBT_ANAPAK -> SBT_INSTRUMENT (C80 -> B5)
                print("\n📤 PROSES 5.3: ANAPAK → INSTRUMENT (C80 → B5)")
                print("-" * 50)
                
                # Read from ANAPAK DATA_OUTPUT.C80
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                sheet_anapak_output_read = wb_anapak_read["DATA_OUTPUT"]
                anapak_c80_value = sheet_anapak_output_read['C80'].value
                
                print(f"📊 Retrieved from ANAPAK DATA_OUTPUT.C80: {anapak_c80_value}")
                wb_anapak_read.close()
                
                # Write to INSTRUMENT DATA INPUT.B5
                wb_instrument = load_workbook(sbt_instrument_path, keep_vba=True)
                sheet_instrument_input = wb_instrument[instrument_input_sheet_name]
                
                old_instrument_b5_2 = sheet_instrument_input['B5'].value
                sheet_instrument_input['B5'] = anapak_c80_value
                
                print(f"📊 Updated INSTRUMENT DATA_INPUT.B5: {old_instrument_b5_2} → {anapak_c80_value}")
                
                wb_instrument.save(sbt_instrument_path)
                wb_instrument.close()
                print("✅ INSTRUMENT updated for Process 5.3")
                
                if progress_callback:
                    progress_callback(70, "Process 5.4: INSTRUMENT calculation and data return...")
                
                # PROSES 5.4: SBT_INSTRUMENT -> SBT_ANAPAK (Force Calculating + Data Transfer)
                print("\n📥 PROSES 5.4: INSTRUMENT → ANAPAK (Force Calc + Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada INSTRUMENT
                print("🧮 Force calculating SBT_INSTRUMENT (round 2)...")
                try:
                    self.force_excel_calculation(sbt_instrument_path,
                        lambda pct, msg: progress_callback(70 + (pct * 0.03), f"INSTRUMENT Calc 2: {msg}") if progress_callback else None)
                    print("✅ SBT_INSTRUMENT calculation completed (round 2)")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_INSTRUMENT calculation (round 2): {str(e)}")
                
                # Read calculated values from INSTRUMENT (round 2)
                wb_instrument_read = load_workbook(sbt_instrument_path, data_only=True)
                sheet_instrument_proposal = wb_instrument_read[instrument_proposal_sheet_name]
                
                # Read values: DATA PROPOSAL.B8 and B5 (round 2)
                instrument_b8_value_2 = sheet_instrument_proposal['B8'].value
                instrument_b5_value_2 = sheet_instrument_proposal['B5'].value
                
                print(f"📊 Retrieved from INSTRUMENT DATA_PROPOSAL (round 2):")
                print(f"   B8: {instrument_b8_value_2}")
                print(f"   B5: {instrument_b5_value_2}")
                
                wb_instrument_read.close()
                
                # Write to ANAPAK DATA_OUTPUT: B8->C81, B5->C82
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                old_c81 = sheet_anapak_output['C81'].value
                old_c82 = sheet_anapak_output['C82'].value
                
                sheet_anapak_output['C81'] = instrument_b8_value_2  # DATA PROPOSAL.B8 -> DATA_OUTPUT.C81
                sheet_anapak_output['C82'] = instrument_b5_value_2  # DATA PROPOSAL.B5 -> DATA_OUTPUT.C82
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 5.4):")
                print(f"   C81: {old_c81} → {instrument_b8_value_2}")
                print(f"   C82: {old_c82} → {instrument_b5_value_2}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 5.4")
                
                if progress_callback:
                    progress_callback(75, "Process 5.5: ANAPAK → INSTRUMENT (C86 → B5)...")
                
                # PROSES 5.5: SBT_ANAPAK -> SBT_INSTRUMENT (C86 -> B5)
                print("\n📤 PROSES 5.5: ANAPAK → INSTRUMENT (C86 → B5)")
                print("-" * 50)
                
                # Read from ANAPAK DATA_OUTPUT.C86
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                sheet_anapak_output_read = wb_anapak_read["DATA_OUTPUT"]
                anapak_c86_value = sheet_anapak_output_read['C86'].value
                
                print(f"📊 Retrieved from ANAPAK DATA_OUTPUT.C86: {anapak_c86_value}")
                wb_anapak_read.close()
                
                # Write to INSTRUMENT DATA INPUT.B5
                wb_instrument = load_workbook(sbt_instrument_path, keep_vba=True)
                sheet_instrument_input = wb_instrument[instrument_input_sheet_name]
                
                old_instrument_b5_3 = sheet_instrument_input['B5'].value
                sheet_instrument_input['B5'] = anapak_c86_value
                
                print(f"📊 Updated INSTRUMENT DATA_INPUT.B5: {old_instrument_b5_3} → {anapak_c86_value}")
                
                wb_instrument.save(sbt_instrument_path)
                wb_instrument.close()
                print("✅ INSTRUMENT updated for Process 5.5")
                
                if progress_callback:
                    progress_callback(85, "Process 5.6: Final INSTRUMENT calculation and data return...")
                
                # PROSES 5.6: SBT_INSTRUMENT -> SBT_ANAPAK (Force Calculating + Data Transfer - Final)
                print("\n📥 PROSES 5.6: INSTRUMENT → ANAPAK (Final Force Calc + Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada INSTRUMENT (final)
                print("🧮 Force calculating SBT_INSTRUMENT (final round)...")
                try:
                    self.force_excel_calculation(sbt_instrument_path,
                        lambda pct, msg: progress_callback(85 + (pct * 0.05), f"INSTRUMENT Final Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_INSTRUMENT calculation completed (final round)")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_INSTRUMENT calculation (final round): {str(e)}")
                
                # Read calculated values from INSTRUMENT (final)
                wb_instrument_read = load_workbook(sbt_instrument_path, data_only=True)
                sheet_instrument_proposal = wb_instrument_read[instrument_proposal_sheet_name]
                
                # Read values: DATA PROPOSAL.B8 and B5 (final)
                instrument_b8_value_final = sheet_instrument_proposal['B8'].value
                instrument_b5_value_final = sheet_instrument_proposal['B5'].value
                
                print(f"📊 Retrieved from INSTRUMENT DATA_PROPOSAL (final):")
                print(f"   B8: {instrument_b8_value_final}")
                print(f"   B5: {instrument_b5_value_final}")
                
                wb_instrument_read.close()
                
                # Write to ANAPAK DATA_OUTPUT: B8->C87, B5->C88
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                old_c87 = sheet_anapak_output['C87'].value
                old_c88 = sheet_anapak_output['C88'].value
                
                sheet_anapak_output['C87'] = instrument_b8_value_final  # DATA PROPOSAL.B8 -> DATA_OUTPUT.C87
                sheet_anapak_output['C88'] = instrument_b5_value_final  # DATA PROPOSAL.B5 -> DATA_OUTPUT.C88
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 5.6 - Final):")
                print(f"   C87: {old_c87} → {instrument_b8_value_final}")
                print(f"   C88: {old_c88} → {instrument_b5_value_final}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 5.6 (Final)")
                
                if progress_callback:
                    progress_callback(90, "Starting Process 6: SET_BDU → SBT_DOSINGPUMP...")
                
                print("\n💊 PROSES 6: SET_BDU → SBT_DOSINGPUMP")
                print("-" * 50)
            
                # PROSES 6: SET_BDU -> SBT_DOSINGPUMP
                print(f"🔓 Opening SBT_DOSINGPUMP: {sbt_dosingpump_path}")
                wb_dosingpump = load_workbook(sbt_dosingpump_path, keep_vba=True)
                print(f"📋 Available sheets: {wb_dosingpump.sheetnames}")
                
                # Cari sheet 'DATA INPUT'
                dosingpump_input_sheet_name = "DATA INPUT"
                found_dosingpump_sheet = False
                
                for sheet_name in wb_dosingpump.sheetnames:
                    if sheet_name.upper() == dosingpump_input_sheet_name.upper():
                        dosingpump_input_sheet_name = sheet_name
                        found_dosingpump_sheet = True
                        print(f"✅ Found DOSINGPUMP sheet: {sheet_name}")
                        break
                
                if not found_dosingpump_sheet:
                    print(f"❌ Sheet 'DATA INPUT' not found in SBT_DOSINGPUMP!")
                    print(f"📋 Available sheets: {wb_dosingpump.sheetnames}")
                    raise Exception("Sheet 'DATA INPUT' not found in SBT_DOSINGPUMP")
                
                sheet_dosingpump_input = wb_dosingpump[dosingpump_input_sheet_name]
                
                # Transfer data: DIP_Project Information.B45 -> DATA INPUT.B6
                # Re-read B45 value dari SET_BDU
                wb_bdu_read = load_workbook(set_bdu_path, data_only=True)
                sheet_project_read = wb_bdu_read["DIP_Project Information"]
                project_b45_value = sheet_project_read['B45'].value
                wb_bdu_read.close()
                
                old_dosingpump_b6 = sheet_dosingpump_input['B6'].value
                sheet_dosingpump_input['B6'] = project_b45_value
                
                print(f"📊 Updated DOSINGPUMP DATA_INPUT:")
                print(f"   B6: {old_dosingpump_b6} → {project_b45_value}")
                
                # Simpan SBT_DOSINGPUMP
                print("💾 Saving SBT_DOSINGPUMP...")
                wb_dosingpump.save(sbt_dosingpump_path)
                wb_dosingpump.close()
                print("✅ SBT_DOSINGPUMP saved and closed")
                
                if progress_callback:
                    progress_callback(90, "Starting Process 7: Complex ANAPAK ↔ DOSINGPUMP iterations...")
                
                print("\n💊 PROSES 7: COMPLEX ANAPAK ↔ DOSINGPUMP ITERATIONS")
                print("=" * 60)
                
                # PROSES 7.1: SBT_ANAPAK -> SBT_DOSINGPUMP
                print("\n📤 PROSES 7.1: ANAPAK → DOSINGPUMP (Q10 → B4)")
                print("-" * 50)
                
                # Force calculation pada ANAPAK sebelum membaca Q10
                print("🧮 Force calculating SBT_ANAPAK before reading Q10...")
                try:
                    self.force_excel_calculation(sbt_anapak_path,
                        lambda pct, msg: progress_callback(90 + (pct * 0.02), f"ANAPAK Pre-Q10 Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_ANAPAK calculation completed before Q10 read")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_ANAPAK calculation: {str(e)}")
                
                # Read from ANAPAK CHEMICAL DOSAGE CALC_ANAPAK.Q10
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                
                # Cari sheet 'CHEMICAL DOSAGE CALC_ANAPAK'
                chemical_dosage_sheet_name = "CHEMICAL DOSAGE CALC_ANAPAK"
                found_chemical_sheet = False
                
                for sheet_name in wb_anapak_read.sheetnames:
                    if "CHEMICAL DOSAGE CALC" in sheet_name.upper() and "ANAPAK" in sheet_name.upper():
                        chemical_dosage_sheet_name = sheet_name
                        found_chemical_sheet = True
                        print(f"✅ Found CHEMICAL DOSAGE sheet: {sheet_name}")
                        break
                
                if not found_chemical_sheet:
                    print(f"❌ Sheet 'CHEMICAL DOSAGE CALC_ANAPAK' not found!")
                    print(f"📋 Available sheets: {wb_anapak_read.sheetnames}")
                    raise Exception("Sheet 'CHEMICAL DOSAGE CALC_ANAPAK' not found in SBT_ANAPAK")
                
                sheet_chemical_dosage = wb_anapak_read[chemical_dosage_sheet_name]
                anapak_q10_value = sheet_chemical_dosage['Q10'].value
                
                print(f"📊 Retrieved from ANAPAK CHEMICAL_DOSAGE.Q10: {anapak_q10_value}")
                wb_anapak_read.close()
                
                # Write to DOSINGPUMP DATA INPUT.B4
                wb_dosingpump = load_workbook(sbt_dosingpump_path, keep_vba=True)
                sheet_dosingpump_input = wb_dosingpump[dosingpump_input_sheet_name]
                
                old_dosingpump_b4 = sheet_dosingpump_input['B4'].value
                sheet_dosingpump_input['B4'] = anapak_q10_value
                
                print(f"📊 Updated DOSINGPUMP DATA_INPUT.B4: {old_dosingpump_b4} → {anapak_q10_value}")
                
                wb_dosingpump.save(sbt_dosingpump_path)
                wb_dosingpump.close()
                print("✅ DOSINGPUMP updated for Process 7.1")
                
                if progress_callback:
                    progress_callback(92, "Process 7.2: DOSINGPUMP calculation and complex data return...")
                
                # PROSES 7.2: SBT_DOSINGPUMP -> SBT_ANAPAK (Force Calculating + Multiple Data Transfer)
                print("\n📥 PROSES 7.2: DOSINGPUMP → ANAPAK (Force Calc + Multiple Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada DOSINGPUMP
                print("🧮 Force calculating SBT_DOSINGPUMP...")
                try:
                    self.force_excel_calculation(sbt_dosingpump_path,
                        lambda pct, msg: progress_callback(92 + (pct * 0.01), f"DOSINGPUMP Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_DOSINGPUMP calculation completed")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_DOSINGPUMP calculation: {str(e)}")
                
                # Read calculated values from DOSINGPUMP DATA PROPOSAL
                wb_dosingpump_read = load_workbook(sbt_dosingpump_path, data_only=True)
                
                # Cari sheet 'DATA PROPOSAL'
                dosingpump_proposal_sheet_name = "DATA PROPOSAL"
                found_dosingpump_proposal = False
                
                for sheet_name in wb_dosingpump_read.sheetnames:
                    if sheet_name.upper() == dosingpump_proposal_sheet_name.upper():
                        dosingpump_proposal_sheet_name = sheet_name
                        found_dosingpump_proposal = True
                        print(f"✅ Found DOSINGPUMP PROPOSAL sheet: {sheet_name}")
                        break
                
                if not found_dosingpump_proposal:
                    print(f"❌ Sheet 'DATA PROPOSAL' not found in SBT_DOSINGPUMP!")
                    print(f"📋 Available sheets: {wb_dosingpump_read.sheetnames}")
                    raise Exception("Sheet 'DATA PROPOSAL' not found in SBT_DOSINGPUMP")
                
                sheet_dosingpump_proposal = wb_dosingpump_read[dosingpump_proposal_sheet_name]
                
                # Read multiple values from DATA PROPOSAL
                dosingpump_b5_value = sheet_dosingpump_proposal['B5'].value
                dosingpump_b6_value = sheet_dosingpump_proposal['B6'].value
                dosingpump_b7_value = sheet_dosingpump_proposal['B7'].value
                dosingpump_b8_value = sheet_dosingpump_proposal['B8'].value
                dosingpump_b9_value = sheet_dosingpump_proposal['B9'].value
                dosingpump_b10_value = sheet_dosingpump_proposal['B10'].value
                
                print(f"📊 Retrieved from DOSINGPUMP DATA_PROPOSAL:")
                print(f"   B5: {dosingpump_b5_value}")
                print(f"   B6: {dosingpump_b6_value}")
                print(f"   B7: {dosingpump_b7_value}")
                print(f"   B8: {dosingpump_b8_value}")
                print(f"   B9: {dosingpump_b9_value}")
                print(f"   B10: {dosingpump_b10_value}")
                
                wb_dosingpump_read.close()
                
                # Write to ANAPAK CHEMICAL DOSAGE CALC and DATA_OUTPUT
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_chemical_dosage_write = wb_anapak[chemical_dosage_sheet_name]
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                # Update CHEMICAL DOSAGE CALC_ANAPAK (Process 7.2)
                old_s10 = sheet_chemical_dosage_write['S10'].value
                old_r10 = sheet_chemical_dosage_write['R10'].value
                old_t10 = sheet_chemical_dosage_write['T10'].value
                old_u10 = sheet_chemical_dosage_write['U10'].value
                old_v10 = sheet_chemical_dosage_write['V10'].value
                old_w10 = sheet_chemical_dosage_write['W10'].value
                
                sheet_chemical_dosage_write['S10'] = dosingpump_b5_value   # B5 -> S10
                sheet_chemical_dosage_write['R10'] = dosingpump_b6_value   # B6 -> R10
                sheet_chemical_dosage_write['T10'] = dosingpump_b7_value   # B7 -> T10
                sheet_chemical_dosage_write['U10'] = dosingpump_b9_value   # B9 -> U10
                sheet_chemical_dosage_write['V10'] = dosingpump_b8_value   # B8 -> V10
                sheet_chemical_dosage_write['W10'] = dosingpump_b10_value  # B10 -> W10
                
                print(f"📊 Updated ANAPAK CHEMICAL_DOSAGE (Process 7.2):")
                print(f"   S10: {old_s10} → {dosingpump_b5_value}")
                print(f"   R10: {old_r10} → {dosingpump_b6_value}")
                print(f"   T10: {old_t10} → {dosingpump_b7_value}")
                print(f"   U10: {old_u10} → {dosingpump_b9_value}")
                print(f"   V10: {old_v10} → {dosingpump_b8_value}")
                print(f"   W10: {old_w10} → {dosingpump_b10_value}")
                
                # Update DATA_OUTPUT (Process 7.2)
                old_c38 = sheet_anapak_output['C38'].value
                old_c39 = sheet_anapak_output['C39'].value
                old_c41 = sheet_anapak_output['C41'].value
                
                sheet_anapak_output['C38'] = dosingpump_b8_value   # B8 -> C38
                sheet_anapak_output['C39'] = dosingpump_b10_value  # B10 -> C39
                sheet_anapak_output['C41'] = dosingpump_b6_value   # B6 -> C41
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 7.2):")
                print(f"   C38: {old_c38} → {dosingpump_b8_value}")
                print(f"   C39: {old_c39} → {dosingpump_b10_value}")
                print(f"   C41: {old_c41} → {dosingpump_b6_value}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 7.2")
                
                if progress_callback:
                    progress_callback(94, "Process 7.3-7.4: Second DOSINGPUMP iteration...")
                
                # PROSES 7.3: SBT_ANAPAK -> SBT_DOSINGPUMP (Q11 -> B4)
                print("\n📤 PROSES 7.3: ANAPAK → DOSINGPUMP (Q11 → B4)")
                print("-" * 50)
                
                # Force calculation pada ANAPAK sebelum membaca Q11
                print("🧮 Force calculating SBT_ANAPAK before reading Q11...")
                try:
                    self.force_excel_calculation(sbt_anapak_path,
                        lambda pct, msg: progress_callback(94 + (pct * 0.01), f"ANAPAK Pre-Q11 Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_ANAPAK calculation completed before Q11 read")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_ANAPAK calculation: {str(e)}")
                
                # Read from ANAPAK CHEMICAL DOSAGE CALC_ANAPAK.Q11
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                sheet_chemical_dosage_read = wb_anapak_read[chemical_dosage_sheet_name]
                anapak_q11_value = sheet_chemical_dosage_read['Q11'].value
                
                print(f"📊 Retrieved from ANAPAK CHEMICAL_DOSAGE.Q11: {anapak_q11_value}")
                wb_anapak_read.close()
                
                # Write to DOSINGPUMP DATA INPUT.B4
                wb_dosingpump = load_workbook(sbt_dosingpump_path, keep_vba=True)
                sheet_dosingpump_input = wb_dosingpump[dosingpump_input_sheet_name]
                
                old_dosingpump_b4_2 = sheet_dosingpump_input['B4'].value
                sheet_dosingpump_input['B4'] = anapak_q11_value
                
                print(f"📊 Updated DOSINGPUMP DATA_INPUT.B4: {old_dosingpump_b4_2} → {anapak_q11_value}")
                
                wb_dosingpump.save(sbt_dosingpump_path)
                wb_dosingpump.close()
                print("✅ DOSINGPUMP updated for Process 7.3")
                
                # PROSES 7.4: SBT_DOSINGPUMP -> SBT_ANAPAK (Force Calculating + Multiple Data Transfer)
                print("\n📥 PROSES 7.4: DOSINGPUMP → ANAPAK (Force Calc + Multiple Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada DOSINGPUMP
                print("🧮 Force calculating SBT_DOSINGPUMP (round 2)...")
                try:
                    self.force_excel_calculation(sbt_dosingpump_path,
                        lambda pct, msg: progress_callback(95 + (pct * 0.01), f"DOSINGPUMP Calc 2: {msg}") if progress_callback else None)
                    print("✅ SBT_DOSINGPUMP calculation completed (round 2)")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_DOSINGPUMP calculation (round 2): {str(e)}")
                
                # Read calculated values from DOSINGPUMP DATA PROPOSAL (round 2)
                wb_dosingpump_read = load_workbook(sbt_dosingpump_path, data_only=True)
                sheet_dosingpump_proposal = wb_dosingpump_read[dosingpump_proposal_sheet_name]
                
                # Read multiple values from DATA PROPOSAL (round 2)
                dosingpump_b5_value_2 = sheet_dosingpump_proposal['B5'].value
                dosingpump_b6_value_2 = sheet_dosingpump_proposal['B6'].value
                dosingpump_b7_value_2 = sheet_dosingpump_proposal['B7'].value
                dosingpump_b8_value_2 = sheet_dosingpump_proposal['B8'].value
                dosingpump_b9_value_2 = sheet_dosingpump_proposal['B9'].value
                dosingpump_b10_value_2 = sheet_dosingpump_proposal['B10'].value
                
                print(f"📊 Retrieved from DOSINGPUMP DATA_PROPOSAL (round 2):")
                print(f"   B5: {dosingpump_b5_value_2}")
                print(f"   B6: {dosingpump_b6_value_2}")
                print(f"   B7: {dosingpump_b7_value_2}")
                print(f"   B8: {dosingpump_b8_value_2}")
                print(f"   B9: {dosingpump_b9_value_2}")
                print(f"   B10: {dosingpump_b10_value_2}")
                
                wb_dosingpump_read.close()
                
                # Write to ANAPAK CHEMICAL DOSAGE CALC and DATA_OUTPUT (round 2)
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_chemical_dosage_write = wb_anapak[chemical_dosage_sheet_name]
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                # Update CHEMICAL DOSAGE CALC_ANAPAK (Process 7.4)
                old_s11 = sheet_chemical_dosage_write['S11'].value
                old_r11 = sheet_chemical_dosage_write['R11'].value
                old_t11 = sheet_chemical_dosage_write['T11'].value
                old_u11 = sheet_chemical_dosage_write['U11'].value
                old_v11 = sheet_chemical_dosage_write['V11'].value
                old_w11 = sheet_chemical_dosage_write['W11'].value
                
                sheet_chemical_dosage_write['S11'] = dosingpump_b5_value_2   # B5 -> S11
                sheet_chemical_dosage_write['R11'] = dosingpump_b6_value_2   # B6 -> R11
                sheet_chemical_dosage_write['T11'] = dosingpump_b7_value_2   # B7 -> T11
                sheet_chemical_dosage_write['U11'] = dosingpump_b9_value_2   # B9 -> U11
                sheet_chemical_dosage_write['V11'] = dosingpump_b8_value_2   # B8 -> V11
                sheet_chemical_dosage_write['W11'] = dosingpump_b10_value_2  # B10 -> W11
                
                print(f"📊 Updated ANAPAK CHEMICAL_DOSAGE (Process 7.4):")
                print(f"   S11: {old_s11} → {dosingpump_b5_value_2}")
                print(f"   R11: {old_r11} → {dosingpump_b6_value_2}")
                print(f"   T11: {old_t11} → {dosingpump_b7_value_2}")
                print(f"   U11: {old_u11} → {dosingpump_b9_value_2}")
                print(f"   V11: {old_v11} → {dosingpump_b8_value_2}")
                print(f"   W11: {old_w11} → {dosingpump_b10_value_2}")
                
                # Update DATA_OUTPUT (Process 7.4)
                old_c43 = sheet_anapak_output['C43'].value
                old_c44 = sheet_anapak_output['C44'].value
                old_c46 = sheet_anapak_output['C46'].value
                
                sheet_anapak_output['C43'] = dosingpump_b8_value_2   # B8 -> C43
                sheet_anapak_output['C44'] = dosingpump_b10_value_2  # B10 -> C44
                sheet_anapak_output['C46'] = dosingpump_b6_value_2   # B6 -> C46
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 7.4):")
                print(f"   C43: {old_c43} → {dosingpump_b8_value_2}")
                print(f"   C44: {old_c44} → {dosingpump_b10_value_2}")
                print(f"   C46: {old_c46} → {dosingpump_b6_value_2}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 7.4")
                
                if progress_callback:
                    progress_callback(96, "Process 7.5-7.6: Final DOSINGPUMP iteration...")
                
                # PROSES 7.5: SBT_ANAPAK -> SBT_DOSINGPUMP (Q13 -> B4)
                print("\n📤 PROSES 7.5: ANAPAK → DOSINGPUMP (Q13 → B4)")
                print("-" * 50)
                
                # Force calculation pada ANAPAK sebelum membaca Q13
                print("🧮 Force calculating SBT_ANAPAK before reading Q13...")
                try:
                    self.force_excel_calculation(sbt_anapak_path,
                        lambda pct, msg: progress_callback(96 + (pct * 0.01), f"ANAPAK Pre-Q13 Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_ANAPAK calculation completed before Q13 read")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_ANAPAK calculation: {str(e)}")
                
                # Read from ANAPAK CHEMICAL DOSAGE CALC_ANAPAK.Q13
                wb_anapak_read = load_workbook(sbt_anapak_path, data_only=True)
                sheet_chemical_dosage_read = wb_anapak_read[chemical_dosage_sheet_name]
                anapak_q13_value = sheet_chemical_dosage_read['Q13'].value
                
                print(f"📊 Retrieved from ANAPAK CHEMICAL_DOSAGE.Q13: {anapak_q13_value}")
                wb_anapak_read.close()
                
                # Write to DOSINGPUMP DATA INPUT.B4
                wb_dosingpump = load_workbook(sbt_dosingpump_path, keep_vba=True)
                sheet_dosingpump_input = wb_dosingpump[dosingpump_input_sheet_name]
                
                old_dosingpump_b4_3 = sheet_dosingpump_input['B4'].value
                sheet_dosingpump_input['B4'] = anapak_q13_value
                
                print(f"📊 Updated DOSINGPUMP DATA_INPUT.B4: {old_dosingpump_b4_3} → {anapak_q13_value}")
                
                wb_dosingpump.save(sbt_dosingpump_path)
                wb_dosingpump.close()
                print("✅ DOSINGPUMP updated for Process 7.5")
                
                # PROSES 7.6: SBT_DOSINGPUMP -> SBT_ANAPAK (Final Force Calculating + Multiple Data Transfer)
                print("\n📥 PROSES 7.6: DOSINGPUMP → ANAPAK (Final Force Calc + Multiple Data Transfer)")
                print("-" * 50)
                
                # Force calculation pada DOSINGPUMP (final)
                print("🧮 Force calculating SBT_DOSINGPUMP (final round)...")
                try:
                    self.force_excel_calculation(sbt_dosingpump_path,
                        lambda pct, msg: progress_callback(97 + (pct * 0.01), f"DOSINGPUMP Final Calc: {msg}") if progress_callback else None)
                    print("✅ SBT_DOSINGPUMP calculation completed (final round)")
                    time.sleep(2)
                except Exception as e:
                    print(f"⚠️ Warning during SBT_DOSINGPUMP calculation (final round): {str(e)}")
                
                # Read calculated values from DOSINGPUMP DATA PROPOSAL (final)
                wb_dosingpump_read = load_workbook(sbt_dosingpump_path, data_only=True)
                sheet_dosingpump_proposal = wb_dosingpump_read[dosingpump_proposal_sheet_name]
                
                # Read multiple values from DATA PROPOSAL (final)
                dosingpump_b5_value_final = sheet_dosingpump_proposal['B5'].value
                dosingpump_b6_value_final = sheet_dosingpump_proposal['B6'].value
                dosingpump_b7_value_final = sheet_dosingpump_proposal['B7'].value
                dosingpump_b8_value_final = sheet_dosingpump_proposal['B8'].value
                dosingpump_b9_value_final = sheet_dosingpump_proposal['B9'].value
                dosingpump_b10_value_final = sheet_dosingpump_proposal['B10'].value
                
                print(f"📊 Retrieved from DOSINGPUMP DATA_PROPOSAL (final):")
                print(f"   B5: {dosingpump_b5_value_final}")
                print(f"   B6: {dosingpump_b6_value_final}")
                print(f"   B7: {dosingpump_b7_value_final}")
                print(f"   B8: {dosingpump_b8_value_final}")
                print(f"   B9: {dosingpump_b9_value_final}")
                print(f"   B10: {dosingpump_b10_value_final}")
                
                wb_dosingpump_read.close()
                
                # Write to ANAPAK CHEMICAL DOSAGE CALC and DATA_OUTPUT (final)
                wb_anapak = load_workbook(sbt_anapak_path)
                sheet_chemical_dosage_write = wb_anapak[chemical_dosage_sheet_name]
                sheet_anapak_output = wb_anapak["DATA_OUTPUT"]
                
                # Update CHEMICAL DOSAGE CALC_ANAPAK (Process 7.6)
                old_s13 = sheet_chemical_dosage_write['S13'].value
                old_r13 = sheet_chemical_dosage_write['R13'].value
                old_t13 = sheet_chemical_dosage_write['T13'].value
                old_u13 = sheet_chemical_dosage_write['U13'].value
                old_v13 = sheet_chemical_dosage_write['V13'].value
                old_w13 = sheet_chemical_dosage_write['W13'].value
                
                sheet_chemical_dosage_write['S13'] = dosingpump_b5_value_final   # B5 -> S13
                sheet_chemical_dosage_write['R13'] = dosingpump_b6_value_final   # B6 -> R13
                sheet_chemical_dosage_write['T13'] = dosingpump_b7_value_final   # B7 -> T13
                sheet_chemical_dosage_write['U13'] = dosingpump_b9_value_final   # B9 -> U13
                sheet_chemical_dosage_write['V13'] = dosingpump_b8_value_final   # B8 -> V13
                sheet_chemical_dosage_write['W13'] = dosingpump_b10_value_final  # B10 -> W13
                
                print(f"📊 Updated ANAPAK CHEMICAL_DOSAGE (Process 7.6 - Final):")
                print(f"   S13: {old_s13} → {dosingpump_b5_value_final}")
                print(f"   R13: {old_r13} → {dosingpump_b6_value_final}")
                print(f"   T13: {old_t13} → {dosingpump_b7_value_final}")
                print(f"   U13: {old_u13} → {dosingpump_b9_value_final}")
                print(f"   V13: {old_v13} → {dosingpump_b8_value_final}")
                print(f"   W13: {old_w13} → {dosingpump_b10_value_final}")
                
                # Update DATA_OUTPUT (Process 7.6)
                old_c48 = sheet_anapak_output['C48'].value
                old_c49 = sheet_anapak_output['C49'].value
                old_c51 = sheet_anapak_output['C51'].value
                
                sheet_anapak_output['C48'] = dosingpump_b8_value_final   # B8 -> C48
                sheet_anapak_output['C49'] = dosingpump_b10_value_final  # B10 -> C49
                sheet_anapak_output['C51'] = dosingpump_b6_value_final   # B6 -> C51
                
                print(f"📊 Updated ANAPAK DATA_OUTPUT (Process 7.6 - Final):")
                print(f"   C48: {old_c48} → {dosingpump_b8_value_final}")
                print(f"   C49: {old_c49} → {dosingpump_b10_value_final}")
                print(f"   C51: {old_c51} → {dosingpump_b6_value_final}")
                
                wb_anapak.save(sbt_anapak_path)
                wb_anapak.close()
                print("✅ ANAPAK updated for Process 7.6 (Final)")
                
                if progress_callback:
                    progress_callback(98, "Performing final calculations and cleanup...")
                
                print("\n🏁 FINAL CALCULATIONS AND CLEANUP")
                print("-" * 50)
                
                # Final force calculation pada ANAPAK untuk memastikan semua nilai terupdate
                print("🧮 Final force calculation on SBT_ANAPAK...")
                try:
                    self.force_excel_calculation(sbt_anapak_path,
                        lambda pct, msg: progress_callback(98 + (pct * 0.015), f"Final ANAPAK Calc: {msg}") if progress_callback else None)
                    print("✅ Final SBT_ANAPAK calculation completed")
                    time.sleep(3)
                except Exception as e:
                    print(f"⚠️ Warning during final SBT_ANAPAK calculation: {str(e)}")
                
                if progress_callback:
                    progress_callback(99, "Validating all data transfers...")
                
                print("\n📋 COMPREHENSIVE VALIDATION SUMMARY")
                print("-" * 50)
                
                # Validation summary - read final values to confirm everything worked
                wb_anapak_final = load_workbook(sbt_anapak_path, data_only=True)
                sheet_anapak_output_final = wb_anapak_final["DATA_OUTPUT"]
                
                final_values = {
                    # Process 3 results
                    'C33': sheet_anapak_output_final['C33'].value,  # From Process 3.3
                    'C34': sheet_anapak_output_final['C34'].value,  # From Process 3.4
                    # Process 5 results
                    'C75': sheet_anapak_output_final['C75'].value,  # From Process 5.2
                    'C76': sheet_anapak_output_final['C76'].value,  # From Process 5.2
                    'C81': sheet_anapak_output_final['C81'].value,  # From Process 5.4
                    'C82': sheet_anapak_output_final['C82'].value,  # From Process 5.4
                    'C87': sheet_anapak_output_final['C87'].value,  # From Process 5.6
                    'C88': sheet_anapak_output_final['C88'].value,  # From Process 5.6
                    # Process 7 results (DOSINGPUMP)
                    'C38': sheet_anapak_output_final['C38'].value,  # From Process 7.2
                    'C39': sheet_anapak_output_final['C39'].value,  # From Process 7.2
                    'C41': sheet_anapak_output_final['C41'].value,  # From Process 7.2
                    'C43': sheet_anapak_output_final['C43'].value,  # From Process 7.4
                    'C44': sheet_anapak_output_final['C44'].value,  # From Process 7.4
                    'C46': sheet_anapak_output_final['C46'].value,  # From Process 7.4
                    'C48': sheet_anapak_output_final['C48'].value,  # From Process 7.6
                    'C49': sheet_anapak_output_final['C49'].value,  # From Process 7.6
                    'C51': sheet_anapak_output_final['C51'].value,  # From Process 7.6
                }
                
                print("📊 Final ANAPAK DATA_OUTPUT values:")
                print("   📈 Process 3 (PUMP) results:")
                print(f"     C33: {final_values['C33']}")
                print(f"     C34: {final_values['C34']}")
                print("   🔧 Process 5 (INSTRUMENT) results:")
                print(f"     C75: {final_values['C75']}")
                print(f"     C76: {final_values['C76']}")
                print(f"     C81: {final_values['C81']}")
                print(f"     C82: {final_values['C82']}")
                print(f"     C87: {final_values['C87']}")
                print(f"     C88: {final_values['C88']}")
                print("   💊 Process 7 (DOSINGPUMP) results:")
                print(f"     C38: {final_values['C38']}")
                print(f"     C39: {final_values['C39']}")
                print(f"     C41: {final_values['C41']}")
                print(f"     C43: {final_values['C43']}")
                print(f"     C44: {final_values['C44']}")
                print(f"     C46: {final_values['C46']}")
                print(f"     C48: {final_values['C48']}")
                print(f"     C49: {final_values['C49']}")
                print(f"     C51: {final_values['C51']}")
                
                wb_anapak_final.close()
                
                if progress_callback:
                    progress_callback(100, "Complete projection process with all 7 processes finished successfully!")
                
                print("\n" + "=" * 80)
                print("🎉 COMPLETE PROJECTION PROCESS WITH ALL 7 PROCESSES FINISHED SUCCESSFULLY!")
                print("=" * 80)
                print("🏆 ALL MODULES PROCESSED: ANAPAK ↔ PUMP ↔ INSTRUMENT ↔ DOSINGPUMP")
                print("=" * 80)
                
                return "Complete projection process with all 7 main processes has been executed successfully!"
                
            except Exception as e:
                error_msg = f"Error during complete projection: {str(e)}"
                print(f"\n🚨 CRITICAL ERROR:")
                print("=" * 80)
                print(f"❌ {error_msg}")
                print("=" * 80)
                import traceback
                traceback.print_exc()
                
                if progress_callback:
                    progress_callback(100, f"Error: {str(e)}")
                return error_msg
        
        # Show loading screen
        loading_screen = LoadingScreen(
            parent=self,
            title="Running Complete Advanced Projection",
            message="Processing data through ANAPAK, PUMP, and INSTRUMENT modules with all 5 main processes..."
        )
        loading_screen.show()
        loading_screen.start_loading(projection_process)
        
        # Connect completion handler
        def on_projection_complete(success, message):
            if success:
                print("\n🎯 COMPLETE PROJECTION FINISHED - UI NOTIFICATION")
                print("-" * 50)
                QMessageBox.information(
                    self, 
                    "Complete Projection Finished", 
                    "Complete advanced projection process with all 5 main processes has been completed successfully. All data has been processed through ANAPAK, PUMP, and INSTRUMENT modules."
                )
                print("🔄 Refreshing UI display...")
                # Refresh display
                self.load_excel_data()
                print("✅ UI refresh completed")
            else:
                print(f"\n🚨 COMPLETE PROJECTION FAILED - UI NOTIFICATION")
                print("-" * 50)
                print(f"❌ Error message: {message}")
                QMessageBox.critical(
                    self, 
                    "Complete Projection Error", 
                    f"An error occurred during complete projection:\n\n{message}"
                )
            
            self.statusBar().clearMessage()
            print("🏁 Complete projection process fully finished")
        
        loading_screen.worker.task_completed.connect(on_projection_complete)
            
    def run_generate_proposal(self):
        """Fungsi untuk menjalankan generate proposal dengan loading screen dan dynamic filename"""
        
        def generate_proposal_process(progress_callback=None):
            try:
                if progress_callback:
                    progress_callback(10, "Checking file paths...")
                
                # Use the customer-specific Excel file
                if not hasattr(self, 'customer_name') or not self.customer_name:
                    excel_path = self.excel_path
                    output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    customer_name = None
                else:
                    from modules.fix_customer_system import clean_folder_name
                    
                    customer_folder = os.path.join(
                        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "data", "customers", clean_folder_name(self.customer_name)
                    )
                    
                    excel_path = os.path.join(customer_folder, "SET_BDU.xlsx")
                    output_dir = customer_folder
                    customer_name = self.customer_name
                
                # Path to the Word template
                template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                                        "data", "Trial WWTP ANP Quotation Template.docx")
                
                if progress_callback:
                    progress_callback(20, "Getting selected user code...")
                
                # Ambil user code yang dipilih dari dropdown
                selected_user_code = None
                if hasattr(self, 'user_code_dropdown') and self.user_code_dropdown:
                    current_selection = self.user_code_dropdown.currentText()
                    if current_selection and current_selection not in ["-- Select User Code --", "-- No User Codes Available --"]:
                        selected_user_code = current_selection
                
                if progress_callback:
                    progress_callback(25, "Preparing dynamic filename...")
                
                # Preview nama file yang akan dibuat
                try:
                    import openpyxl
                    from modules.generate_proposal import generate_dynamic_filename
                    
                    wb_preview = openpyxl.load_workbook(excel_path, data_only=True)
                    preview_filename = generate_dynamic_filename(
                        wb_preview, 
                        fallback_customer_name=customer_name, 
                        version="01"
                    )
                    wb_preview.close()
                    
                    if progress_callback:
                        progress_callback(30, f"Will create: {preview_filename}")
                    
                    print(f"Preview filename: {preview_filename}")
                    
                except Exception as e:
                    print(f"Error creating filename preview: {str(e)}")
                    if progress_callback:
                        progress_callback(30, "Preparing proposal...")
                
                if progress_callback:
                    progress_callback(35, "Loading generate_proposal module...")
                
                # Import the generate_proposal module
                import importlib.util
                spec = importlib.util.spec_from_file_location(
                    "generate_proposal", 
                    os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                                "modules", "generate_proposal.py")
                )
                generate_proposal = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(generate_proposal)
                
                if progress_callback:
                    progress_callback(40, "Processing Excel data...")
                
                if progress_callback:
                    progress_callback(50, "Creating Word document from template...")
                
                # Call the updated function with directory instead of full path
                success, output_path = generate_proposal.generate_proposal(
                    excel_path=excel_path,
                    template_path=template_path, 
                    output_dir=output_dir,
                    selected_user_code=selected_user_code,
                    customer_name=customer_name,
                    version="01"  # Bisa dibuat dynamic jika diperlukan
                )
                
                if progress_callback:
                    progress_callback(95, "Finalizing document...")
                
                if success:
                    if progress_callback:
                        progress_callback(100, "Proposal generated successfully!")
                    
                    # Extract just the filename from the full path
                    filename = os.path.basename(output_path)
                    
                    result_message = f"Proposal successfully generated!\n\n"
                    result_message += f"File: {filename}\n"
                    result_message += f"Location: {output_path}\n"
                    
                    if selected_user_code:
                        result_message += f"\nContact Person: User Code {selected_user_code}"
                    
                    return result_message
                else:
                    return "Failed to generate proposal. Check the console for more details."
                
            except Exception as e:
                if progress_callback:
                    progress_callback(100, f"Error: {str(e)}")
                return f"An error occurred while generating proposal: {str(e)}"
        
        # Show loading screen
        from views.loading_screen import LoadingScreen
        loading_screen = LoadingScreen(
            parent=self,
            title="Generating Proposal",
            message="Creating proposal document with dynamic filename..."
        )
        loading_screen.show()
        loading_screen.start_loading(generate_proposal_process)
        
        # Connect completion handler
        def on_proposal_complete(success, message):
            if success and "successfully generated" in message:
                from PyQt5.QtWidgets import QMessageBox
                QMessageBox.information(
                    self,
                    "Proposal Generated Successfully",
                    message
                )
                # Refresh display to show the new file
                self.load_excel_data()
            else:
                from PyQt5.QtWidgets import QMessageBox
                QMessageBox.critical(
                    self,
                    "Error",
                    message
                )
            
            self.statusBar().clearMessage()
        
        loading_screen.worker.task_completed.connect(on_proposal_complete)

    def preview_proposal_filename(self):
        """Preview nama file proposal yang akan dibuat"""
        try:
            # Get necessary data
            excel_path = self.excel_path
            customer_name = getattr(self, 'customer_name', None)
            
            # Get selected user code
            selected_user_code = None
            if hasattr(self, 'user_code_dropdown') and self.user_code_dropdown:
                current_selection = self.user_code_dropdown.currentText()
                if current_selection and current_selection not in ["-- Select User Code --", "-- No User Codes Available --"]:
                    selected_user_code = current_selection
            
            # Generate preview filename
            import openpyxl
            from modules.generate_proposal import generate_dynamic_filename, clean_filename
            
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            
            # Simulate the user code being saved to DATA_TEMP.B1
            if selected_user_code and 'DATA_TEMP' in wb.sheetnames:
                sheet = wb['DATA_TEMP']
                sheet['B1'] = selected_user_code
            
            preview_filename = generate_dynamic_filename(
                wb, 
                fallback_customer_name=customer_name, 
                version="01"
            )
            
            wb.close()
            
            return preview_filename
            
        except Exception as e:
            print(f"Error generating filename preview: {str(e)}")
            fallback_name = f"WWTP_Quotation_{customer_name or 'Result'}_Ver.01.docx"
            return clean_filename(fallback_name)

    def show_filename_preview(self):
        """Tampilkan preview nama file di UI"""
        try:
            filename = self.preview_proposal_filename()
            
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.information(
                self,
                "Filename Preview",
                f"Proposal akan disimpan dengan nama:\n\n{filename}"
            )
            
        except Exception as e:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.warning(
                self,
                "Preview Error",
                f"Cannot generate filename preview: {str(e)}"
            )
                    
    def on_generate_proposal_finished(self, success, output):
        """Handler ketika proses generate proposal selesai"""
        if success:
            QMessageBox.information(
                self,
                "Success",
                "The proposal has been successfully generated. Please refresh the page to view the result."
            )
            print("Output generate-proposal.py:", output)
        else:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to generate the proposal. {output}"
            )
            print("Error generate-proposal.py:", output)
        
        # Hapus pesan status
        self.statusBar().clearMessage()

    def process_proposal_document(self, file_path):
        try:
            # Pastikan path adalah absolut
            abs_file_path = self.get_absolute_path(file_path)
            
            # Periksa apakah file ada
            if not os.path.exists(abs_file_path):
                print(f"Word document not found: {abs_file_path}")
                return False
                    
            # Dapatkan ekstensi file
            _, ext = os.path.splitext(abs_file_path)
            ext = ext.lower()
            
            # Periksa apakah itu dokumen Word
            if ext not in ['.docx', '.doc']:
                print(f"Not a Word document:: {abs_file_path}")
                return False

            # Simpan path dokumen untuk referensi
            self.proposal_document_path = abs_file_path
                    
            # Buat widget untuk menampilkan dokumen
            document_widget = QWidget()
            document_layout = QVBoxLayout(document_widget)
            document_layout.setContentsMargins(0, 0, 0, 0)
            
            # Buat widget info file - layout lebih kompak 
            file_header = QWidget()
            file_header.setStyleSheet("background-color: white; border: 1px solid #ddd; border-radius: 4px;")
            file_header_layout = QHBoxLayout(file_header)
            file_header_layout.setContentsMargins(10, 5, 10, 5)
            
            # Ikon file
            file_icon = QLabel("📄")
            file_icon.setFont(QFont("Segoe UI", 14))
            file_icon.setStyleSheet("color: #3498DB;")
            
            # Nama file
            file_name = os.path.basename(file_path)
            file_label = QLabel(f"{file_name}")
            file_label.setFont(QFont("Segoe UI", 11))
            file_label.setStyleSheet("color: #333;")
            
            # Tombol-tombol menu
            run_projection_btn = QPushButton("Re-run Projection")
            run_projection_btn.setFont(QFont("Segoe UI", 10))
            run_projection_btn.setCursor(Qt.PointingHandCursor)
            run_projection_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: #3498db;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 5px 10px;
                    margin-right: 5px;
                }}
                QPushButton:hover {{
                    background-color: #2980b9;
                }}
            """)
            run_projection_btn.clicked.connect(self.run_projection)
            
            generate_btn = QPushButton("Regenerate Proposal")
            generate_btn.setFont(QFont("Segoe UI", 10))
            generate_btn.setCursor(Qt.PointingHandCursor)
            generate_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: #27ae60;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 5px 10px;
                    margin-right: 5px;
                }}
                QPushButton:hover {{
                    background-color: #2ecc71;
                }}
            """)
            generate_btn.clicked.connect(self.run_generate_proposal)
            
            open_btn = QPushButton("Open in Word")
            open_btn.setFont(QFont("Segoe UI", 10))
            open_btn.setCursor(Qt.PointingHandCursor)
            open_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {SECONDARY_COLOR};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 5px 10px;
                }}
                QPushButton:hover {{
                    background-color: #2980B9;
                }}
            """)
            open_btn.clicked.connect(lambda: self.open_document(file_path))
            
            # Tambahkan widget ke layout header
            file_header_layout.addWidget(file_icon)
            file_header_layout.addWidget(file_label)
            file_header_layout.addStretch()
            file_header_layout.addWidget(run_projection_btn)
            file_header_layout.addWidget(generate_btn)
            file_header_layout.addWidget(open_btn)
            
            document_layout.addWidget(file_header)
            
            # Alternatif untuk pratinjau: Pesan sederhana dengan ikon dokumen
            preview_container = QWidget()
            preview_layout = QVBoxLayout(preview_container)
            preview_layout.setAlignment(Qt.AlignCenter)
            preview_container.setStyleSheet("""
                background-color: #f9f9f9;
                border: 1px solid #ddd;
                border-radius: 4px;
            """)
            
            # Info pesan
            info_label = QLabel("Preview tidak tersedia.")
            info_label.setFont(QFont("Segoe UI", 12))
            info_label.setAlignment(Qt.AlignCenter)
            info_label.setStyleSheet("color: #666; margin: 20px;")
            
            # Ikon dokumen besar
            doc_icon = QLabel("📄")
            doc_icon.setFont(QFont("Segoe UI", 48))
            doc_icon.setAlignment(Qt.AlignCenter)
            doc_icon.setStyleSheet("color: #3498DB; margin: 20px;")
            
            # Pesan untuk membuka di Word
            tip_label = QLabel("Gunakan tombol 'Open in Word' untuk melihat dokumen.")
            tip_label.setFont(QFont("Segoe UI", 10))
            tip_label.setAlignment(Qt.AlignCenter)
            tip_label.setStyleSheet("color: #666; margin: 10px;")
            
            # Tambahkan ke container
            preview_layout.addWidget(info_label)
            preview_layout.addWidget(doc_icon)
            preview_layout.addWidget(tip_label)
            
            document_layout.addWidget(preview_container, 1)  # Stretch factor 1
            
            # Coba gunakan metode thumbnail generator jika tersedia modul yang diperlukan
            try:
                # Cek apakah library yang diperlukan tersedia
                have_required_modules = False
                try:
                    import pythoncom
                    import docx2pdf
                    import fitz
                    have_required_modules = True
                except ImportError:
                    have_required_modules = False
                
                if have_required_modules:
                    # Membuat area scroll untuk thumbnail
                    scroll_area = QScrollArea()
                    scroll_area.setWidgetResizable(True)
                    scroll_area.setStyleSheet("""
                        QScrollArea {
                            border: 1px solid #ddd;
                            background-color: #f9f9f9;
                            border-radius: 4px;
                        }
                    """)
                    
                    # Container untuk thumbnail
                    thumbnail_container = QWidget()
                    thumbnail_layout = QVBoxLayout(thumbnail_container)
                    thumbnail_layout.setAlignment(Qt.AlignHCenter)
                    thumbnail_layout.setSpacing(20)
                    
                    # Label loading
                    loading_label = QLabel("Loading preview document...")
                    loading_label.setFont(QFont("Segoe UI", 12))
                    loading_label.setAlignment(Qt.AlignCenter)
                    loading_label.setStyleSheet("color: #666; margin: 20px;")
                    thumbnail_layout.addWidget(loading_label)
                    
                    # Simpan referensi layout thumbnail untuk digunakan nanti
                    self.thumbnail_layout = thumbnail_layout
                    self.loading_label = loading_label
                    
                    # Tambahkan container ke area gulir
                    scroll_area.setWidget(thumbnail_container)
                    document_layout.removeWidget(preview_container)
                    preview_container.deleteLater()
                    document_layout.addWidget(scroll_area, 1)  # Stretch factor 1
                    
                    # Inisialisasi COM library untuk main thread
                    pythoncom.CoInitialize()
                    
                    # Mulai thread generator thumbnail
                    self.thumbnail_thread = ThumbnailGeneratorThread(file_path, max_pages=100)
                    self.thumbnail_thread.thumbnail_ready.connect(self.add_thumbnail)
                    self.thumbnail_thread.finished.connect(self.generation_finished)
                    self.thumbnail_thread.error.connect(self.generation_error)
                    self.thumbnail_thread.start()
                else:
                    # Jika modul tidak tersedia, gunakan pratinjau sederhana
                    print("Library yang diperlukan tidak tersedia. Menggunakan pratinjau sederhana.")
            except Exception as e:
                print(f"Error saat membuat pratinjau: {e}")
            
            # Simpan widget untuk digunakan nanti
            self.proposal_document_widget = document_widget
            
            return True
                
        except Exception as e:
            print(f"Error memproses dokumen Word: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
            
    def add_thumbnail(self, page_num, pixmap):
        """Menambahkan thumbnail ke UI"""
        try:
            # Hapus label loading saat thumbnail pertama ditambahkan
            if page_num == 0 and self.loading_label:
                self.loading_label.setVisible(False)
            
            # Buat frame untuk thumbnail
            thumbnail_frame = QFrame()
            thumbnail_frame.setFrameShape(QFrame.StyledPanel)
            thumbnail_frame.setStyleSheet("""
                QFrame {
                    border: 1px solid #ddd;
                    border-radius: 4px;
                    background-color: white;
                    padding: 10px;
                }
            """)
            frame_layout = QVBoxLayout(thumbnail_frame)
            
            # Tambahkan label halaman
            page_label = QLabel(f"Page {page_num + 1}")
            page_label.setFont(QFont("Segoe UI", 10, QFont.Bold))
            page_label.setAlignment(Qt.AlignCenter)
            frame_layout.addWidget(page_label)
            
            # Skala thumbnail untuk fit dalam lebar 600px
            scaled_pixmap = pixmap.scaledToWidth(600, Qt.SmoothTransformation)
            
            # Buat label untuk thumbnail
            thumbnail_label = QLabel()
            thumbnail_label.setPixmap(scaled_pixmap)
            thumbnail_label.setAlignment(Qt.AlignCenter)
            thumbnail_label.setStyleSheet("background-color: white;")
            frame_layout.addWidget(thumbnail_label)
            
            # Tambahkan thumbnail ke layout
            self.thumbnail_layout.addWidget(thumbnail_frame)
        except Exception as e:
            print(f"Error menambahkan thumbnail: {str(e)}")

    def generation_finished(self):
        """Dipanggil saat generasi thumbnail selesai"""
        try:
            # Hapus label loading jika masih ada
            if self.loading_label:
                self.loading_label.setVisible(False)
                
        except Exception as e:
            print(f"Error dalam generation_finished: {str(e)}")

    def generation_error(self, error_message):
        """Dipanggil jika ada error saat generasi thumbnail"""
        try:
            if self.loading_label:
                self.loading_label.setText(f"Error saat memuat pratinjau: {error_message}")
                self.loading_label.setStyleSheet("color: #E74C3C; margin: 20px;")
        except Exception as e:
            print(f"Error dalam generation_error: {str(e)}")

    def open_document(self, file_path):
        """Membuka dokumen di aplikasi aslinya"""
        try:
            # Konversi ke path absolut
            abs_file_path = self.get_absolute_path(file_path)
                
            # Untuk Windows
            if sys.platform == 'win32':
                os.startfile(abs_file_path)
            # Untuk macOS
            elif sys.platform == 'darwin':
                subprocess.call(('open', abs_file_path))
            # Untuk Linux
            else:
                subprocess.call(('xdg-open', abs_file_path))
                    
            print(f"Membuka dokumen: {abs_file_path}")
        except Exception as e:
            QMessageBox.critical(
                self, 
                "Error Membuka File", 
                f"Tidak dapat membuka file: {str(e)}"
            )
            print(f"Error membuka dokumen: {str(e)}")
                            
    def process_excel_images(self, sheet_name, layout):
        """Extract and display images from Excel sheet"""
        try:
            from openpyxl import load_workbook
            from openpyxl.drawing.image import Image
            from io import BytesIO
            from PIL import Image as PILImage
            
            # Load workbook
            wb = load_workbook(self.excel_path)
            if sheet_name not in wb.sheetnames:
                return False
                
            sheet = wb[sheet_name]
            
            # Create a frame for images
            images_frame = QWidget()
            images_layout = QVBoxLayout(images_frame)
            images_layout.setContentsMargins(10, 10, 10, 10)
            
            # Add a title for images section
            images_title = QLabel("Diagrams and Images")
            images_title.setFont(QFont("Segoe UI", 14, QFont.Bold))
            images_title.setStyleSheet(f"color: {PRIMARY_COLOR};")
            images_layout.addWidget(images_title)
            
            # Track if we found any images
            found_images = False
            
            # Process all images in the sheet
            for image in sheet._images:
                found_images = True
                
                # Create a label to display the image
                img_label = QLabel()
                img_label.setAlignment(Qt.AlignCenter)
                img_label.setStyleSheet("background-color: white; border: 1px solid #ddd; padding: 10px;")
                
                # Extract image data
                img_data = image._data()
                
                # Convert to QPixmap and set to label
                pixmap = QPixmap()
                pixmap.loadFromData(img_data)
                
                # Scale image if too large
                if pixmap.width() > 800:
                    pixmap = pixmap.scaledToWidth(800, Qt.SmoothTransformation)
                    
                img_label.setPixmap(pixmap)
                images_layout.addWidget(img_label)
                
                # Add some spacing between images
                images_layout.addSpacing(20)
            
            # Add the images frame to the main layout if we found any
            if found_images:
                layout.addWidget(images_frame)
                return True
            
            return False
        except Exception as e:
            print(f"Error processing images from sheet {sheet_name}: {str(e)}")
            return False
    
    def process_sheet_data(self, df, sheet_name, layout):
        """Process the data from a sheet and create UI elements in a grid layout similar to Excel"""
        # Check if the sheet is a DATA sheet (just display as a table)
        if sheet_name.startswith("DATA_"):
            self.create_data_table(df, layout)
            self.process_excel_images(sheet_name, layout)
            return

        # For DIP sheets or other sheets, process as forms
        # Initialize variables
        current_section = None
        section_layout = None
        section_grid = None
        current_row = 0  # Track the current row in the grid
        current_header_labels = []  # For storing column headers from ch_
        has_column_headers = False
        
        # Track sections by column position
        left_section = None  # Main left section (columns 0-1)
        right_section = None  # Right section (columns 2-3)
        right_section_title = None

        # Field identification
        field_count = 0
        
        # Variables to track industry dropdown fields
        industry_dropdown = None
        sub_industry_dropdown = None
        
        # Untuk melacak pasangan field dropdown province-city
        province1_dropdown = None
        city1_dropdown = None
        province2_dropdown = None
        city2_dropdown = None
        province1_field_key = None
        city1_field_key = None
        province2_field_key = None
        city2_field_key = None
        
        # Untuk melacak pasangan field dropdown industry-subindustry
        industry_field_key = None
        sub_industry_field_key = None
        
        # Variabel untuk melacak pasangan field dropdown pump (brand-type-model)
        pump_brand_dropdown = None
        pump_type_dropdown = None  
        pump_model_dropdown = None
        pump_brand_field_key = None
        pump_type_field_key = None
        pump_model_field_key = None
        
        # Create a section for table if we find table formatting
        table_section = None
        table_grid = None
        in_table = False
        table_row = 0

        # Dictionary to track rowspans
        rowspans = {}
        tdm_counter = 0  # Counter for tdm_ cells in current row
        last_header = None  # Track last header for rowspans
        current_tdm_row = 0  # Track current row for tdm counters

        # Check if the dataframe is empty
        if df.empty:
            empty_label = QLabel("No data found in this sheet.")
            empty_label.setAlignment(Qt.AlignCenter)
            empty_label.setStyleSheet("color: #666; margin: 20px;")
            layout.addWidget(empty_label)
            return

        # First pass: pre-scan for column headers
        # This will help us collect all ch_ headers before processing fields
        for index, row in df.iterrows():
            # Skip empty rows
            if pd.isna(row).all():
                continue
            
            first_col = row.iloc[0] if not pd.isna(row.iloc[0]) else ""
            if not isinstance(first_col, str):
                try:
                    first_col = str(first_col)
                except:
                    continue
                    
            # If this is a 'Contact' type row with headers
            if first_col.startswith('fh_') and any(isinstance(cell, str) and cell.startswith('ch_') for cell in row if not pd.isna(cell)):
                # Process all cells in this row to find ch_ headers
                header_row_labels = []
                for col_idx in range(df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = str(row[col_idx]).strip() if isinstance(row[col_idx], str) else ""
                        if col_value.startswith('ch_'):
                            header_text = col_value[3:].strip()  # Remove 'ch_' prefix
                            header_row_labels.append(header_text)
                
                if len(header_row_labels) > 0:
                    current_header_labels = header_row_labels
                    has_column_headers = True
                    
                # Also look for any right headers (fh_) in the same row
                for col_idx in range(1, df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = str(row[col_idx]).strip() if isinstance(row[col_idx], str) else ""
                        if col_value.startswith('fh_'):
                            right_header_text = col_value[3:].strip()  # Remove 'fh_' prefix
                            # If this is first header for right section, treat it as section title
                            if right_section is None:
                                right_section = right_header_text

        # Process each row
        for index, row in df.iterrows():
            # Skip empty rows
            if pd.isna(row).all():
                continue
            
            # Get the first column to determine the type
            first_col = row.iloc[0] if not pd.isna(row.iloc[0]) else ""

            # Convert to string for startswith checks, but only if we're not in a DATA_ sheet
            if not isinstance(first_col, str):
                try:
                    first_col = str(first_col)
                except:
                    continue  # Skip if can't convert to string
            
            # Ambil nilai kolom A (indeks 0)
            col_a = row.iloc[0] if not pd.isna(row.iloc[0]) else ""
            
            # Ubah ke string untuk pemeriksaan dengan startswith
            if not isinstance(col_a, str):
                try:
                    col_a = str(col_a)
                except:
                    continue  # Lewati jika tidak bisa dikonversi ke string
            
            # Pemrosesan header (th_ atau thr_) di kolom A
            if col_a.startswith('th_') or col_a.startswith('thr_'):
                # Jika kita sedang dalam tabel tapi ini header baru, reset penghitung
                if in_table and table_grid is None:
                    # Buat struktur tabel baru
                    table_frame = QWidget()
                    table_frame.setStyleSheet("background-color: white; border: 1px solid #ddd;")
                    
                    table_section = QVBoxLayout(table_frame)
                    table_section.setContentsMargins(0, 0, 0, 0)
                    table_section.setSpacing(0)
                    
                    table_grid = QGridLayout()
                    table_grid.setSpacing(0)
                    table_grid.setContentsMargins(0, 0, 0, 0)
                    
                    section_layout.addWidget(table_frame)
                    table_section.addLayout(table_grid)
                    
                    table_row = 0
                elif not in_table:
                    # Buat tabel baru jika belum ada
                    in_table = True
                    
                    # Buat section untuk tabel jika belum ada
                    if section_layout is None:
                        section_frame = QWidget()
                        section_frame.setStyleSheet("background-color: white;")
                        section_layout = QVBoxLayout(section_frame)
                        section_layout.setContentsMargins(15, 15, 15, 15)
                        section_layout.setSpacing(15)
                        layout.addWidget(section_frame)
                    
                    # Buat frame tabel
                    table_frame = QWidget()
                    table_frame.setStyleSheet("background-color: white; border: 1px solid #ddd;")
                    
                    table_section = QVBoxLayout(table_frame)
                    table_section.setContentsMargins(0, 0, 0, 0)
                    table_section.setSpacing(0)
                    
                    # Buat grid untuk tabel
                    table_grid = QGridLayout()
                    table_grid.setSpacing(0)
                    table_grid.setContentsMargins(0, 0, 0, 0)
                    
                    # Tambahkan ke section
                    section_layout.addWidget(table_frame)
                    table_section.addLayout(table_grid)
                    
                    # Reset penghitung tabel
                    table_row = 0
                
                # Header biasa (th_)
                if col_a.startswith('th_'):
                    header_text = col_a[3:].strip()  # Hapus awalan 'th_'
                    header_label = QLabel(header_text)
                    header_label.setFont(QFont("Segoe UI", 11, QFont.Bold))
                    header_label.setStyleSheet("""
                        background-color: #f5f5f5; 
                        color: #333;
                        padding: 8px;
                        border: 1px solid #ddd;
                    """)
                    header_label.setMinimumWidth(250)
                    
                    # Tambahkan ke grid
                    table_grid.addWidget(header_label, table_row, 0)
                    
                    # Cek kolom B (indeks 1) untuk data
                    if len(row) > 1 and not pd.isna(row.iloc[1]):
                        col_b = row.iloc[1]
                        
                        # Ubah ke string untuk pemeriksaan
                        if not isinstance(col_b, str):
                            try:
                                col_b = str(col_b)
                            except:
                                col_b = ""
                        
                        # Proses data berdasarkan jenisnya
                        if col_b.startswith('td_'):
                            # Data biasa
                            data_text = col_b[3:].strip()  # Hapus awalan 'td_'
                            data_label = QLabel(data_text)
                            data_label.setFont(QFont("Segoe UI", 11))
                            data_label.setStyleSheet("""
                                padding: 8px;
                                border: 1px solid #ddd;
                            """)
                            data_label.setWordWrap(True)
                            
                            # Tambahkan ke grid, span 2 kolom
                            table_grid.addWidget(data_label, table_row, 1, 1, 2)
                        
                        elif col_b.startswith('tdi_'):
                            # Data dengan input field
                            data_text = col_b[4:].strip()  # Hapus awalan 'tdi_'
                            
                            # Cek placeholder seperti $P1$
                            import re
                            placeholders = re.findall(r'\$P\d+\$', data_text)
                            
                            if placeholders:
                                # Buat container untuk input field + teks
                                container = QWidget()
                                container_layout = QHBoxLayout(container)
                                container_layout.setContentsMargins(8, 8, 8, 8)
                                
                                # Pisah teks berdasarkan placeholder
                                parts = re.split(r'(\$P\d+\$)', data_text)
                                
                                for part in parts:
                                    if re.match(r'\$P\d+\$', part):
                                        # Ini placeholder, buat input field
                                        input_field = QLineEdit()
                                        input_field.setFixedWidth(50)
                                        input_field.setStyleSheet("""
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                        """)
                                        container_layout.addWidget(input_field)
                                        
                                        # Daftarkan field di data_fields
                                        field_key = f"tdi_{part}_{header_text}"
                                        self.data_fields[field_key] = input_field
                                    else:
                                        # Teks biasa
                                        text_label = QLabel(part)
                                        text_label.setFont(QFont("Segoe UI", 11))
                                        container_layout.addWidget(text_label)
                                
                                container_layout.addStretch()
                                
                                # Tambahkan ke grid, span 2 kolom
                                table_grid.addWidget(container, table_row, 1, 1, 2)
                            else:
                                # Tidak ada placeholder, teks biasa
                                data_label = QLabel(data_text)
                                data_label.setFont(QFont("Segoe UI", 11))
                                data_label.setStyleSheet("""
                                    padding: 8px;
                                    border: 1px solid #ddd;
                                """)
                                data_label.setWordWrap(True)
                                
                                # Tambahkan ke grid, span 2 kolom
                                table_grid.addWidget(data_label, table_row, 1, 1, 2)
                        
                        elif col_b.startswith('tdm_'):
                            # Multi-kolom data (2 kolom)
                            data_text_b = col_b[4:].strip()  # Hapus awalan 'tdm_'
                            
                            # Buat label untuk kolom pertama (B)
                            data_label_b = QLabel(data_text_b)
                            data_label_b.setFont(QFont("Segoe UI", 11))
                            data_label_b.setStyleSheet("""
                                padding: 8px;
                                border: 1px solid #ddd;
                            """)
                            data_label_b.setWordWrap(True)
                            
                            # Tambahkan ke grid di kolom 1
                            table_grid.addWidget(data_label_b, table_row, 1)
                            
                            # Cek kolom C (indeks 2) untuk data kedua
                            if len(row) > 2 and not pd.isna(row.iloc[2]):
                                col_c = row.iloc[2]
                                
                                # Ubah ke string untuk pemeriksaan
                                if not isinstance(col_c, str):
                                    try:
                                        col_c = str(col_c)
                                    except:
                                        col_c = ""
                                
                                if col_c.startswith('tdm_'):
                                    # Multi-kolom data kedua
                                    data_text_c = col_c[4:].strip()  # Hapus awalan 'tdm_'
                                    
                                    # Buat label untuk kolom kedua (C)
                                    data_label_c = QLabel(data_text_c)
                                    data_label_c.setFont(QFont("Segoe UI", 11))
                                    data_label_c.setStyleSheet("""
                                        padding: 8px;
                                        border: 1px solid #ddd;
                                    """)
                                    data_label_c.setWordWrap(True)
                                    
                                    # Tambahkan ke grid di kolom 2
                                    table_grid.addWidget(data_label_c, table_row, 2)
                        
                        else:
                            # Data tanpa awalan khusus, tampilkan sebagai teks biasa
                            data_text = str(col_b).strip()
                            data_label = QLabel(data_text)
                            data_label.setFont(QFont("Segoe UI", 11))
                            data_label.setStyleSheet("""
                                padding: 8px;
                                border: 1px solid #ddd;
                            """)
                            data_label.setWordWrap(True)
                            
                            # Tambahkan ke grid, span 2 kolom
                            table_grid.addWidget(data_label, table_row, 1, 1, 2)
                    
                    # Simpan header terakhir
                    last_header = header_text
                    table_row += 1
                
                # Untuk header dengan rowspan (thr_)
                elif col_a.startswith('thr_'):
                    header_text = col_a[4:].strip()  # Hapus awalan 'thr_'
                    header_label = QLabel(header_text)
                    header_label.setFont(QFont("Segoe UI", 11, QFont.Bold))
                    header_label.setStyleSheet("""
                        background-color: #f5f5f5; 
                        color: #333;
                        padding: 8px;
                        border: 1px solid #ddd;
                    """)
                    header_label.setMinimumWidth(250)
                    header_label.setAlignment(Qt.AlignLeft | Qt.AlignTop)
                    
                    # Hitung rowspan dengan melihat baris-baris berikutnya
                    rowspan = 0
                    tdm_count = 0
                    current_row = index
                    
                    # Simpan baris-baris data yang akan diproses nanti
                    data_rows = []
                    
                    # Periksa baris-baris berikutnya sampai ketemu header baru
                    while current_row + 1 < len(df):
                        next_row = current_row + 1
                        next_row_data = df.iloc[next_row]
                        next_col_a = next_row_data.iloc[0] if not pd.isna(next_row_data.iloc[0]) else ""
                        
                        # Ubah ke string jika perlu
                        if not isinstance(next_col_a, str):
                            try:
                                next_col_a = str(next_col_a)
                            except:
                                next_col_a = ""
                        
                        # Jika ketemu header baru, berhenti
                        if next_col_a.startswith('th_') or next_col_a.startswith('thr_'):
                            break
                        
                        # Periksa kolom B untuk data, meskipun kolom A kosong
                        if len(next_row_data) > 1:
                            next_col_b = next_row_data.iloc[1] if not pd.isna(next_row_data.iloc[1]) else ""
                            
                            # Ubah ke string jika perlu
                            if not isinstance(next_col_b, str):
                                try:
                                    next_col_b = str(next_col_b)
                                except:
                                    next_col_b = ""
                            
                            # Jika ada data di kolom B, tambah rowspan dan simpan data
                            if (next_col_b.startswith('td_') or 
                                next_col_b.startswith('tdi_') or
                                next_col_b.startswith('tdm_')):
                                rowspan += 1
                                data_rows.append(next_row_data)
                        
                        current_row = next_row
                    
                    # Perlu ditambahkan 1 untuk baris pertama
                    rowspan = max(1, rowspan + 1)
                    
                    # Tambahkan baris saat ini ke daftar data
                    data_rows.insert(0, row)
                    
                    # Simpan informasi rowspan
                    rowspans[table_row] = {
                        'header': header_text,
                        'rowspan': rowspan,
                        'widget': header_label
                    }
                    
                    # Tambahkan ke grid dengan rowspan
                    table_grid.addWidget(header_label, table_row, 0, rowspan, 1)
                    
                    # Proses semua data dalam rowspan
                    current_data_row = table_row
                    
                    # Proses semua baris data yang telah dikumpulkan
                    for data_row in data_rows:
                        # Cek kolom B untuk data
                        if len(data_row) > 1 and not pd.isna(data_row.iloc[1]):
                            col_b = data_row.iloc[1]
                            
                            # Ubah ke string jika perlu
                            if not isinstance(col_b, str):
                                try:
                                    col_b = str(col_b)
                                except:
                                    col_b = ""
                            
                            # Proses berdasarkan jenis data
                            if col_b.startswith('td_'):
                                data_text = col_b[3:].strip()
                                data_label = QLabel(data_text)
                                data_label.setFont(QFont("Segoe UI", 11))
                                data_label.setStyleSheet("""
                                    padding: 8px;
                                    border: 1px solid #ddd;
                                """)
                                data_label.setWordWrap(True)
                                table_grid.addWidget(data_label, current_data_row, 1, 1, 2)
                                current_data_row += 1
                            
                            # Untuk memproses tdi_
                            elif col_b.startswith('tdi_'):
                                data_text = col_b[4:].strip()  # Hapus awalan 'tdi_'
                                
                                import re
                                placeholders = re.findall(r'\$P\d+\$', data_text)
                                
                                if placeholders:
                                    # Buat container untuk input field + teks
                                    container = QWidget()
                                    container_layout = QHBoxLayout(container)
                                    container_layout.setContentsMargins(8, 8, 8, 8)
                                    
                                    # Split teks berdasarkan placeholder
                                    parts = re.split(r'(\$P\d+\$)', data_text)
                                    
                                    for part in parts:
                                        if re.match(r'\$P\d+\$', part):
                                            # Ini placeholder, buat input field
                                            input_field = QLineEdit()
                                            input_field.setFixedWidth(50)
                                            input_field.setStyleSheet("""
                                                padding: 5px;
                                                border: 1px solid #ccc;
                                                border-radius: 4px;
                                            """)
                                            
                                            # Validator untuk membatasi input hanya angka 0-100
                                            from PyQt5.QtGui import QIntValidator
                                            validator = QIntValidator(0, 100)
                                            input_field.setValidator(validator)
                                            
                                            # Set nilai default berdasarkan placeholder
                                            if part == "$P1$":
                                                input_field.setText("30")
                                            elif part == "$P2$":
                                                input_field.setText("50")
                                            elif part == "$P3$":
                                                input_field.setText("15")
                                            elif part == "$P4$":
                                                input_field.setText("5")
                                            
                                            container_layout.addWidget(input_field)
                                            
                                            # Daftarkan field di data_fields
                                            field_key = f"tdi_{part}_{header_text}_{current_data_row}"
                                            self.data_fields[field_key] = input_field
                                        else:
                                            # Teks biasa
                                            text_label = QLabel(part)
                                            text_label.setFont(QFont("Segoe UI", 11))
                                            # Tidak ada border pada teks
                                            container_layout.addWidget(text_label)
                                    
                                    container_layout.addStretch()
                                    
                                    # Container tidak memiliki border
                                    container.setStyleSheet("background-color: transparent; border: none;")
                                    
                                    # Tambahkan ke grid, span seluruh kolom
                                    table_grid.addWidget(container, current_data_row, 1, 1, 2)
                                else:
                                    # Tidak ada placeholder, tampilkan sebagai teks biasa
                                    data_label = QLabel(data_text)
                                    data_label.setFont(QFont("Segoe UI", 11))
                                    # Tidak ada border pada teks
                                    data_label.setStyleSheet("padding: 8px; border: none;")
                                    data_label.setWordWrap(True)
                                    
                                    # Tambahkan ke grid, span seluruh kolom
                                    table_grid.addWidget(data_label, current_data_row, 1, 1, 2)
                                
                                current_data_row += 1
                            
                            elif col_b.startswith('tdm_'):
                                data_text_b = col_b[4:].strip()
                                
                                data_label_b = QLabel(data_text_b)
                                data_label_b.setFont(QFont("Segoe UI", 11))
                                data_label_b.setStyleSheet("""
                                    padding: 8px;
                                    border: 1px solid #ddd;
                                """)
                                data_label_b.setWordWrap(True)
                                table_grid.addWidget(data_label_b, current_data_row, 1)
                                
                                # Cek kolom C untuk tdm_ pasangan
                                if len(data_row) > 2 and not pd.isna(data_row.iloc[2]):
                                    col_c = data_row.iloc[2]
                                    
                                    if not isinstance(col_c, str):
                                        try:
                                            col_c = str(col_c)
                                        except:
                                            col_c = ""
                                    
                                    if col_c.startswith('tdm_'):
                                        data_text_c = col_c[4:].strip()
                                        
                                        data_label_c = QLabel(data_text_c)
                                        data_label_c.setFont(QFont("Segoe UI", 11))
                                        data_label_c.setStyleSheet("""
                                            padding: 8px;
                                            border: 1px solid #ddd;
                                        """)
                                        data_label_c.setWordWrap(True)
                                        table_grid.addWidget(data_label_c, current_data_row, 2)
                                
                                current_data_row += 1
                            
                            else:
                                # Data biasa tanpa prefix
                                data_text = str(col_b).strip()
                                data_label = QLabel(data_text)
                                data_label.setFont(QFont("Segoe UI", 11))
                                data_label.setStyleSheet("""
                                    padding: 8px;
                                    border: 1px solid #ddd;
                                """)
                                data_label.setWordWrap(True)
                                table_grid.addWidget(data_label, current_data_row, 1, 1, 2)
                                current_data_row += 1
                    
                    # Update table_row setelah memproses semua data
                    table_row = current_data_row
                    
                    # Simpan header terakhir
                    last_header = header_text
            else:
                # Ini bukan header di kolom A, mungkin baris lanjutan atau bukan bagian tabel
                # Jika sebelumnya kita dalam tabel tapi sekarang tidak ada header/data yang sesuai,
                # tandai bahwa kita keluar dari tabel
                if in_table and not col_a:
                    in_table = False
                    table_section = None
                    table_grid = None
                
            # Check if it's a section header (sub_)
            if isinstance(first_col, str) and first_col.startswith('sub_'):
                # Create a new section
                section_title = first_col[4:].strip()  # Remove 'sub_' prefix

                # Create section frame - removed border
                section_frame = QWidget()
                section_frame.setStyleSheet("""
                    background-color: white;
                """)

                section_layout = QVBoxLayout(section_frame)
                section_layout.setContentsMargins(15, 15, 15, 15)
                section_layout.setSpacing(15)

                # Add section title - ensure it's aligned with the left margin, no indent
                title_label = QLabel(section_title)
                title_label.setFont(QFont("Segoe UI", 14, QFont.Bold))
                title_label.setStyleSheet(f"color: {PRIMARY_COLOR}; background-color: transparent;")
                title_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                title_label.setIndent(0)  # Ensure no text indentation

                section_layout.addWidget(title_label)
                
                # Create a grid for this section
                section_grid = QGridLayout()
                section_grid.setHorizontalSpacing(30)  # Increase horizontal spacing
                section_grid.setVerticalSpacing(10)
                section_grid.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                section_layout.addLayout(section_grid)

                # Add section to main layout
                layout.addWidget(section_frame)
                layout.addSpacing(20)

                current_section = section_title
                left_section = section_title  # Track as left section
                current_row = 0  # Reset row counter for new section
                field_count = 0

                # Also check if there are additional sections in this row (columns to the right)
                for col_idx in range(1, df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        right_col_value = str(row[col_idx]).strip() if isinstance(row[col_idx], str) else ""
                        if right_col_value.startswith('sub_'):
                            right_section_title = right_col_value[4:].strip()  # Remove 'sub_' prefix
                            right_section = right_section_title  # Track as right section
                            break

                continue
                
            # Ensure we have a section grid to add fields to
            if section_layout is None:
                # If no section is defined yet, create a default one
                section_frame = QWidget()
                section_frame.setStyleSheet("""
                    background-color: white;
                """)

                section_layout = QVBoxLayout(section_frame)
                section_layout.setContentsMargins(15, 15, 15, 15)
                section_layout.setSpacing(15)
                
                # Create a grid for this section
                section_grid = QGridLayout()
                section_grid.setHorizontalSpacing(30)  # Increase horizontal spacing
                section_grid.setVerticalSpacing(10)
                section_layout.addLayout(section_grid)

                # Add to main layout
                layout.addWidget(section_frame)
                current_section = "Default"
                current_row = 0
            
            # Check if it's a field header (fh_)
            if first_col.startswith('fh_'):
                field_header = first_col[3:].strip()  # Remove 'fh_' prefix
                
                # Field header label
                header_label = QLabel(field_header)
                header_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
                header_label.setStyleSheet("color: #555; margin-top: 5px;")
                
                # Crucial settings to prevent indentation
                header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                header_label.setIndent(0)  # Set text indent to 0
                
                # Add header to the grid - in column 0, same as regular field labels
                # Don't span columns to keep consistent alignment with field labels
                section_grid.addWidget(header_label, current_row, 0, 1, 1)
                
                # Check if there are column headers (ch_) in the same row
                has_ch_in_row = False
                ch_headers = []
                ch_columns = []  # Track the column positions of the headers
                
                for col_idx in range(1, df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            col_value = str(row[col_idx]).strip()
                        
                        if col_value.startswith('ch_'):
                            has_ch_in_row = True
                            ch_header_text = col_value[3:].strip()
                            ch_headers.append(ch_header_text)
                            ch_columns.append(col_idx)  # Save the actual column position
                
                # If we found ch_ headers in this row, update our current headers
                if has_ch_in_row:
                    current_header_labels = ch_headers
                    has_column_headers = True
                    
                    # Add the column headers to the grid in the same row as the field header
                    # This is the key change - align column headers horizontally with the field header
                    for i, header_text in enumerate(current_header_labels):
                        col_header_label = QLabel(header_text)
                        col_header_label.setFont(QFont("Segoe UI", 11, QFont.Bold))
                        col_header_label.setStyleSheet("color: #555; margin-top: 5px; padding-left: 0px;")
                        col_header_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                        col_header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                        col_header_label.setIndent(0)  # Prevent text indentation
                        
                        # Position headers at the same row as the field header
                        # For the first header (Name), put it at column 1
                        # For the second header (Phone No/Email), put it at column 2
                        section_grid.addWidget(col_header_label, current_row, i+1)
                
                # Check if there are fields or headers in columns C and beyond in the same row
                # Keep track of found right header to avoid duplicates
                right_header_found = False
                
                for col_idx in range(1, df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # Skip ch_ headers as we've already processed them
                        if col_value.startswith('ch_'):
                            continue
                            
                        # Check for any field type in the right columns - not just fh_
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fh_') or 
                            col_value.startswith('fm_')) and not right_header_found:
                            
                            # Handle right headers by prefix type
                            prefix = col_value[:2] if col_value.startswith('f_') else col_value[:3]
                            suffix = col_value[2:] if col_value.startswith('f_') else col_value[3:]
                            right_content = suffix.strip()
                            
                            right_header_found = True  # Mark that we found a right content
                            
                            # For header fields (fh_)
                            if col_value.startswith('fh_'):
                                # It's another header in the same row - this will be for the right section
                                right_header = right_content
                                
                                # If this is first header for right section, treat it as section title if we don't have one yet
                                if right_section is None:
                                    right_section = right_header
                                    
                                right_header_label = QLabel(right_header)
                                right_header_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
                                right_header_label.setStyleSheet("color: #555; margin-top: 5px;")
                                right_header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                                right_header_label.setIndent(0)
                                # Position in the grid correctly - at the same row level as the left header
                                # But in columns 3-4 (index 3-4) to create proper separation
                                section_grid.addWidget(right_header_label, current_row, 3, 1, 2)
                            else:
                                # Handle non-header fields in the right section immediately after a header in the left
                                right_section_name = right_section if right_section else current_section
                                right_field_key = f"{sheet_name}_{right_section_name}_{field_count}"
                                field_count += 1
                                
                                # Create right field label
                                right_label = QLabel(right_content)
                                right_label.setFont(QFont("Segoe UI", 11))
                                right_label.setStyleSheet("color: #333; background-color: transparent;")
                                right_label.setMinimumWidth(250)  # Set minimum width for consistent layout
                                
                                # Add label to grid - position at the same row level as current header
                                section_grid.addWidget(right_label, current_row, 3)
                                
                                # Create input field based on type
                                if col_value.startswith('fd_'):
                                    # Create dropdown for right field
                                    right_input_field = QComboBox()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setMinimumWidth(200)
                                    right_input_field.setStyleSheet("""
                                        QComboBox {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QComboBox:hover {
                                            border: 1px solid #3498DB;
                                        }
                                        QComboBox::drop-down {
                                            subcontrol-origin: padding;
                                            subcontrol-position: top right;
                                            width: 20px;
                                            border-left-width: 1px;
                                            border-left-color: #ccc;
                                            border-left-style: solid;
                                            border-top-right-radius: 4px;
                                            border-bottom-right-radius: 4px;
                                        }
                                    """)
                                    
                                    # Get options for dropdown
                                    right_options = []
                                    right_cell_col = col_idx + 1
                                    right_cell_address = f"{chr(ord('A') + right_cell_col)}{index + 1}"
                                    
                                    right_validation_options = self.get_validation_values(self.excel_path, sheet_name, right_cell_address)
                                    
                                    if right_validation_options:
                                        right_options = right_validation_options
                                    else:
                                        # Fallback if data validation not found
                                        if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                            right_options_str = str(row.iloc[col_idx + 1]).strip()
                                            right_options = [opt.strip() for opt in right_options_str.split(',')]
                                        
                                    # Add placeholder item "-- Select value --" first (KEY ADDITION)
                                    display_options = ["-- Select Value --"] + right_options
                                    right_input_field.addItems(display_options)
                                    
                                    # Style the placeholder item
                                    right_input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                                    right_input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                                    right_input_field.setCurrentIndex(0)  # Select placeholder by default
                                    
                                    # Only set a saved value if it exists and matches an option
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        saved_value = str(row.iloc[col_idx + 1]).strip()
                                        if saved_value in right_options:
                                            # Find the index in the display_options list (add 1 because of placeholder)
                                            option_index = right_options.index(saved_value) + 1
                                            right_input_field.setCurrentIndex(option_index)
                                    
                                    # Add to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                else:
                                    # It's a regular input field (f_ or fm_)
                                    right_input_field = QLineEdit()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setPlaceholderText(f"Enter {right_content}")
                                    right_input_field.setMinimumWidth(200)
                                    right_input_field.setStyleSheet("""
                                        QLineEdit {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QLineEdit:hover {
                                            border: 1px solid #3498DB;
                                        }
                                    """)
                                    
                                    # Set value if available
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        right_input_field.setText(str(row.iloc[col_idx + 1]).strip())
                                    
                                    # Add to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                
                                # Register the field
                                self.data_fields[right_field_key] = right_input_field
                
                current_row += 1
                continue

            # Check if it's a column header row (first cell starts with ch_ or has ch_ cells in the row)
            elif first_col.startswith('ch_') or any(isinstance(cell, str) and cell.startswith('ch_') for cell in row if not pd.isna(cell)):
                has_column_headers = True
                current_header_labels = []
                
                # Process header row and collect all ch_ columns
                for col_idx in range(df.shape[1]):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = str(row[col_idx]).strip() if isinstance(row[col_idx], str) else str(row[col_idx]).strip()
                        if col_value.startswith('ch_'):
                            header_text = col_value[3:].strip()  # Remove 'ch_' prefix
                            current_header_labels.append(header_text)
                
                # If we have header labels, create a header row
                if len(current_header_labels) > 0:
                    for col_idx, header_text in enumerate(current_header_labels):
                        col_header_label = QLabel(header_text)
                        col_header_label.setFont(QFont("Segoe UI", 11, QFont.Bold))
                        col_header_label.setStyleSheet("color: #555; margin-top: 5px; padding-left: 0px;")
                        col_header_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                        col_header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                        col_header_label.setIndent(0)  # Prevent text indentation
                        
                        # Grid column position depends on the header position
                        grid_col = col_idx  # Each field takes its own column in our grid
                        section_grid.addWidget(col_header_label, current_row, grid_col + 1)  # +1 to leave room for labels
                    
                    current_row += 1
                
                continue
            
            # Check if it's a field (f_)
            if first_col.startswith('f_'):
                field_name = first_col[2:].strip()  # Remove 'f_' prefix
        
                # Extract display name by removing numeric prefix if present
                display_name = field_name
                if field_name and field_name[0].isdigit():
                    # Remove digit prefix from display name
                    for i, char in enumerate(field_name):
                        if not char.isdigit():
                            display_name = field_name[i:]
                            break
                
                field_key = f"{sheet_name}_{current_section}_{field_count}"
                field_count += 1

                # Create field label
                label = QLabel(display_name)
                label.setFont(QFont("Segoe UI", 11))
                label.setStyleSheet("color: #333; background-color: transparent;")
                label.setMinimumWidth(250)  # Set minimum width for consistent layout
                label.setContentsMargins(0, 0, 0, 0)  # Ensure no margins
                label.setIndent(0)  # Set text indent to 0 for proper alignment
                
                # Add label to grid
                section_grid.addWidget(label, current_row, 0)

                # Create input field
                input_field = QLineEdit()
                input_field.setFont(QFont("Segoe UI", 11))
                input_field.setPlaceholderText(f"Enter {display_name}")
                input_field.setMinimumWidth(200)  # Set minimum width for consistent layout
                input_field.setStyleSheet("""
                    QLineEdit {
                        padding: 5px;
                        border: 1px solid #ccc;
                        border-radius: 4px;
                        background-color: white;
                        min-height: 28px;
                    }
                    QLineEdit:hover {
                        border: 1px solid #3498DB;
                    }
                """)
                
                # Set default value if available
                if len(row) > 1 and not pd.isna(row.iloc[1]):
                    input_field.setText(str(row.iloc[1]).strip())
                
                # Check if there's a unit value (fu_) in the next column
                has_unit = False
                unit_value = ""
                
                if len(row) > 2 and not pd.isna(row.iloc[2]):
                    unit_cell = str(row.iloc[2]).strip() if not pd.isna(row.iloc[2]) else ""
                    # Only process unit values with fu_ prefix
                    if isinstance(unit_cell, str) and unit_cell.startswith('fu_'):
                        has_unit = True
                        unit_value = unit_cell[3:].strip()  # Remove 'fu_' prefix
                
                # Check if there's any field in the right columns (columns C and beyond)
                has_right_field = False
                for col_idx in range(2, min(len(row), df.shape[1])):  # Start checking from column 2
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # If we find a unit (fu_), skip it for right field check
                        if isinstance(col_value, str) and col_value.startswith('fu_'):
                            continue
                            
                        # Check for any field prefix in the right columns
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fh_') or 
                            col_value.startswith('fm_')):
                            has_right_field = True
                            break
                
                # If there's a unit value, add it to the layout
                if has_unit:
                    # Create unit label
                    unit_label = QLabel(unit_value)
                    unit_label.setFont(QFont("Segoe UI", 11))
                    unit_label.setStyleSheet("color: #666; margin-left: 5px;")
                    unit_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                    
                    # Add unit label without changing the original input field position
                    section_grid.addWidget(input_field, current_row, 1)
                    section_grid.addWidget(unit_label, current_row, 2)
                else:
                    # If there's no field in the right columns, make this field span wider
                    if not has_right_field:
                        section_grid.addWidget(input_field, current_row, 1, 1, 2)  # Span 2 columns
                    else:
                        # Add input field to grid normally
                        section_grid.addWidget(input_field, current_row, 1)
                
                # Register the field
                self.data_fields[field_key] = input_field
                
                # Check for fields in right section (columns to the right)
                # Process any field type (f_, fd_, fh_) in the right section
                right_field_found = False
                
                for col_idx in range(2, min(len(row), df.shape[1])):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            # Convert non-string values to string
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # Check for any field prefix in the right section (f_, fd_, fh_)
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fh_') or 
                            col_value.startswith('fm_')):
                            
                            # Extract the right field prefix and name accordingly
                            prefix = col_value[:2] if col_value.startswith('f_') else col_value[:3]
                            suffix = col_value[2:] if col_value.startswith('f_') else col_value[3:]
                            right_field_name = suffix.strip()
                            
                            # If it's a header field (fh_)
                            if col_value.startswith('fh_'):
                                if not right_field_found:  # Only process the first header in this row
                                    right_header_label = QLabel(right_field_name)
                                    right_header_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
                                    right_header_label.setStyleSheet("color: #555; margin-top: 5px;")
                                    right_header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                                    right_header_label.setIndent(0)  # Prevent text indentation
                                    
                                    # Add header to the right section - position on column 3 only (don't span)
                                    section_grid.addWidget(right_header_label, current_row, 3, 1, 1)
                                    right_field_found = True
                                continue  # Skip further processing for headers
                                
                            # For regular fields or dropdowns
                            if not right_field_found:  # Only process the first field in this row
                                right_section_name = right_section if right_section else current_section
                                right_field_key = f"{sheet_name}_{right_section_name}_{field_count}"
                                field_count += 1
                                right_field_found = True
                                
                                # Create right field label
                                right_label = QLabel(right_field_name)
                                right_label.setFont(QFont("Segoe UI", 11))
                                right_label.setStyleSheet("color: #333; background-color: transparent;")
                                right_label.setMinimumWidth(250)  # Set minimum width for consistent layout
                                right_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                                right_label.setIndent(0)  # Prevent text indentation
                                
                                # Add label to grid - position at the same row level as current field
                                section_grid.addWidget(right_label, current_row, 3)
                                
                                # Handle different input field types based on prefix
                                if col_value.startswith('fd_'):
                                    # Create dropdown for right field
                                    right_input_field = QComboBox()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setMinimumWidth(200)  # Set minimum width for consistent layout
                                    right_input_field.setStyleSheet("""
                                        QComboBox {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QComboBox:hover {
                                            border: 1px solid #3498DB;
                                        }
                                        QComboBox::drop-down {
                                            subcontrol-origin: padding;
                                            subcontrol-position: top right;
                                            width: 20px;
                                            border-left-width: 1px;
                                            border-left-color: #ccc;
                                            border-left-style: solid;
                                            border-top-right-radius: 4px;
                                            border-bottom-right-radius: 4px;
                                        }
                                    """)
                                    
                                    # Get options for right dropdown
                                    right_options = []
                                    right_cell_col = col_idx + 1
                                    right_cell_address = f"{chr(ord('A') + right_cell_col)}{index + 1}"
                                    
                                    right_validation_options = self.get_validation_values(self.excel_path, sheet_name, right_cell_address)
                                    
                                    if right_validation_options:
                                        right_options = right_validation_options
                                    else:
                                        # Fallback if data validation not found
                                        if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                            right_options_str = str(row.iloc[col_idx + 1]).strip()
                                            right_options = [opt.strip() for opt in right_options_str.split(',')]
                                    
                                    # Add options to right dropdown
                                    display_options = ["-- Select Value --"] + right_options
                                    right_input_field.addItems(display_options)
                                    
                                    # Style the placeholder item
                                    right_input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                                    right_input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                                    right_input_field.setCurrentIndex(0)  # Select placeholder by default

                                    # Only set a saved value if it exists and matches an option
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        saved_value = str(row.iloc[col_idx + 1]).strip()
                                        if saved_value in right_options:
                                            # Find the index in the display_options list (add 1 because of placeholder)
                                            option_index = right_options.index(saved_value) + 1
                                            right_input_field.setCurrentIndex(option_index)
                                    
                                    # Add dropdown to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                else:
                                    # Regular input field (f_ or fm_)
                                    right_input_field = QLineEdit()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setPlaceholderText(f"Enter {right_field_name}")
                                    right_input_field.setMinimumWidth(200)  # Set minimum width for consistent layout
                                    right_input_field.setStyleSheet("""
                                        QLineEdit {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QLineEdit:hover {
                                            border: 1px solid #3498DB;
                                        }
                                    """)
                                    
                                    # Set default value if available
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        right_input_field.setText(str(row.iloc[col_idx + 1]).strip())
                                    
                                    # Add input field to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                
                                # Register the field
                                self.data_fields[right_field_key] = right_input_field
                
                current_row += 1
                continue
            
            if first_col.startswith('fd_'):
                field_name = first_col[3:].strip()  # Remove 'fd_' prefix

                # Extract display name by removing numeric prefix
                display_name = field_name
                if field_name and field_name[0].isdigit():
                    for i, char in enumerate(field_name):
                        if not char.isdigit():
                            display_name = field_name[i:]
                            break
                            
                field_key = f"{sheet_name}_{current_section}_{field_count}"
                field_count += 1

                # Create field label
                label = QLabel(display_name)
                label.setFont(QFont("Segoe UI", 11))
                label.setStyleSheet("color: #333; background-color: transparent;")
                label.setMinimumWidth(250)
                
                # Add label to grid
                section_grid.addWidget(label, current_row, 0)

                # Create dropdown
                input_field = QComboBox()
                input_field.setFont(QFont("Segoe UI", 11))
                input_field.setMinimumWidth(200)
                input_field.setStyleSheet("""
                    QComboBox {
                        padding: 5px;
                        border: 1px solid #ccc;
                        border-radius: 4px;
                        background-color: white;
                        min-height: 28px;
                    }
                    QComboBox:hover {
                        border: 1px solid #3498DB;
                    }
                    QComboBox::drop-down {
                        subcontrol-origin: padding;
                        subcontrol-position: top right;
                        width: 20px;
                        border-left-width: 1px;
                        border-left-color: #ccc;
                        border-left-style: solid;
                        border-top-right-radius: 4px;
                        border-bottom-right-radius: 4px;
                    }
                """)
                
                # BAGIAN BARU: Deteksi field "Effluent Warranty" dan gunakan hard coded options
                is_effluent_warranty_field = "Effluent Warranty" in field_name
                is_industry_field = field_name == "Industry Classification"
                is_sub_industry_field = field_name == "Sub Industry Specification"
                is_province1_field = field_name == "1Province"
                is_city1_field = field_name == "1City"
                is_province2_field = field_name == "2Province"
                is_city2_field = field_name == "2City"

                # Try to get options from data validation or from second column
                options = []
                
                # Calculate cell address more accurately
                excel_row = index + 1  # Convert to 1-based Excel row
                excel_col = 2  # Column B
                cell_address = f"B{excel_row}"

                # Untuk field Effluent Warranty
                if is_effluent_warranty_field:
                    options = EFFLUENT_WARRANTY_OPTIONS                    
                # Untuk field Industry Classification
                elif is_industry_field:
                    options = list(INDUSTRY_SUBTYPE_MAPPING.keys())
                    industry_dropdown = input_field
                    industry_field_key = field_key
                # Untuk field Sub Industry Specification
                elif is_sub_industry_field:
                    sub_industry_dropdown = input_field
                    sub_industry_field_key = field_key
                    options = ["-- Select Industry First --"]
                # Untuk field 1Province
                elif is_province1_field:
                    options = INDONESIA_PROVINCES
                    province1_dropdown = input_field
                    province1_field_key = field_key
                    input_field.setProperty("dropdown_type", "province")
                    input_field.setProperty("province_number", "1")
                # Untuk field 1City
                elif is_city1_field:
                    city1_dropdown = input_field
                    city1_field_key = field_key
                    input_field.setProperty("placeholder_type", "city")
                    input_field.setProperty("city_number", "1")
                    options = ["-- Select Province First --"]
                # Untuk field 2Province
                elif is_province2_field:
                    options = INDONESIA_PROVINCES
                    province2_dropdown = input_field
                    province2_field_key = field_key
                    input_field.setProperty("dropdown_type", "province")
                    input_field.setProperty("province_number", "2")
                # Untuk field 2City
                elif is_city2_field:
                    city2_dropdown = input_field
                    city2_field_key = field_key
                    input_field.setProperty("placeholder_type", "city")
                    input_field.setProperty("city_number", "2")
                    options = ["-- Select Province First --"]
                elif "Pump Brand" in field_name or field_name == "Pump Brand Selection":
                    options = ["GRUNDFOS", "KSB", "XYLEM", "ITT GOULDS", "EBARA", "WILO", "FLOWREX", "CNP", "LEO"]
                    pump_brand_dropdown = input_field
                    pump_brand_field_key = field_key
                    input_field.setProperty("pump_dropdown_type", "pump_brand")
                elif "Pump Type" in field_name:
                    pump_type_dropdown = input_field
                    pump_type_field_key = field_key
                    input_field.setProperty("pump_dropdown_type", "pump_type")
                    options = ["-- Select Pump Brand First --"]
                elif "Pump Model" in field_name:
                    pump_model_dropdown = input_field  
                    pump_model_field_key = field_key
                    input_field.setProperty("pump_dropdown_type", "pump_model")
                    options = ["-- Select Pump Brand and Type First --"]
                else:
                    # Try multiple cell addresses in case the calculation is off
                    possible_addresses = [
                        cell_address,
                        f"C{excel_row}",
                        f"B{excel_row + 1}",
                        f"B{excel_row - 1}"
                    ]
                    
                    validation_options = []
                    for addr in possible_addresses:
                        validation_options = self.get_validation_values(self.excel_path, sheet_name, addr)
                        if validation_options:
                            break
                    
                    if validation_options:
                        options = validation_options
                    else:
                        # Fallback to old method if data validation not found
                        if len(row) > 1 and not pd.isna(row.iloc[1]):
                            options_str = str(row.iloc[1]).strip()
                            if options_str and options_str != "nan":
                                options = [opt.strip() for opt in options_str.split(',')]
                        
                    # Process options untuk field Process yang sama
                    if "Process" in field_name and field_name.replace("Process ", "").strip().isdigit():
                        if "Process 1" in field_name:
                            self.process_field_options = options
                        elif hasattr(self, 'process_field_options') and self.process_field_options:
                            options = self.process_field_options

                # Set default value if available
                default_value = ""
                if len(row) > 1 and not pd.isna(row.iloc[1]):
                    default_value = str(row.iloc[1]).strip()
                
                # Populate dropdown options
                if options:
                    # Add placeholder if no default value or if we want to force selection
                    if not default_value or default_value == "nan":
                        # For dependent fields, don't add extra placeholder
                        if any(msg in options[0] for msg in ["Select Industry First", "Select Province First"]):
                            input_field.addItems(options)
                            input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                            input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                            input_field.setCurrentIndex(0)
                        else:
                            display_options = ["-- Select Value --"] + options
                            input_field.addItems(display_options)
                            
                            # Style placeholder
                            input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                            input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                            input_field.setCurrentIndex(0)
                            
                            # Set default value if it exists and matches an option
                            if default_value and default_value in options:
                                index = options.index(default_value) + 1  # +1 for placeholder
                                input_field.setCurrentIndex(index)
                    else:
                        # No placeholder needed, add options directly
                        input_field.addItems(options)
                        
                        # Set default value if it exists
                        if default_value and default_value in options:
                            input_field.setCurrentText(default_value)
                        elif len(options) > 0:
                            input_field.setCurrentText(options[0])
                    
                else:
                    # No options found - add placeholder only
                    input_field.addItem("-- No Options Available --")
                    input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                    input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                
                # Stylesheet untuk tooltip
                app = QApplication.instance()
                app.setStyleSheet("""
                    QToolTip {
                        background-color: #F5F5F5;
                        color: #333333;
                        border: 1px solid #CCCCCC;
                        padding: 5px;
                        font: 10pt "Segoe UI";
                        opacity: 255;
                    }
                """)
                
                if is_effluent_warranty_field:
                    # Setup tooltip pada setiap item dropdown
                    setup_effluent_warranty_item_tooltips(input_field)
                    
                    # PERBAIKAN STYLING: Set warna teks menjadi hitam untuk semua item
                    for i in range(input_field.count()):
                        # Set warna teks hitam untuk setiap item
                        input_field.setItemData(i, QtGui.QColor("#000000"), Qt.ForegroundRole)
                        
                        # Kecuali untuk placeholder, tetap abu-abu
                        if input_field.itemText(i) == "-- Select Value --":
                            input_field.setItemData(i, QtGui.QColor("#999999"), Qt.ForegroundRole)
                            input_field.setItemData(i, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                    
                # Add tooltips for special fields (existing code continues...)
                if field_name == "Seismic Hazard Zone":
                    for i in range(input_field.count()):
                        zone_text = input_field.itemText(i)
                        if zone_text in SEISMIC_ZONE_DESCRIPTIONS:
                            input_field.setItemData(i, SEISMIC_ZONE_DESCRIPTIONS[zone_text], Qt.ToolTipRole)
                elif field_name == "Wind Speed Zone":
                    for i in range(input_field.count()):
                        level_text = input_field.itemText(i)
                        if level_text in WIND_SPEED_DESCRIPTIONS:
                            input_field.setItemData(i, WIND_SPEED_DESCRIPTIONS[level_text], Qt.ToolTipRole)

                # Check if there's any field in the right columns (columns C and beyond)
                has_right_field = False
                for col_idx in range(2, min(len(row), df.shape[1])):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # Check for any field prefix in the right columns
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fh_') or 
                            col_value.startswith('fm_')):
                            has_right_field = True
                            break

                # Position dropdown in grid
                if not has_right_field:
                    # If there's no field in the right columns, make this field span to match fm_ fields
                    section_grid.addWidget(input_field, current_row, 1, 1, 2)  # Span 2 columns
                else:
                    # Add dropdown to grid normally
                    section_grid.addWidget(input_field, current_row, 1)
                    
                # Register the field
                self.data_fields[field_key] = input_field
                
                # Check for fields in right section (columns to the right)
                # Process any field type (f_, fd_, fh_, fm_) in the right section
                right_field_found = False
                right_header_found = False
                
                for col_idx in range(2, min(len(row), df.shape[1])):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            # Convert non-string values to string
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # Check for any field prefix in the right section
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fh_') or 
                            col_value.startswith('fm_')):
                            
                            # Extract the right field prefix and name accordingly
                            prefix = col_value[:2] if col_value.startswith('f_') else col_value[:3]
                            suffix = col_value[2:] if col_value.startswith('f_') else col_value[3:]
                            right_field_name = suffix.strip()
                            
                            # Special handling for header fields (fh_) in the right section
                            if col_value.startswith('fh_') and not right_header_found:
                                right_header_found = True  # Track that we found a header
                                right_field_found = True   # Consider this as a field being processed
                                
                                # Create right field header
                                right_header_label = QLabel(right_field_name)
                                right_header_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
                                right_header_label.setStyleSheet("color: #555; margin-top: 5px;")
                                right_header_label.setContentsMargins(0, 0, 0, 0)  # Remove any default margins
                                right_header_label.setIndent(0)  # Prevent text indentation
                                
                                # Add header to the grid at the same row level - don't span columns
                                section_grid.addWidget(right_header_label, current_row, 3, 1, 1)
                                
                                # If this is first header for right section, treat it as section title if we don't have one yet
                                if right_section is None:
                                    right_section = right_field_name
                            
                            # For regular fields or dropdowns (if no header was found yet)
                            elif not right_field_found:
                                right_section_name = right_section if right_section else current_section
                                right_field_key = f"{sheet_name}_{right_section_name}_{field_count}"
                                field_count += 1
                                right_field_found = True
                                
                                # Create right field label
                                right_label = QLabel(right_field_name)
                                right_label.setFont(QFont("Segoe UI", 11))
                                right_label.setStyleSheet("color: #333; background-color: transparent;")
                                right_label.setMinimumWidth(250)  # Set minimum width for consistent layout
                                
                                # Add label to grid - position at the same row level as current field
                                section_grid.addWidget(right_label, current_row, 3)
                                
                                # Handle different input field types based on prefix
                                if col_value.startswith('fd_'):
                                    # Create dropdown for right field
                                    right_input_field = QComboBox()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setMinimumWidth(200)
                                    right_input_field.setStyleSheet("""
                                        QComboBox {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QComboBox:hover {
                                            border: 1px solid #3498DB;
                                        }
                                        QComboBox::drop-down {
                                            subcontrol-origin: padding;
                                            subcontrol-position: top right;
                                            width: 20px;
                                            border-left-width: 1px;
                                            border-left-color: #ccc;
                                            border-left-style: solid;
                                            border-top-right-radius: 4px;
                                            border-bottom-right-radius: 4px;
                                        }
                                    """)
                                    
                                    # Get options for dropdown
                                    right_options = []
                                    right_cell_col = col_idx + 1
                                    right_cell_address = f"{chr(ord('A') + right_cell_col)}{index + 1}"
                                    
                                    right_validation_options = self.get_validation_values(self.excel_path, sheet_name, right_cell_address)
                                    
                                    if right_validation_options:
                                        right_options = right_validation_options
                                    else:
                                        # Fallback if data validation not found
                                        if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                            right_options_str = str(row.iloc[col_idx + 1]).strip()
                                            right_options = [opt.strip() for opt in right_options_str.split(',')]
                                        
                                    # Add placeholder item "-- Select value --" first (KEY ADDITION)
                                    display_options = ["-- Select Value --"] + right_options
                                    right_input_field.addItems(display_options)
                                    
                                    # Style the placeholder item
                                    right_input_field.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                                    right_input_field.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                                    right_input_field.setCurrentIndex(0)  # Select placeholder by default
                                    
                                    # Only set a saved value if it exists and matches an option
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        saved_value = str(row.iloc[col_idx + 1]).strip()
                                        if saved_value in right_options:
                                            # Find the index in the display_options list (add 1 because of placeholder)
                                            option_index = right_options.index(saved_value) + 1
                                            right_input_field.setCurrentIndex(option_index)
                                    
                                    # Add to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                else:
                                    # Regular input field
                                    right_input_field = QLineEdit()
                                    right_input_field.setFont(QFont("Segoe UI", 11))
                                    right_input_field.setPlaceholderText(f"Enter {right_field_name}")
                                    right_input_field.setMinimumWidth(200)
                                    right_input_field.setStyleSheet("""
                                        QLineEdit {
                                            padding: 5px;
                                            border: 1px solid #ccc;
                                            border-radius: 4px;
                                            background-color: white;
                                            min-height: 28px;
                                        }
                                        QLineEdit:hover {
                                            border: 1px solid #3498DB;
                                        }
                                    """)
                                    
                                    # Set value if available
                                    if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                        right_input_field.setText(str(row.iloc[col_idx + 1]).strip())
                                    
                                    # Add to grid
                                    section_grid.addWidget(right_input_field, current_row, 4)
                                
                                # Register field
                                self.data_fields[right_field_key] = right_input_field
                
                current_row += 1
                continue
                
            # Check if it's a field multiple (fm_)
            if first_col.startswith('fm_'):
                field_name = first_col[3:].strip()  # Remove 'fd_' prefix
        
                # Extract display name by removing numeric prefix
                display_name = field_name
                if field_name and field_name[0].isdigit():
                    for i, char in enumerate(field_name):
                        if not char.isdigit():
                            display_name = field_name[i:]
                            break
                        
                field_key_base = f"{sheet_name}_{current_section}_{field_name}"
                
                # Create field label
                label = QLabel(display_name)
                label.setFont(QFont("Segoe UI", 11))
                label.setStyleSheet("color: #333; background-color: transparent;")
                label.setMinimumWidth(250)  # Set minimum width for consistent layout
                
                # Add label to grid
                section_grid.addWidget(label, current_row, 0)
                
                # Determine if we have column headers to use for field names
                header_names = []
                if has_column_headers and len(current_header_labels) > 0:
                    header_names = current_header_labels
                else:
                    # Default fallback column names based on the screenshot
                    header_names = ["Name", "Phone No/Email"]
                
                # KEY CHANGE: Instead of creating a container widget, we'll place fields directly in the grid
                # This aligns the fields with the column headers above
                
                # Create input fields based on the number of headers
                for i, header in enumerate(header_names):
                    if i >= 2:  # Limit to 2 columns in case there are more headers
                        break
                        
                    input_field = QLineEdit()
                    input_field.setFont(QFont("Segoe UI", 11))
                    input_field.setMinimumWidth(180)
                    
                    # Include the field name in the placeholder text
                    input_field.setPlaceholderText(f"Enter {display_name} {header}")
                    
                    input_field.setStyleSheet("""
                        QLineEdit {
                            padding: 5px;
                            border: 1px solid #ccc;
                            border-radius: 4px;
                            background-color: white;
                            min-height: 28px;
                        }
                        QLineEdit:hover {
                            border: 1px solid #3498DB;
                        }
                    """)
                    
                    # Set value if available
                    if i+1 < len(row) and not pd.isna(row.iloc[i+1]):
                        input_field.setText(str(row.iloc[i+1]).strip())
                    
                    # Place each field directly in its corresponding column position
                    # First field (Name) goes in column 1, second field (Phone No/Email) goes in column 2
                    section_grid.addWidget(input_field, current_row, i+1)
                    
                    # Register field
                    self.data_fields[f"{field_key_base}_{i}"] = input_field
                    field_count += 1
                
                # Check for fields in right section (columns to the right)
                right_field_found = False
                
                for col_idx in range(3, min(len(row), df.shape[1])):
                    if col_idx < len(row) and not pd.isna(row[col_idx]):
                        col_value = ""
                        if isinstance(row[col_idx], str):
                            col_value = row[col_idx].strip()
                        elif not pd.isna(row[col_idx]):
                            col_value = str(row[col_idx]).strip()
                        else:
                            continue
                            
                        # Check for any field prefix in right section
                        if (col_value.startswith('f_') or 
                            col_value.startswith('fd_') or 
                            col_value.startswith('fm_') or
                            col_value.startswith('fh_')):
                            
                            # Extract the prefix and name
                            if col_value.startswith('f_'):
                                prefix = col_value[:2]
                                suffix = col_value[2:]
                            else:
                                prefix = col_value[:3]
                                suffix = col_value[3:]
                            
                            right_field_name = suffix.strip()
                            
                            # Skip headers
                            if col_value.startswith('fh_'):
                                continue
                                
                            # Process the first field found in right section
                            if not right_field_found:
                                right_field_found = True
                                
                                right_section_name = right_section if right_section else current_section
                                right_field_key_base = f"{sheet_name}_{right_section_name}_{right_field_name}"
                                
                                # Create right field label
                                right_label = QLabel(right_field_name)
                                right_label.setFont(QFont("Segoe UI", 11))
                                right_label.setStyleSheet("color: #333; background-color: transparent;")
                                right_label.setMinimumWidth(250)
                                
                                section_grid.addWidget(right_label, current_row, 3)
                                
                                # For multiple fields (fm_) in right section
                                if col_value.startswith('fm_'):
                                    # Create container for right fields
                                    right_container = QWidget()
                                    right_layout = QHBoxLayout(right_container)
                                    right_layout.setContentsMargins(0, 0, 0, 0)
                                    right_layout.setSpacing(10)
                                    right_layout.setAlignment(Qt.AlignLeft)
                                    
                                    # Create fields based on number of headers
                                    for i, header in enumerate(header_names):
                                        if i >= 2:  # Limit to 2 columns
                                            break
                                            
                                        right_input = QLineEdit()
                                        right_input.setFont(QFont("Segoe UI", 11))
                                        right_input.setPlaceholderText(f"Enter {header}")
                                        right_input.setStyleSheet("""
                                            QLineEdit {
                                                padding: 5px;
                                                border: 1px solid #ccc;
                                                border-radius: 4px;
                                                background-color: white;
                                                min-height: 28px;
                                            }
                                            QLineEdit:hover {
                                                border: 1px solid #3498DB;
                                            }
                                        """)
                                        
                                        # Set value if available
                                        if col_idx + i + 1 < len(row) and not pd.isna(row.iloc[col_idx + i + 1]):
                                            right_input.setText(str(row.iloc[col_idx + i + 1]).strip())
                                        
                                        right_layout.addWidget(right_input)
                                        
                                        # Register field
                                        self.data_fields[f"{right_field_key_base}_{i}"] = right_input
                                        field_count += 1
                                    
                                    # Add container to grid
                                    section_grid.addWidget(right_container, current_row, 4, 1, 1)
                                else:
                                    # For other field types
                                    if col_value.startswith('fd_'):
                                        # Create dropdown
                                        right_input = QComboBox()
                                        right_input.setFont(QFont("Segoe UI", 11))
                                        right_input.setMinimumWidth(200)
                                        right_input.setStyleSheet("""
                                            QComboBox {
                                                padding: 5px;
                                                border: 1px solid #ccc;
                                                border-radius: 4px;
                                                background-color: white;
                                                min-height: 28px;
                                            }
                                            QComboBox:hover {
                                                border: 1px solid #3498DB;
                                            }
                                            QComboBox::drop-down {
                                                subcontrol-origin: padding;
                                                subcontrol-position: top right;
                                                width: 20px;
                                                border-left-width: 1px;
                                                border-left-color: #ccc;
                                                border-left-style: solid;
                                                border-top-right-radius: 4px;
                                                border-bottom-right-radius: 4px;
                                            }
                                        """)
                                        
                                        # Get options
                                        right_options = []
                                        if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                            right_options_str = str(row.iloc[col_idx + 1]).strip()
                                            right_options = [opt.strip() for opt in right_options_str.split(',')]
                                        
                                        right_input.addItems(right_options)
                                        
                                        if len(right_options) > 0:
                                            right_input.setCurrentText(right_options[0])
                                        
                                        section_grid.addWidget(right_input, current_row, 4)
                                    else:
                                        # Regular input field
                                        right_input = QLineEdit()
                                        right_input.setFont(QFont("Segoe UI", 11))
                                        right_input.setPlaceholderText(f"Enter {right_field_name}")
                                        right_input.setStyleSheet("""
                                            QLineEdit {
                                                padding: 5px;
                                                border: 1px solid #ccc;
                                                border-radius: 4px;
                                                background-color: white;
                                                min-height: 28px;
                                            }
                                            QLineEdit:hover {
                                                border: 1px solid #3498DB;
                                            }
                                        """)
                                        
                                        # Set value
                                        if col_idx + 1 < len(row) and not pd.isna(row.iloc[col_idx + 1]):
                                            right_input.setText(str(row.iloc[col_idx + 1]).strip())
                                        
                                        section_grid.addWidget(right_input, current_row, 4)
                                    
                                    # Register field
                                    self.data_fields[f"{right_field_key_base}_0"] = right_input
                                    field_count += 1
                            
                            break  # Only process the first field with a prefix
                
                current_row += 1
                continue
            
            # Check if it's a table group (ftg_)
            if first_col.startswith('ftg_'):
                group_name = first_col[4:].strip()  # Remove 'ftg_' prefix
                
                # Create group header label - span across all columns
                group_label = QLabel(group_name)
                group_label.setFont(QFont("Segoe UI", 11, QFont.Bold))  # Make it bold
                group_label.setStyleSheet("""
                    color: #333; 
                    background-color: #f0f0f0; 
                    padding: 5px;
                    border-bottom: 1px solid #ccc;
                """)
                group_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                
                # Add label to grid - span across all columns (4 columns: task, client, contractor, remarks)
                section_grid.addWidget(group_label, current_row, 0, 1, 4)
                
                current_row += 1
                continue
            
            # Check if it's a table note (ftn_)
            if first_col.startswith('ftn_'):
                note_text = first_col[4:].strip()  # Remove 'ftn_' prefix
                
                # Create note label with italic style
                note_label = QLabel(note_text)
                note_label.setFont(QFont("Segoe UI", 10))
                note_label.setStyleSheet("""
                    color: #666; 
                    background-color: #f9f9f9; 
                    padding: 3px 5px;
                    font-style: italic;
                """)
                note_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                
                # Add label to grid - span across all columns
                section_grid.addWidget(note_label, current_row, 0, 1, 4)
                
                current_row += 1
                continue

            # Check if it's a table item (ft_)
            elif first_col.startswith('ft_'):
                task_name = first_col[3:].strip()  # Remove 'ft_' prefix
                field_key_base = f"{sheet_name}_{current_section}_{field_count}"
                field_count += 1
                
                # Create task label
                task_label = QLabel(task_name)
                task_label.setFont(QFont("Segoe UI", 11))
                task_label.setStyleSheet("color: #333; background-color: transparent;")
                task_label.setMinimumWidth(1000)  # Wider for task text
                task_label.setWordWrap(True)  # Allow text wrapping for long task names
                
                # Add label to grid - span only column 0
                section_grid.addWidget(task_label, current_row, 0)
                
                # Create checkbox for Client
                client_checkbox = QCheckBox()
                client_checkbox.setStyleSheet("""
                    QCheckBox {
                        min-height: 20px;
                    }
                    QCheckBox::indicator {
                        width: 18px;
                        height: 18px;
                    }
                """)
                
                # Create container widget for client checkbox to center it
                client_container = QWidget()
                client_layout = QHBoxLayout(client_container)
                client_layout.setAlignment(Qt.AlignCenter)
                client_layout.setContentsMargins(0, 0, 0, 0)
                client_layout.addWidget(client_checkbox)
                
                section_grid.addWidget(client_container, current_row, 1)
                
                # Create checkbox for Contractor
                contractor_checkbox = QCheckBox()
                contractor_checkbox.setStyleSheet("""
                    QCheckBox {
                        min-height: 20px;
                    }
                    QCheckBox::indicator {
                        width: 18px;
                        height: 18px;
                    }
                """)
                
                # Create container widget for contractor checkbox to center it
                contractor_container = QWidget()
                contractor_layout = QHBoxLayout(contractor_container)
                contractor_layout.setAlignment(Qt.AlignCenter)
                contractor_layout.setContentsMargins(0, 0, 0, 0)
                contractor_layout.addWidget(contractor_checkbox)
                
                section_grid.addWidget(contractor_container, current_row, 2)
                
                # Create input field for Remarks
                remarks_input = QLineEdit()
                remarks_input.setFont(QFont("Segoe UI", 11))
                remarks_input.setPlaceholderText("Enter remarks")
                remarks_input.setStyleSheet("""
                    QLineEdit {
                        padding: 5px;
                        border: 1px solid #ccc;
                        border-radius: 4px;
                        background-color: white;
                        min-height: 28px;
                    }
                    QLineEdit:hover {
                        border: 1px solid #3498DB;
                    }
                """)
                section_grid.addWidget(remarks_input, current_row, 3)
                
                # Set default values if available
                if len(row) > 1 and not pd.isna(row.iloc[1]):
                    # For Client checkbox
                    client_value = str(row.iloc[1]).strip()
                    if client_value == 'ü' or client_value.lower() == 'true':
                        client_checkbox.setChecked(True)
                        
                if len(row) > 2 and not pd.isna(row.iloc[2]):
                    # For Contractor checkbox
                    contractor_value = str(row.iloc[2]).strip()
                    if contractor_value == 'ü' or contractor_value.lower() == 'true':
                        contractor_checkbox.setChecked(True)
                        
                if len(row) > 3 and not pd.isna(row.iloc[3]):
                    # For Remarks input
                    remarks_input.setText(str(row.iloc[3]).strip())
                
                def create_checkbox_handler(client_cb, contractor_cb, source):
                    def handler():
                        if source == 'client' and client_cb.isChecked():
                            contractor_cb.setChecked(False)
                        elif source == 'contractor' and contractor_cb.isChecked():
                            client_cb.setChecked(False)
                    return handler

                # Hubungkan dengan cara yang lebih eksplisit
                client_handler = create_checkbox_handler(client_checkbox, contractor_checkbox, 'client')
                contractor_handler = create_checkbox_handler(client_checkbox, contractor_checkbox, 'contractor')

                client_checkbox.clicked.connect(client_handler)
                contractor_checkbox.clicked.connect(contractor_handler)
                
                # Register the fields
                self.data_fields[f"{field_key_base}_client"] = client_checkbox
                self.data_fields[f"{field_key_base}_contractor"] = contractor_checkbox
                self.data_fields[f"{field_key_base}_remarks"] = remarks_input
                
                current_row += 1
                continue

        # Setelah loop pemrosesan baris selesai dan sebelum return, tambahkan:

        # Hubungkan dropdown industry dan sub-industry
        if industry_dropdown and sub_industry_dropdown:
            # Simpan pasangan dropdown untuk referensi nanti
            self.linked_dropdowns[industry_field_key] = sub_industry_field_key
            
            # Hubungkan event dropdown industry ke fungsi update untuk dropdown sub-industry
            industry_dropdown.currentTextChanged.connect(
                lambda text, child=sub_industry_dropdown: self.update_dependent_dropdown(text, child))
            
            # Trigger update awal untuk mengisi sub-industry berdasarkan nilai industry yang sudah dipilih
            industry_value = industry_dropdown.currentText()
            if industry_value and industry_value != "-- Select Value --":
                self.update_dependent_dropdown(industry_value, sub_industry_dropdown)
                
                # Cari nilai sub-industry yang tersimpan di Excel
                sub_industry_value = None
                sub_industry_row = None
                
                # Cari baris yang berisi 'fd_Sub Industry Specification'
                for row_idx, row in df.iterrows():
                    if pd.isna(row.iloc[0]):
                        continue
                        
                    first_col = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                    if first_col == 'fd_Sub Industry Specification':
                        sub_industry_row = row_idx
                        # Ambil nilai dari kolom B (indeks 1)
                        if len(row) > 1 and not pd.isna(row.iloc[1]):
                            sub_industry_value = str(row.iloc[1]).strip()
                        break
                
                # Jika nilai sub-industry ditemukan, set ke dropdown
                if sub_industry_value and sub_industry_value in [sub_industry_dropdown.itemText(i) for i in range(sub_industry_dropdown.count())]:
                    sub_industry_dropdown.setCurrentText(sub_industry_value)

        # Hubungkan dropdown province1 dan city1
        if province1_dropdown and city1_dropdown:
            # Simpan pasangan dropdown untuk referensi nanti
            self.linked_dropdowns[province1_field_key] = city1_field_key
            
            # Hubungkan event dropdown province ke fungsi update untuk dropdown city
            province1_dropdown.currentTextChanged.connect(
                lambda text, child=city1_dropdown: self.update_dependent_dropdown(text, child))
            
            # Trigger update awal untuk mengisi city berdasarkan nilai province yang sudah dipilih
            province1_value = province1_dropdown.currentText()
            if province1_value and province1_value != "-- Select Value --":
                self.update_dependent_dropdown(province1_value, city1_dropdown)
                
                # Cari nilai city yang tersimpan di Excel
                city1_value = None
                city1_row = None
                
                # Cari baris yang berisi 'fd_1City'
                for row_idx, row in df.iterrows():
                    if pd.isna(row.iloc[0]):
                        continue
                        
                    first_col = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                    if "1City" in first_col and first_col.startswith('fd_'):
                        city1_row = row_idx
                        # Ambil nilai dari kolom B (indeks 1)
                        if len(row) > 1 and not pd.isna(row.iloc[1]):
                            city1_value = str(row.iloc[1]).strip()
                        break
                
                # Jika nilai city ditemukan, set ke dropdown
                if city1_value and province1_value in INDONESIA_CITIES and city1_value in INDONESIA_CITIES[province1_value]:
                    city1_dropdown.setCurrentText(city1_value)

        # Hubungkan dropdown province2 dan city2
        if province2_dropdown and city2_dropdown:
            # Simpan pasangan dropdown untuk referensi nanti
            self.linked_dropdowns[province2_field_key] = city2_field_key
            
            # Hubungkan event dropdown province ke fungsi update untuk dropdown city
            province2_dropdown.currentTextChanged.connect(
                lambda text, child=city2_dropdown: self.update_dependent_dropdown(text, child))
            
            # Trigger update awal untuk mengisi city berdasarkan nilai province yang sudah dipilih
            province2_value = province2_dropdown.currentText()
            if province2_value and province2_value != "-- Select Value --":
                self.update_dependent_dropdown(province2_value, city2_dropdown)
                
                # Cari nilai city yang tersimpan di Excel
                city2_value = None
                city2_row = None
                
                # Cari baris yang berisi 'fd_2City'
                for row_idx, row in df.iterrows():
                    if pd.isna(row.iloc[0]):
                        continue
                        
                    first_col = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                    if "2City" in first_col and first_col.startswith('fd_'):
                        city2_row = row_idx
                        # Ambil nilai dari kolom B (indeks 1)
                        if len(row) > 1 and not pd.isna(row.iloc[1]):
                            city2_value = str(row.iloc[1]).strip()
                        break
                
                # Jika nilai city ditemukan, set ke dropdown
                if city2_value and province2_value in INDONESIA_CITIES and city2_value in INDONESIA_CITIES[province2_value]:
                    city2_dropdown.setCurrentText(city2_value)
                    
        # Hubungkan pump dropdowns (triple dependency: brand -> type -> model)
        if pump_brand_dropdown and pump_type_dropdown:
            # Simpan referensi untuk triple dependency
            self.linked_dropdowns[pump_brand_field_key] = pump_type_field_key
            
            # Hubungkan brand -> type
            pump_brand_dropdown.currentTextChanged.connect(
                lambda text, child=pump_type_dropdown: self.update_dependent_dropdown(text, child))

        if pump_brand_dropdown and pump_type_dropdown and pump_model_dropdown:
            # Fungsi khusus untuk menangani triple dependency
            def update_pump_model():
                brand = pump_brand_dropdown.currentText()
                pump_type = pump_type_dropdown.currentText()
                
                # Cek apakah brand memiliki turunan
                if brand not in PUMP_BRAND_TYPE_MAPPING:
                    # Brand tidak memiliki turunan - disable model dropdown
                    pump_model_dropdown.clear()
                    pump_model_dropdown.addItem("-- Not Required for this Brand --")
                    pump_model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                    pump_model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                    pump_model_dropdown.setEnabled(False)
                elif (brand and brand != "-- Select Value --" and 
                    pump_type and pump_type != "-- Select Value --" and
                    pump_type != "-- Not Required for this Brand --"):
                    self.update_pump_model_dropdown(brand, pump_type, pump_model_dropdown)
                else:
                    pump_model_dropdown.clear()
                    pump_model_dropdown.addItem("-- Select Pump Brand and Type First --")
                    pump_model_dropdown.setItemData(0, QtGui.QColor("#999999"), Qt.ForegroundRole)
                    pump_model_dropdown.setItemData(0, QtGui.QFont("Segoe UI", 10, QtGui.QFont.StyleItalic), Qt.FontRole)
                    pump_model_dropdown.setEnabled(True)
            
            # Hubungkan kedua parent dropdown ke function update model
            pump_brand_dropdown.currentTextChanged.connect(lambda: update_pump_model())
            pump_type_dropdown.currentTextChanged.connect(lambda: update_pump_model())
            
        # Trigger update awal untuk pump brand -> type
        if pump_brand_dropdown and pump_type_dropdown:
            brand_value = pump_brand_dropdown.currentText()
            if brand_value and brand_value != "-- Select Value --":
                self.update_dependent_dropdown(brand_value, pump_type_dropdown)
                
                # Cari nilai pump type yang tersimpan di Excel
                pump_type_value = None
                for row_idx, row in df.iterrows():
                    if pd.isna(row.iloc[0]):
                        continue
                        
                    first_col = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                    if "Pump Type" in first_col and first_col.startswith('fd_'):
                        if len(row) > 1 and not pd.isna(row.iloc[1]):
                            pump_type_value = str(row.iloc[1]).strip()
                        break
                
                # Set nilai pump type jika ditemukan
                if (pump_type_value and pump_type_value in 
                    [pump_type_dropdown.itemText(i) for i in range(pump_type_dropdown.count())]):
                    pump_type_dropdown.setCurrentText(pump_type_value)
                    
                    # Trigger update untuk pump model
                    if pump_model_dropdown:
                        self.update_pump_model_dropdown(brand_value, pump_type_value, pump_model_dropdown)
                        
                        # Cari nilai pump model yang tersimpan di Excel
                        pump_model_value = None
                        for row_idx, row in df.iterrows():
                            if pd.isna(row.iloc[0]):
                                continue
                                
                            first_col = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ""
                            if "Pump Model" in first_col and first_col.startswith('fd_'):
                                if len(row) > 1 and not pd.isna(row.iloc[1]):
                                    pump_model_value = str(row.iloc[1]).strip()
                                break
                        
                        # Set nilai pump model jika ditemukan
                        if (pump_model_value and pump_model_value in 
                            [pump_model_dropdown.itemText(i) for i in range(pump_model_dropdown.count())]):
                            pump_model_dropdown.setCurrentText(pump_model_value)
        
        # Process any excel images that may exist in this sheet
        if section_layout:
            self.process_excel_images(sheet_name, section_layout)
            
        # Add a save button for the sheet at the end (only for DIP sheets)
        if section_layout and sheet_name.startswith("DIP_"):
            # Add spacer
            section_layout.addSpacing(10)

            # Save button
            save_btn = QPushButton("Save Changes")
            save_btn.setFont(QFont("Segoe UI", 11, QFont.Bold))
            save_btn.setCursor(Qt.PointingHandCursor)
            save_btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {SECONDARY_COLOR};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 8px 15px;
                }}
                QPushButton:hover {{
                    background-color: #2980B9;
                }}
            """)
            save_btn.clicked.connect(lambda: self.save_sheet_data(sheet_name))

            button_layout = QHBoxLayout()
            button_layout.addStretch()
            button_layout.addWidget(save_btn)

            section_layout.addLayout(button_layout)
            
        # If no content was added, add a default message
        if not layout.count():
            no_data_label = QLabel("No form fields found in this sheet.")
            no_data_label.setAlignment(Qt.AlignCenter)
            no_data_label.setStyleSheet("color: #666; margin: 20px;")
            layout.addWidget(no_data_label)
    
    # Implementasi method lainnya yang perlu tetap ada (dari kode asli)
    def get_validation_values(self, excel_path, sheet_name, cell_address):
        """Mengambil nilai dari data validation di sebuah sel Excel"""
        from openpyxl import load_workbook
        import re
        
        workbook = None
        try:
            # Pastikan untuk memuat dengan data_only=False agar kita bisa mengakses validasi
            workbook = load_workbook(excel_path, data_only=False)
            
            if sheet_name not in workbook.sheetnames:
                print(f"Sheet {sheet_name} not found in workbook")
                return []
                
            sheet = workbook[sheet_name]
            
            # Parse cell address manually (e.g., "B5" -> column=2, row=5)
            try:
                # Use regex to parse cell address
                match = re.match(r'^([A-Z]+)(\d+)$', cell_address.upper())
                if not match:
                    print(f"Invalid cell address format: {cell_address}")
                    return []
                
                col_letters = match.group(1)
                row_num = int(match.group(2))
                
                # Convert column letters to number (A=1, B=2, etc.)
                col_num = 0
                for i, letter in enumerate(reversed(col_letters)):
                    col_num += (ord(letter) - ord('A') + 1) * (26 ** i)
                
                target_cell = sheet.cell(row=row_num, column=col_num)
            except Exception as e:
                print(f"Error parsing cell address {cell_address}: {str(e)}")
                return []
            
            # Method 1: Check data validations collection
            if hasattr(sheet, 'data_validations') and sheet.data_validations:
                for validation in sheet.data_validations.dataValidation:
                    if validation.type == "list":
                        # Check if our cell is in the validation range
                        for coord_range in validation.sqref.ranges:
                            if target_cell.coordinate in str(coord_range):
                                formula = validation.formula1
                                
                                if formula:
                                    # Handle different formula formats
                                    if formula.startswith('"') and formula.endswith('"'):
                                        # Direct list: "option1,option2,option3"
                                        formula = formula[1:-1]
                                        options = [val.strip() for val in formula.split(',')]
                                        return options
                                    elif formula.startswith('='):
                                        # Reference to another range
                                        try:
                                            ref_range = formula[1:]  # Remove '='
                                            # Handle sheet references like 'Sheet1!A1:A10'
                                            if '!' in ref_range:
                                                ref_sheet_name, ref_range = ref_range.split('!', 1)
                                                ref_sheet = workbook[ref_sheet_name]
                                            else:
                                                ref_sheet = sheet
                                            
                                            # Get values from the referenced range
                                            ref_cells = ref_sheet[ref_range]
                                            options = []
                                            if hasattr(ref_cells, '__iter__'):
                                                for cell_row in ref_cells:
                                                    if hasattr(cell_row, '__iter__'):
                                                        for cell in cell_row:
                                                            if cell.value is not None:
                                                                options.append(str(cell.value).strip())
                                                    else:
                                                        if cell_row.value is not None:
                                                            options.append(str(cell_row.value).strip())
                                            else:
                                                if ref_cells.value is not None:
                                                    options.append(str(ref_cells.value).strip())
                                            
                                            return options
                                        except Exception as ref_e:
                                            print(f"Error processing reference formula {formula}: {str(ref_e)}")
                                    else:
                                        # Simple list without quotes
                                        options = [val.strip() for val in formula.split(',')]
                                        return options
            
            # Method 2: Check if cell has direct validation (alternative approach)
            try:
                if hasattr(target_cell, 'data_validation') and target_cell.data_validation:
                    cell_validation = target_cell.data_validation
                    if hasattr(cell_validation, 'type') and cell_validation.type == 'list':
                        formula = cell_validation.formula1
                        if formula:
                            if formula.startswith('"') and formula.endswith('"'):
                                formula = formula[1:-1]
                            options = [val.strip() for val in formula.split(',')]
                            return options
            except Exception as e:
                print(f"Error checking direct cell validation: {str(e)}")
            
            # Method 3: Scan all validations and find matching cell coordinates
            if hasattr(sheet, 'data_validations') and sheet.data_validations:
                for validation in sheet.data_validations.dataValidation:
                    if validation.type == "list":
                        try:
                            for coord_range in validation.sqref.ranges:
                                # Check if our target coordinates fall within this range
                                if (coord_range.min_row <= row_num <= coord_range.max_row and
                                    coord_range.min_col <= col_num <= coord_range.max_col):
                                    formula = validation.formula1
                                    if formula:
                                        if formula.startswith('"') and formula.endswith('"'):
                                            formula = formula[1:-1]
                                        options = [val.strip() for val in formula.split(',')]
                                        return options
                        except Exception as range_e:
                            print(f"Error processing validation range: {str(range_e)}")
                            continue
            
            return []
                    
        except Exception as e:
            print(f"Error reading data validation for {cell_address}: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
        finally:
            # Always close the workbook if it was opened
            if workbook is not None:
                try:
                    workbook.close()
                except Exception as close_e:
                    print(f"Error closing workbook: {str(close_e)}")
                                            
    def save_sheet_data(self, sheet_name):
        """Save the form data back to the Excel file with loading screen"""
        
        def save_process(progress_callback=None):
            try:
                if progress_callback:
                    progress_callback(5, "Validating sheet data...")
                
                # Check if the sheet name exists
                if not sheet_name.startswith("DIP_"):
                    return "Only DIP sheets can be saved."
                
                import pandas as pd
                from openpyxl import load_workbook
                import os
                import time
                
                if progress_callback:
                    progress_callback(10, "Checking file accessibility...")
                
                # Check if file exists and is accessible
                if not os.path.exists(self.excel_path):
                    return f"Excel file not found: {self.excel_path}"
                    
                # Check if file is not opened by another process
                try:
                    with open(self.excel_path, 'a'):
                        pass
                except PermissionError:
                    return "Excel file is currently opened by another application. Please close it and try again."
                
                if progress_callback:
                    progress_callback(15, "Creating backup...")
                
                # Create a backup of the Excel file
                backup_path = self.excel_path + ".bak"
                try:
                    import shutil
                    shutil.copy2(self.excel_path, backup_path)
                    print(f"Backup created at: {backup_path}")
                except Exception as e:
                    print(f"Warning: Could not create backup: {str(e)}")
                
                if progress_callback:
                    progress_callback(20, "Loading Excel workbook...")
                
                # Load the Excel workbook with openpyxl
                wb = load_workbook(self.excel_path)
                
                if sheet_name not in wb.sheetnames:
                    return f"Sheet '{sheet_name}' not found in the Excel file."
                
                # Get the sheet
                sheet = wb[sheet_name]
                
                if progress_callback:
                    progress_callback(25, "Reading current data structure...")
                
                # Also load with pandas to help us find the field positions
                df = pd.read_excel(self.excel_path, sheet_name=sheet_name, header=None)
                
                # Create validation data maps to help match dropdowns with their correct options
                validation_data = {}
                
                if progress_callback:
                    progress_callback(30, "Processing data validations...")
                
                # Load validation info from workbook
                from openpyxl.worksheet.datavalidation import DataValidation
                
                # Process all data validations in the sheet
                if hasattr(sheet, 'data_validations'):
                    for validation in sheet.data_validations.dataValidation:
                        if validation.type == 'list':
                            # This is a dropdown validation
                            formula = validation.formula1
                            options = []
                            
                            # Extract options from formula
                            if formula.startswith('"') and formula.endswith('"'):
                                formula = formula[1:-1]
                                options = [opt.strip() for opt in formula.split(',')]
                            
                            # Get all cells this validation applies to
                            for coord in validation.sqref.ranges:
                                for row in range(coord.min_row, coord.max_row + 1):
                                    for col in range(coord.min_col, coord.max_col + 1):
                                        cell_key = f"row_{row-1}_col_{col-1}"  # Adjust for 0-based indexing
                                        cell_address = f"{chr(64 + col)}{row}"  # Convert to A1 notation
                                        validation_data[cell_key] = {
                                            'options': options,
                                            'cell': cell_address
                                        }

                if progress_callback:
                    progress_callback(35, "Building field position mapping...")
                
                # Create mapping of field identifiers to their Excel row positions
                field_positions = {}
                
                # First pass: scan the Excel to build a map of field identifiers to row numbers
                for row_idx, row in df.iterrows():
                    # Skip empty rows
                    if pd.isna(row).all():
                        continue
                    
                    # Get first column value
                    first_col = row.iloc[0] if not pd.isna(row.iloc[0]) else ""
                    
                    # Convert to string if not already
                    if not isinstance(first_col, str):
                        try:
                            first_col = str(first_col)
                        except:
                            continue
                    
                    # Store the row index for field identifiers
                    if first_col.startswith('fd_'):
                        field_name = first_col[3:].strip()  # Remove 'fd_' prefix
                        field_positions[field_name] = {
                            'row_idx': row_idx,
                            'excel_row': row_idx + 1,
                            'field_id': first_col,
                            'cell_address': f"B{row_idx + 1}"  # Column B
                        }

                if progress_callback:
                    progress_callback(40, "Mapping specific dropdown widgets...")
                
                # Create specific widget mapping based on field names and positions
                widget_mapping = {}
                
                # Find industry and sub-industry dropdown widgets
                industry_dropdown = None
                sub_industry_dropdown = None
                province1_dropdown = None
                city1_dropdown = None
                province2_dropdown = None
                city2_dropdown = None
                pump_brand_dropdown = None
                pump_type_dropdown = None
                pump_model_dropdown = None
                effluent_warranty_dropdown = None
                
                # Find specific dropdown widgets
                for key, widget in self.data_fields.items():
                    if not key.startswith(sheet_name) or not isinstance(widget, QComboBox):
                        continue
                    
                    # Get the widget's label by looking at the grid layout
                    widget_label = self._get_widget_label(widget)
                    
                    # Map based on widget labels or properties
                    dropdown_type = widget.property("dropdown_type") or ""
                    placeholder_type = widget.property("placeholder_type") or ""
                    province_number = widget.property("province_number") or ""
                    city_number = widget.property("city_number") or ""
                    pump_dropdown_type = widget.property("pump_dropdown_type") or ""
                     
                    # Identify specific dropdowns
                    if "Industry Classification" in widget_label:
                        industry_dropdown = widget
                        widget_mapping['Industry Classification'] = widget
                    elif "Sub Industry" in widget_label:
                        sub_industry_dropdown = widget
                        widget_mapping['Sub Industry Specification'] = widget
                    elif dropdown_type == "province" and province_number == "1":
                        province1_dropdown = widget
                        widget_mapping['1Province'] = widget
                    elif placeholder_type == "city" and city_number == "1":
                        city1_dropdown = widget
                        widget_mapping['1City'] = widget
                    elif dropdown_type == "province" and province_number == "2":
                        province2_dropdown = widget
                        widget_mapping['2Province'] = widget
                    elif placeholder_type == "city" and city_number == "2":
                        city2_dropdown = widget
                        widget_mapping['2City'] = widget
                    elif "Effluent Warranty" in widget_label:
                        effluent_warranty_dropdown = widget
                        widget_mapping['Effluent Warranty'] = widget
                    elif pump_dropdown_type == "pump_brand" or "Pump Brand" in widget_label:
                        pump_brand_dropdown = widget
                        widget_mapping['Pump Brand Selection'] = widget
                    elif pump_dropdown_type == "pump_type" or "Pump Type" in widget_label:
                        pump_type_dropdown = widget
                        widget_mapping['Pump Type'] = widget
                    elif pump_dropdown_type == "pump_model" or "Pump Model" in widget_label:
                        pump_model_dropdown = widget
                        widget_mapping['Pump Model'] = widget

                if progress_callback:
                    progress_callback(45, "Processing remaining dropdown widgets...")
                
                # Map remaining dropdown widgets based on field positions and validation data
                remaining_widgets = []
                for key, widget in self.data_fields.items():
                    if (not key.startswith(sheet_name) or not isinstance(widget, QComboBox) or
                        widget in [industry_dropdown, sub_industry_dropdown, province1_dropdown, 
                                city1_dropdown, province2_dropdown, city2_dropdown,
                                effluent_warranty_dropdown, pump_brand_dropdown, 
                                pump_type_dropdown, pump_model_dropdown]):  # TAMBAHAN: Include pump dropdowns
                        continue
                    remaining_widgets.append((key, widget))

                # Match remaining widgets to field positions
                for field_name, position_info in field_positions.items():
                    if field_name in widget_mapping:
                        continue  # Already mapped
                    
                    row_idx = position_info['row_idx']
                    cell_address = position_info['cell_address']
                    
                    # Get validation data for this cell
                    cell_key = f"row_{row_idx}_col_1"  # Column B (0-based: col_1)
                    expected_options = []
                    
                    if cell_key in validation_data:
                        expected_options = validation_data[cell_key]['options']
                    
                    # TAMBAHAN: Untuk field hardcoded, gunakan options dari konstanta
                    if "Effluent Warranty" in field_name:
                        expected_options = EFFLUENT_WARRANTY_OPTIONS
                    elif "Pump Brand" in field_name:
                        expected_options = ["GRUNDFOS", "KSB", "XYLEM", "ITT GOULDS", "EBARA", "WILO", "FLOWREX", "CNP", "LEO"]
                    elif "Pump Type" in field_name and pump_brand_dropdown:
                        # Get pump type options based on selected brand
                        selected_brand = pump_brand_dropdown.currentText()
                        if selected_brand in PUMP_BRAND_TYPE_MAPPING:
                            expected_options = PUMP_BRAND_TYPE_MAPPING[selected_brand]
                    elif "Pump Model" in field_name and pump_brand_dropdown and pump_type_dropdown:
                        # Get pump model options based on brand and type
                        selected_brand = pump_brand_dropdown.currentText()
                        selected_type = pump_type_dropdown.currentText()
                        model_key = (selected_brand, selected_type)
                        if model_key in PUMP_BRAND_TYPE_MODEL_MAPPING:
                            expected_options = PUMP_BRAND_TYPE_MODEL_MAPPING[model_key]
                    
                    # Find the best matching widget
                    best_widget = None
                    best_match_score = 0
                    best_widget_key = None
                    
                    for widget_key, widget in remaining_widgets:
                        if widget_key in [w[0] for w in widget_mapping.values() if isinstance(w, tuple)]:
                            continue  # Already used
                        
                        # Get widget options
                        widget_options = [widget.itemText(i) for i in range(widget.count())]
                        # Remove placeholder options
                        clean_widget_options = [opt for opt in widget_options if not opt.startswith("-- ")]
                        
                        # Calculate match score
                        match_score = 0
                        if expected_options:
                            # Count exact matches
                            exact_matches = len(set(expected_options) & set(clean_widget_options))
                            match_score = exact_matches / len(expected_options) if expected_options else 0
                            
                            # Bonus for complete match
                            if exact_matches == len(expected_options):
                                match_score += 1
                        else:
                            # If no expected options, any unmatched widget can be used
                            match_score = 0.5
                        
                        if match_score > best_match_score:
                            best_match_score = match_score
                            best_widget = widget
                            best_widget_key = widget_key
                    
                    # Map the best matching widget
                    if best_widget:
                        widget_mapping[field_name] = best_widget
                        # Remove from remaining widgets
                        remaining_widgets = [(k, w) for k, w in remaining_widgets if k != best_widget_key]

                if progress_callback:
                    progress_callback(50, "Saving dropdown field values...")
                
                # Track changes
                changes_made = 0
                changes_log = []

                # Save values for all mapped fields
                for field_name, widget in widget_mapping.items():
                    if field_name not in field_positions:
                        continue
                        
                    position_info = field_positions[field_name]
                    excel_row = position_info['excel_row']
                    
                    # Get widget value
                    value = widget.currentText()
                    
                    # Skip placeholder values
                    if value == "-- Select Value --":
                        value = ""
                    
                    # Update Excel cell
                    target_cell = sheet.cell(row=excel_row, column=2)  # Column B
                    old_value = target_cell.value
                    target_cell.value = value
                    
                    changes_made += 1
                    changes_log.append(f"Updated cell B{excel_row} ({field_name}): {old_value} -> {value}")

                if progress_callback:
                    progress_callback(55, "Processing text fields and form controls...")
                
                # Create cell-to-widget mapping for non-dropdown fields
                cell_to_widget_map = {}
                
                # Store processed items to avoid duplicates
                processed_table_items = set()
                
                # Map text fields (f_), multiple fields (fm_), and table fields (ft_)
                for row_idx, row in df.iterrows():
                    if pd.isna(row).all():
                        continue
                        
                    first_col = row.iloc[0] if not pd.isna(row.iloc[0]) else ""
                    if not isinstance(first_col, str):
                        try:
                            first_col = str(first_col)
                        except:
                            continue
                    
                    if not (first_col.startswith('f_') or first_col.startswith('fm_') or first_col.startswith('ft_')):
                        continue
                    
                    # Get field properties
                    field_type = first_col[:3] if not first_col.startswith('f_') else first_col[:2]
                    field_name = first_col[3:] if not first_col.startswith('f_') else first_col[2:]
                    
                    # Clean up the display name
                    display_name = field_name.strip()
                    if display_name and display_name[0].isdigit():
                        for i, char in enumerate(display_name):
                            if not char.isdigit():
                                display_name = display_name[i:]
                                break
                    
                    # Get the section for this row
                    section_name = self._get_section_for_row(df, row_idx)
                    
                    # Handle different field types
                    if field_type == 'f_':
                        # Regular input field
                        field_id = first_col
                        original_field_name = field_name.strip()
                        
                        widget_key_prefix = f"{sheet_name}_{section_name}_"
                        possible_keys = []
                        
                        for i in range(100):
                            possible_keys.append(f"{widget_key_prefix}{i}")
                            possible_keys.append(f"{widget_key_prefix}{original_field_name}_{i}")
                            possible_keys.append(f"{widget_key_prefix}{display_name}_{i}")
                        
                        possible_keys.append(f"{widget_key_prefix}{row_idx}")
                        
                        found_widget = None
                        found_key = None
                        
                        for test_key in possible_keys:
                            if test_key in self.data_fields:
                                widget = self.data_fields[test_key]
                                if (isinstance(widget, QLineEdit) and 
                                    test_key not in [info.get('key') for info in cell_to_widget_map.values() if 'key' in info]):
                                    
                                    placeholder = widget.placeholderText()
                                    if placeholder:
                                        if field_id in placeholder:
                                            found_widget = widget
                                            found_key = test_key
                                            break
                                        elif original_field_name in placeholder:
                                            found_widget = widget
                                            found_key = test_key
                                        elif display_name in placeholder and not found_widget:
                                            found_widget = widget
                                            found_key = test_key
                        
                        if found_widget:
                            cell_key = f"row_{row_idx}_col_1"
                            cell_to_widget_map[cell_key] = {
                                'widget': found_widget,
                                'key': found_key,
                                'type': 'text',
                                'display_name': display_name,
                                'original_field': original_field_name,
                                'field_id': field_id
                            }
                    
                    elif field_type == 'fm_':
                        # Multiple field handling
                        field_base = f"{sheet_name}_{section_name}_{field_name}"
                        field_key_0 = f"{field_base}_0"
                        field_key_1 = f"{field_base}_1"
                        
                        if field_key_0 in self.data_fields and field_key_1 in self.data_fields:
                            widget_0 = self.data_fields[field_key_0]
                            widget_1 = self.data_fields[field_key_1]
                            
                            cell_key_0 = f"row_{row_idx}_col_1"
                            cell_key_1 = f"row_{row_idx}_col_2"
                            
                            cell_to_widget_map[cell_key_0] = {
                                'widget': widget_0,
                                'key': field_key_0,
                                'type': 'text'
                            }
                            
                            cell_to_widget_map[cell_key_1] = {
                                'widget': widget_1,
                                'key': field_key_1,
                                'type': 'text'
                            }
                    
                    elif field_type == 'ft_':
                        # Table item handling
                        processed_table_items = getattr(self, '_processed_table_items', set())
                        
                        found_base_key = None
                        
                        for i in range(100):
                            base_key = f"{sheet_name}_{section_name}_{i}"
                            client_key = f"{base_key}_client"
                            contractor_key = f"{base_key}_contractor"
                            remarks_key = f"{base_key}_remarks"
                            
                            if (client_key in self.data_fields and 
                                contractor_key in self.data_fields and 
                                remarks_key in self.data_fields and
                                base_key not in processed_table_items):
                                
                                client_widget = self.data_fields[client_key]
                                contractor_widget = self.data_fields[contractor_key]
                                remarks_widget = self.data_fields[remarks_key]
                                
                                if (isinstance(client_widget, QCheckBox) and 
                                    isinstance(contractor_widget, QCheckBox) and 
                                    isinstance(remarks_widget, QLineEdit)):
                                    
                                    cell_key_client = f"row_{row_idx}_col_1"
                                    cell_key_contractor = f"row_{row_idx}_col_2"
                                    cell_key_remarks = f"row_{row_idx}_col_3"
                                    
                                    cell_to_widget_map[cell_key_client] = {
                                        'widget': client_widget,
                                        'key': client_key,
                                        'type': 'checkbox'
                                    }
                                    
                                    cell_to_widget_map[cell_key_contractor] = {
                                        'widget': contractor_widget,
                                        'key': contractor_key,
                                        'type': 'checkbox'
                                    }
                                    
                                    cell_to_widget_map[cell_key_remarks] = {
                                        'widget': remarks_widget,
                                        'key': remarks_key,
                                        'type': 'text'
                                    }
                                    
                                    found_base_key = base_key
                                    processed_table_items.add(base_key)
                                    break
                        
                        self._processed_table_items = processed_table_items

                if progress_callback:
                    progress_callback(60, "Processing right column fields...")

                # NOW HANDLE RIGHT COLUMN FIELDS (Columns C/D, E/F, etc.)
                # This is crucial for fields like "Competitor Information" 
                for row_idx in range(len(df)):
                    # Skip empty rows
                    if pd.isna(df.iloc[row_idx]).all():
                        continue
                    
                    # Check columns starting from the third column (index 2 = column C)
                    # This is where right fields start
                    for col_idx in range(2, len(df.columns)):
                        if col_idx >= len(df.columns):
                            break
                            
                        # Get the value in this cell
                        cell_value = ""
                        if col_idx < len(df.iloc[row_idx]) and not pd.isna(df.iloc[row_idx, col_idx]):
                            cell_value = df.iloc[row_idx, col_idx]
                            
                            # Convert to string if needed
                            if not isinstance(cell_value, str):
                                try:
                                    cell_value = str(cell_value)
                                except:
                                    continue
                        
                        # Only process cells with field identifiers
                        if not (cell_value.startswith('f_') or 
                            cell_value.startswith('fd_') or 
                            cell_value.startswith('fm_')):
                            continue
                        
                        # Skip header cells
                        if cell_value.startswith('fh_'):
                            continue
                        
                        # Get field properties
                        field_type = cell_value[:3] if not cell_value.startswith('f_') else cell_value[:2]
                        field_name = cell_value[3:] if not cell_value.startswith('f_') else cell_value[2:]
                        
                        # Clean up the display name
                        display_name = field_name.strip()
                        if display_name and display_name[0].isdigit():
                            for i, char in enumerate(display_name):
                                if not char.isdigit():
                                    display_name = display_name[i:]
                                    break
                        
                        # Get the section for this row (either right section or default)
                        section_name = self._get_right_section_for_row(df, row_idx, col_idx)
                        if not section_name:
                            section_name = self._get_section_for_row(df, row_idx)
                        
                        # Find the matching widget based on field type
                        if field_type == 'f_':
                            # Regular input field
                            # Look for a QLineEdit with this name in the placeholder
                            for key, widget in self.data_fields.items():
                                if not key.startswith(sheet_name):
                                    continue
                                    
                                if isinstance(widget, QLineEdit) and key not in [info.get('key') for info in cell_to_widget_map.values() if 'key' in info]:
                                    placeholder = widget.placeholderText()
                                    if placeholder and display_name in placeholder:
                                        target_col = col_idx + 1  # Next column
                                        cell_key = f"row_{row_idx}_col_{target_col}"
                                        cell_to_widget_map[cell_key] = {
                                            'widget': widget,
                                            'key': key,
                                            'type': 'text'
                                        }
                                        break
                        
                        elif field_type == 'fd_':
                            # Dropdown field - check if it's already mapped in our widget_mapping
                            target_col = col_idx + 1  # Next column
                            cell_key = f"row_{row_idx}_col_{target_col}"
                            
                            # Skip if this field is already handled in widget_mapping
                            if field_name in widget_mapping:
                                continue
                            
                            # Find the options for this dropdown
                            dropdown_key = f"{row_idx}_{col_idx}"
                            options = []
                            
                            # Try to get options from validation data
                            validation_key = f"row_{row_idx}_col_{target_col}"
                            if validation_key in validation_data:
                                options = validation_data[validation_key]['options']
                            
                            # Find a QComboBox that isn't already mapped
                            best_widget = None
                            best_widget_key = None
                            best_match_score = 0
                            
                            for key, widget in self.data_fields.items():
                                if not key.startswith(sheet_name):
                                    continue
                                    
                                if isinstance(widget, QComboBox) and key not in [info.get('key') for info in cell_to_widget_map.values() if 'key' in info]:
                                    # Skip hardcode dropdowns that are already mapped
                                    if (widget in [industry_dropdown, sub_industry_dropdown, province1_dropdown, 
                                                city1_dropdown, province2_dropdown, city2_dropdown]):
                                        continue
                                    
                                    # Skip if already used in widget_mapping
                                    if widget in widget_mapping.values():
                                        continue
                                        
                                    # Check if widget options match
                                    widget_options = [widget.itemText(i) for i in range(widget.count())]
                                    # Remove placeholder option for comparison
                                    clean_widget_options = [opt for opt in widget_options if opt != "-- Select Value --"]
                                    
                                    match_score = 0
                                    
                                    if options:
                                        # Calculate exact match score
                                        exact_matches = len(set(options) & set(clean_widget_options))
                                        if len(options) > 0:
                                            match_score = exact_matches / len(options)
                                        
                                        # Bonus for having all expected options
                                        if exact_matches == len(options):
                                            match_score += 1
                                            
                                    else:
                                        # If no validation options found, use any available dropdown
                                        match_score = 0.5
                                    
                                    # Update best match
                                    if match_score > best_match_score:
                                        best_match_score = match_score
                                        best_widget = widget
                                        best_widget_key = key
                            
                            # Map the best matching widget
                            if best_widget:
                                cell_to_widget_map[cell_key] = {
                                    'widget': best_widget,
                                    'key': best_widget_key,
                                    'type': 'dropdown',
                                    'options': options,
                                    'display_name': display_name
                                }
                            
                        elif field_type == 'fm_':
                            # Multiple field (e.g., Name and Phone/Email)
                            found_base_key = None
                            
                            # Look for a pair of QLineEdit widgets for this field
                            for i in range(100):
                                base_key = f"{sheet_name}_{section_name}_{i}"
                                key_0 = f"{base_key}_0"
                                key_1 = f"{base_key}_1"
                                
                                if key_0 in self.data_fields and key_1 in self.data_fields:
                                    widget_0 = self.data_fields[key_0]
                                    widget_1 = self.data_fields[key_1]
                                    
                                    # Skip if already mapped
                                    if (key_0 in [info.get('key') for info in cell_to_widget_map.values() if 'key' in info] or
                                        key_1 in [info.get('key') for info in cell_to_widget_map.values() if 'key' in info]):
                                        continue
                                    
                                    if isinstance(widget_0, QLineEdit) and isinstance(widget_1, QLineEdit):
                                        # Check if placeholders contain the display name
                                        placeholder_0 = widget_0.placeholderText()
                                        placeholder_1 = widget_1.placeholderText()
                                        
                                        if ((placeholder_0 and display_name in placeholder_0) or 
                                            (placeholder_1 and display_name in placeholder_1)):
                                                
                                            # Map both widgets
                                            target_col_0 = col_idx + 1  # Next column
                                            target_col_1 = col_idx + 2  # Next column + 1
                                            
                                            cell_key_0 = f"row_{row_idx}_col_{target_col_0}"
                                            cell_key_1 = f"row_{row_idx}_col_{target_col_1}"
                                            
                                            cell_to_widget_map[cell_key_0] = {
                                                'widget': widget_0,
                                                'key': key_0,
                                                'type': 'text'
                                            }
                                            
                                            cell_to_widget_map[cell_key_1] = {
                                                'widget': widget_1,
                                                'key': key_1,
                                                'type': 'text'
                                            }
                                            
                                            found_base_key = base_key
                                            break

                if progress_callback:
                    progress_callback(70, "Processing tdi_ fields with placeholders...")

                # Handle tdi_ fields with placeholders (existing code)
                placeholder_widgets = {}

                for row_idx, row in df.iterrows():
                    if pd.isna(row).all():
                        continue
                        
                    for col_idx in range(len(row)):
                        cell_value = row.iloc[col_idx] if col_idx < len(row) and not pd.isna(row.iloc[col_idx]) else ""
                        
                        if not isinstance(cell_value, str):
                            try:
                                cell_value = str(cell_value)
                            except:
                                continue
                                
                        if cell_value.startswith('tdi_') and ('$P1$' in cell_value or '$P2$' in cell_value or 
                                                            '$P3$' in cell_value or '$P4$' in cell_value):
                            cell_key = f"row_{row_idx}_col_{col_idx}"
                            placeholder_widgets[cell_key] = {
                                'cell_value': cell_value,
                                'row': row_idx,
                                'col': col_idx,
                                'placeholders': []
                            }
                            
                            import re
                            placeholders = re.findall(r'\$P\d+\$', cell_value)
                            placeholder_widgets[cell_key]['placeholders'] = placeholders

                # Process placeholder widgets (existing code)
                for cell_key, info in placeholder_widgets.items():
                    row_idx = info['row']
                    col_idx = info['col']
                    cell_text = info['cell_value']
                    header_text = ""
                    
                    if col_idx > 0:
                        header_cell = df.iloc[row_idx, 0] if not pd.isna(df.iloc[row_idx, 0]) else ""
                        if isinstance(header_cell, str) and (header_cell.startswith('th_') or header_cell.startswith('thr_')):
                            if header_cell.startswith('th_'):
                                header_text = header_cell[3:].strip()
                            else:
                                header_text = header_cell[4:].strip()
                    
                    placeholders = info['placeholders']
                    
                    for key, widget in self.data_fields.items():
                        for placeholder in placeholders:
                            placeholder_key = f"tdi_{placeholder}_{header_text}"
                            placeholder_key_with_row = f"tdi_{placeholder}_{header_text}_{row_idx}"
                            
                            if (key == placeholder_key or key == placeholder_key_with_row) and isinstance(widget, QLineEdit):
                                value = widget.text()
                                
                                if 'placeholder_inputs' not in placeholder_widgets[cell_key]:
                                    placeholder_widgets[cell_key]['placeholder_inputs'] = {}
                                    
                                placeholder_widgets[cell_key]['placeholder_inputs'][placeholder] = {
                                    'widget': widget,
                                    'value': value,
                                    'key': key
                                }
                                
                                break

                if progress_callback:
                    progress_callback(80, "Updating Excel cells with new values...")

                # Update cells for ALL fields including right dropdown fields
                for cell_key, info in cell_to_widget_map.items():
                    if 'widget' not in info:
                        continue
                        
                    widget = info['widget']
                    widget_type = info.get('type', 'text')
                    
                    # Parse row and column from cell key
                    parts = cell_key.split('_')
                    row_idx = int(parts[1])
                    col_idx = int(parts[3])
                    
                    # Convert to Excel coordinates (1-based)
                    excel_row = row_idx + 1
                    excel_col = col_idx + 1
                    col_letter = chr(64 + excel_col)
                    
                    # Get cell value based on widget type
                    if widget_type == 'text' and isinstance(widget, QLineEdit):
                        value = widget.text()
                    elif widget_type == 'dropdown' and isinstance(widget, QComboBox):
                        # Handle dropdown fields (including right dropdown fields)
                        value = widget.currentText()
                        # Skip placeholder values
                        if value == "-- Select Value --":
                            value = ""
                    elif widget_type == 'checkbox' and isinstance(widget, QCheckBox):
                        value = 'ü' if widget.isChecked() else ''
                        target_cell = sheet.cell(row=excel_row, column=excel_col)
                        old_value = target_cell.value
                        target_cell.value = value
                        
                        from openpyxl.styles import Font
                        target_cell.font = Font(name='Wingdings', size=11)
                    else:
                        continue
                    
                    # Get display name for logging
                    display_name = info.get('display_name', '')
                    if not display_name:
                        # Try to determine field name from Excel content
                        if row_idx < len(df):
                            # For right fields, look at the identifier column
                            identifier_col = col_idx - 1
                            if identifier_col >= 2 and identifier_col < len(df.columns):
                                identifier_cell = df.iloc[row_idx, identifier_col] if not pd.isna(df.iloc[row_idx, identifier_col]) else ""
                                if isinstance(identifier_cell, str):
                                    if identifier_cell.startswith('fd_'):
                                        display_name = identifier_cell[3:].strip()
                                    elif identifier_cell.startswith('f_'):
                                        display_name = identifier_cell[2:].strip()
                                    
                                    # Clean up display name
                                    if display_name and display_name[0].isdigit():
                                        for i, char in enumerate(display_name):
                                            if not char.isdigit():
                                                display_name = display_name[i:]
                                                break
                    
                    # Add "Right" prefix for right columns
                    if col_idx >= 3:
                        display_name = f"Right {display_name}" if display_name else f"Right Field"
                    
                    # Update the cell
                    if widget_type != 'checkbox':  # Checkbox already handled above
                        target_cell = sheet.cell(row=excel_row, column=excel_col)
                        old_value = target_cell.value
                        target_cell.value = value
                    
                    changes_made += 1
                    changes_log.append(f"Updated cell {col_letter}{excel_row} ({display_name}): {old_value} -> {value}")

                if progress_callback:
                    progress_callback(85, "Processing placeholder replacements...")

                # Apply placeholder replacements
                placeholder_replacements = []

                for cell_key, info in placeholder_widgets.items():
                    row_idx = info['row']
                    col_idx = info['col']
                    cell_text = info['cell_value']
                    
                    if 'placeholder_inputs' not in info:
                        continue
                        
                    excel_row = row_idx + 1
                    excel_col = col_idx + 1
                    col_letter = chr(64 + excel_col)
                    
                    new_text = cell_text
                    
                    for placeholder, input_info in info['placeholder_inputs'].items():
                        value = input_info['value']
                        new_text = new_text.replace(placeholder, value)
                    
                    if new_text.startswith('tdi_'):
                        new_text = new_text[4:]
                    
                    placeholder_replacements.append({
                        'row': excel_row,
                        'col': excel_col, 
                        'value': new_text,
                        'original': sheet.cell(row=excel_row, column=excel_col).value
                    })

                for replacement in placeholder_replacements:
                    target_cell = sheet.cell(row=replacement['row'], column=replacement['col'])
                    old_value = replacement['original']
                    new_value = replacement['value']
                    
                    target_cell.value = new_value
                    
                    col_letter = chr(64 + replacement['col'])
                    cell_address = f"{col_letter}{replacement['row']}"
                    changes_made += 1
                    changes_log.append(f"Updated placeholder cell {cell_address}: {old_value} -> {new_value}")

                if progress_callback:
                    progress_callback(90, "Saving Excel workbook...")
                                
                # Save the workbook
                try:
                    wb.save(self.excel_path)
          
                    if progress_callback:
                        progress_callback(98, "Writing change log...")
                    
                    # Write detailed log to file for debugging
                    log_path = os.path.join(os.path.dirname(self.excel_path), "save_changes_log.txt")
                    with open(log_path, 'w') as log_file:
                        log_file.write(f"Save operation at {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                        log_file.write(f"Excel file: {self.excel_path}\n")
                        log_file.write(f"Sheet: {sheet_name}\n")
                        log_file.write(f"Total changes: {changes_made}\n\n")
                        log_file.write("Detailed changes:\n")
                        for change in changes_log:
                            log_file.write(f"- {change}\n")
                    
                    if progress_callback:
                        progress_callback(100, "Save completed successfully!")
                    
                    return f"Data successfully saved to {sheet_name} with {changes_made} changes."
                    
                except PermissionError:
                    return "Could not save the file. Make sure it is not open in Excel or another program."
                except Exception as e:
                    return f"Could not save the file: {str(e)}"
                
            except Exception as e:
                if progress_callback:
                    progress_callback(100, f"Error: {str(e)}")
                return f"An error occurred while processing data: {str(e)}"
        
        # Show loading screen for save operation
        loading_screen = LoadingScreen(
            parent=self,
            title="Saving Data",
            message=f"Saving changes to {sheet_name}..."
        )
        loading_screen.show()
        loading_screen.start_loading(save_process)
        
        # Connect completion handler
        def on_save_complete(success, message):
            if success and "successfully saved" in message:
                QMessageBox.information(
                    self, 
                    "Save Successful", 
                    message
                )
                
                # Reload the data to reflect changes - just reload the current tab
                current_tab_index = self.tab_widget.currentIndex()
                
                # Only reload the data for the current sheet to avoid freezing
                if sheet_name in self.sheet_tabs:
                    try:
                        # Get the tab's scroll area and its widget
                        scroll_area = self.tab_widget.widget(current_tab_index)
                        if isinstance(scroll_area, QScrollArea):
                            # Save scroll position
                            scroll_pos = scroll_area.verticalScrollBar().value()
                            
                            # Clear and reload just this sheet
                            df = pd.read_excel(self.excel_path, sheet_name=sheet_name, header=None)
                            
                            sheet_widget = scroll_area.widget()
                            if sheet_widget:
                                sheet_layout = sheet_widget.layout()
                                if sheet_layout:
                                    # Clear existing layout
                                    while sheet_layout.count():
                                        item = sheet_layout.takeAt(0)
                                        widget = item.widget()
                                        if widget:
                                            widget.deleteLater()
                                    
                                    # Recreate sheet content
                                    self.process_sheet_data(df, sheet_name, sheet_layout)
                                    
                                    # Restore scroll position
                                    QApplication.processEvents()
                                    scroll_area.verticalScrollBar().setValue(scroll_pos)
                                    
                    except Exception as e:
                        import traceback
                        traceback.print_exc()
            else:
                QMessageBox.critical(
                    self, 
                    "Save Error", 
                    message
                )
        
        loading_screen.worker.task_completed.connect(on_save_complete)
    
    def _get_widget_label(self, widget):
        """Helper method to get the label text for a widget by looking at the grid layout"""
        try:
            # Find all grid layouts in the widget hierarchy
            for section_grid in self.findChildren(QGridLayout):
                for row in range(section_grid.rowCount()):
                    for col in range(section_grid.columnCount()):
                        # Check if this position contains our target widget
                        widget_item = section_grid.itemAtPosition(row, col)
                        if widget_item and widget_item.widget() == widget:
                            # Look for a label in the previous column (usually column 0)
                            label_item = section_grid.itemAtPosition(row, col-1)
                            if label_item and label_item.widget():
                                label_widget = label_item.widget()
                                if isinstance(label_widget, QLabel):
                                    return label_widget.text()
            return ""
        except Exception as e:
            print(f"Error getting widget label: {str(e)}")
            return ""

    def _get_right_section_for_row(self, df, row_idx, col_idx):
        """Helper untuk mendapatkan section dari kolom kanan untuk baris tertentu"""
        # Start from the top and look for 'sub_' in the specified column
        current_right_section = None
        
        for i in range(row_idx + 1):  # Include current row
            if i >= len(df):
                break
                
            row = df.iloc[i]
            
            # Skip if column doesn't exist
            if col_idx >= len(row):
                continue
                
            col_value = row.iloc[col_idx] if not pd.isna(row.iloc[col_idx]) else ""
            
            if isinstance(col_value, str) and col_value.startswith('sub_'):
                current_right_section = col_value[4:].strip()
        
        return current_right_section

    def _get_section_for_row(self, df, row_idx):
        """Helper untuk mendapatkan section dari baris tertentu"""
        current_section = "Default"
        
        # Cari section terdekat sebelum row_idx
        for i in range(row_idx + 1):  # Include current row
            if i >= len(df):
                break
                
            row = df.iloc[i]
            first_col = row.iloc[0] if not pd.isna(row.iloc[0]) else ""
            
            if isinstance(first_col, str) and first_col.startswith('sub_'):
                current_section = first_col[4:].strip()
        
        return current_section

    def create_data_table(self, df, layout):
        """Create a table view for DATA sheets"""
        # Create a table widget
        table = QTableWidget()
        
        # Set row and column count
        table.setRowCount(len(df))
        table.setColumnCount(len(df.columns))
        
        # Set headers - convert all column headers to strings
        headers = [str(col) for col in df.columns.tolist()]
        table.setHorizontalHeaderLabels(headers)
        
        # Fill data - convert all values to strings to avoid type issues
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                if pd.isna(value):
                    item = QTableWidgetItem("")
                else:
                    item = QTableWidgetItem(str(value))
                table.setItem(row_idx, col_idx, item)
        
        # Style the table
        table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #E0E0E0;
                border-radius: 8px;
                background-color: white;
                gridline-color: #E0E0E0;
            }
            QHeaderView::section {
                background-color: #F5F5F5;
                padding: 5px;
                border: 1px solid #E0E0E0;
                font-weight: bold;
            }
            QTableWidget::item {
                padding: 5px;
            }
            QTableWidget::item:selected {
                background-color: #D6EAF8;
            }
        """)
        
        # Resize columns to content
        header = table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeToContents)
        header.setStretchLastSection(True)
        
        # Make it read-only for DATA sheets
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        
        # Add to layout
        layout.addWidget(table)
        
        # Add export button
        export_btn = QPushButton("Export to CSV")
        export_btn.setFont(QFont("Segoe UI", 11))
        export_btn.setCursor(Qt.PointingHandCursor)
        export_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {SECONDARY_COLOR};
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 15px;
                margin-top: 10px;
            }}
            QPushButton:hover {{
                background-color: #2980B9;
            }}
        """)
        export_btn.clicked.connect(lambda: self.export_data_table(df))
        
        export_layout = QHBoxLayout()
        export_layout.addStretch()
        export_layout.addWidget(export_btn)
        
        layout.addLayout(export_layout)
        
    def export_data_table(self, df):
        """Export DATA table to CSV"""
        try:
            # Ask for save location
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save CSV File", "", "CSV Files (*.csv);;All Files (*)"
            )
            
            if file_path:
                # Ensure it has .csv extension
                if not file_path.endswith('.csv'):
                    file_path += '.csv'
                
                # Save the dataframe
                df.to_csv(file_path, index=False)
                
                # Show success message
                QMessageBox.information(self, "Success", f"Data exported successfully to {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export data: {str(e)}")
            print(f"Error exporting data: {str(e)}")
            
    def closeEvent(self, event):
        """Cleanup on close"""
        try:
            # Cleanup formula evaluator
            if hasattr(self, 'formula_evaluator') and self.formula_evaluator:
                self.formula_evaluator.close()
                self.formula_evaluator = None
        except Exception as e:
            print(f"Error during cleanup: {str(e)}")
        
        # Call parent close event
        super().closeEvent(event)