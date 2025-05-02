# Excel ke Word dengan Referensi Sel - Mempertahankan Format Font
# Script ini mempertahankan format (bold, italic, dll) saat mengganti placeholder

import pandas as pd
import openpyxl
from docx import Document
import re
import os
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def excel_to_word_by_cell(excel_path, template_path, output_path):
    """
    Mengisi template Word dengan data dari file Excel berdasarkan referensi sel.
    Mempertahankan format font saat mengganti placeholder.
    
    Parameters:
    - excel_path: Path ke file Excel yang berisi data
    - template_path: Path ke template Word
    - output_path: Path untuk menyimpan hasil output Word
    """
    print(f"Membuka file Excel: {excel_path}")
    print(f"Membuka template Word: {template_path}")
    
    # Buka workbook Excel menggunakan openpyxl untuk akses sel langsung
    try:
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        print(f"Berhasil membuka file Excel. Sheet yang tersedia: {workbook.sheetnames}")
    except Exception as e:
        print(f"Error saat membuka file Excel: {e}")
        return
    
    # Fungsi untuk mendapatkan nilai sel
    def get_cell_value(sheet_name, cell_ref):
        if sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            try:
                cell_value = sheet[cell_ref].value
                print(f"Membaca {sheet_name}.{cell_ref} = {cell_value}")
                return cell_value
            except Exception as e:
                print(f"Error pada sel {sheet_name}.{cell_ref}: {e}")
                return f"ERROR: Invalid cell reference {cell_ref}"
        else:
            print(f"Sheet {sheet_name} tidak ditemukan")
            return f"ERROR: Sheet {sheet_name} not found"
    
    # Baca template Word
    try:
        doc = Document(template_path)
        print(f"Berhasil membuka template Word. Memiliki {len(doc.paragraphs)} paragraf dan {len(doc.tables)} tabel.")
    except Exception as e:
        print(f"Error saat membuka template Word: {e}")
        return
    
    # Pola untuk mendeteksi placeholder, mis: {{Sheet1.A1}}
    pattern = r'\{\{([^}]+)\.([A-Z]+[0-9]+)\}\}'
    replacement_count = 0
    
    # Fungsi untuk mengganti placeholder dalam paragraf dengan mempertahankan format
    def replace_in_paragraph_runs(paragraph):
        nonlocal replacement_count
        
        # Kita perlu melacak perubahan pada struktur run, karena ini bisa berubah saat kita memodifikasi
        orig_runs = list(paragraph.runs)
        placeholder_runs = {}  # Menyimpan info tentang placeholder di run mana

        # Identifikasi run mana yang berisi bagian dari placeholder
        for i, run in enumerate(orig_runs):
            if '{{' in run.text:
                # Kemungkinan awal dari placeholder
                full_placeholder = run.text
                run_index = i
                
                # Jika placeholder terbagi di beberapa run, kita perlu menyatukannya
                while '}}' not in full_placeholder and run_index < len(orig_runs) - 1:
                    run_index += 1
                    full_placeholder += orig_runs[run_index].text
                
                # Ekstrak bagian placeholder menggunakan regex
                matches = list(re.finditer(pattern, full_placeholder))
                for match in matches:
                    # Simpan informasi tentang placeholder ini
                    placeholder_text = match.group(0)  # {{Sheet1.A1}}
                    sheet_name = match.group(1)        # Sheet1
                    cell_ref = match.group(2)          # A1
                    
                    # Dapatkan nilai dari Excel
                    value = get_cell_value(sheet_name, cell_ref)
                    if value is not None:
                        value = str(value)
                    else:
                        value = ""
                    
                    # Simpan informasi untuk pemrosesan nanti
                    placeholder_runs[i] = {
                        'start_run': i,
                        'end_run': run_index,
                        'placeholder': placeholder_text,
                        'value': value
                    }
                    replacement_count += 1
        
        # Proses penggantian dimulai dari run terakhir untuk mencegah pergeseran indeks
        for start_run_idx in sorted(placeholder_runs.keys(), reverse=True):
            info = placeholder_runs[start_run_idx]
            
            # Kasus sederhana: placeholder ada dalam satu run
            if info['start_run'] == info['end_run']:
                run = orig_runs[info['start_run']]
                run.text = run.text.replace(info['placeholder'], info['value'])
            
            # Kasus kompleks: placeholder terbagi di beberapa run
            else:
                # Ambil run pertama dan terakhir
                first_run = orig_runs[info['start_run']]
                last_run = orig_runs[info['end_run']]
                
                # Dapatkan teks sebelum placeholder di run pertama
                before_text = first_run.text.split('{{')[0]
                
                # Dapatkan teks setelah placeholder di run terakhir
                after_text = last_run.text.split('}}')[1] if '}}' in last_run.text else ''
                
                # Gabungkan: teks sebelum + nilai + teks setelah
                first_run.text = before_text + info['value']
                
                # Hapus atau kosongkan run lain yang merupakan bagian dari placeholder
                for i in range(info['start_run'] + 1, info['end_run'] + 1):
                    orig_runs[i].text = ''
                
                # Tambahkan teks setelah ke run terakhir jika ada
                if after_text:
                    last_run.text = after_text
    
    # Fungsi untuk memproses tabel
    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph_runs(para)
    
    # Fungsi untuk memproses header dan footer
    def replace_in_section_headers_footers(doc):
        for section in doc.sections:
            # Header
            header = section.header
            for para in header.paragraphs:
                replace_in_paragraph_runs(para)
            replace_in_tables(header.tables)
            
            # Footer
            footer = section.footer
            for para in footer.paragraphs:
                replace_in_paragraph_runs(para)
            replace_in_tables(footer.tables)
    
    # Proses paragraf utama
    for para in doc.paragraphs:
        replace_in_paragraph_runs(para)
    
    # Proses tabel
    replace_in_tables(doc.tables)
    
    # Proses header dan footer
    replace_in_section_headers_footers(doc)
    
    print(f"Total {replacement_count} penggantian dilakukan.")
    
    # Simpan hasil ke file baru
    try:
        doc.save(output_path)
        print(f"Dokumen berhasil dihasilkan: {output_path}")
    except Exception as e:
        print(f"Error saat menyimpan dokumen Word: {e}")
    
    # Tutup workbook Excel
    workbook.close()

if __name__ == "__main__":
    # Folder data
    data_folder = "data"
    
    # Nama file
    excel_filename = "SET_BDU.xlsx"
    template_filename = "Trial WWTP ANP Quotation Template.docx"
    output_filename = "WWTP_Quotation_Result.docx"
    
    # Path lengkap
    excel_file = os.path.join(data_folder, excel_filename)
    template_file = os.path.join(data_folder, template_filename)
    output_file = os.path.join(data_folder, output_filename)
    
    # Pastikan folder data ada
    if not os.path.exists(data_folder):
        os.makedirs(data_folder)
        print(f"Folder {data_folder} dibuat.")
    
    # Periksa apakah file yang diperlukan ada
    if not os.path.exists(excel_file):
        print(f"KESALAHAN: File Excel '{excel_file}' tidak ditemukan!")
    elif not os.path.exists(template_file):
        print(f"KESALAHAN: File template Word '{template_file}' tidak ditemukan!")
    else:
        print(f"Mulai memproses file...")
        excel_to_word_by_cell(excel_file, template_file, output_file)