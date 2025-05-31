import pandas as pd
import openpyxl
from docx import Document
import re
import os
import datetime
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

def clean_filename(filename):
    """
    Clean filename from invalid characters
    
    Parameters:
    - filename: string filename
    
    Returns:
    - string: cleaned filename
    """
    # Invalid characters in Windows filenames
    invalid_chars = '<>:"/\\|?*'
    
    # Replace invalid characters with underscore
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    
    # Replace multiple spaces with single space
    filename = re.sub(r'\s+', ' ', filename)
    
    # Replace multiple underscores with single underscore
    filename = re.sub(r'_+', '_', filename)
    
    # Trim spaces at beginning and end
    filename = filename.strip()
    
    # Limit filename length (Windows limit ~255 characters, we use 200 for safety)
    if len(filename) > 200:
        # Cut but preserve extension
        name, ext = os.path.splitext(filename)
        filename = name[:200-len(ext)] + ext
    
    return filename

def get_quotation_number_by_company(workbook, company_name):
    """
    Find quotation number based on company name from 'No of Quotation' sheet
    
    Parameters:
    - workbook: openpyxl workbook object
    - company_name: company name to search
    
    Returns:
    - string: quotation number or default if not found
    """
    try:
        if 'No of Quotation' not in workbook.sheetnames:
            return "-"  # Default fallback
        
        # Access sheet using openpyxl
        sheet = workbook['No of Quotation']
        
        # Read header from first row to determine columns
        headers = []
        for col in range(1, sheet.max_column + 1):
            cell_value = sheet.cell(row=1, column=col).value
            if cell_value:
                headers.append(str(cell_value).strip())
            else:
                headers.append(f"Col_{col}")
        
        # Find required column indices with exact match
        company_col_idx = None
        quotation_col_idx = None
        
        for idx, header in enumerate(headers):
            col_idx = idx + 1  # openpyxl uses 1-based indexing
            
            # Exact match to avoid ambiguity
            if header == 'Company Name':  # Must be exact match
                company_col_idx = col_idx
            elif header == 'Quotation No.' or header == 'Quotation No':  # Support both formats
                quotation_col_idx = col_idx
        
        if company_col_idx is None or quotation_col_idx is None:
            return "-"
        
        # Clean company name for more accurate search
        company_name_clean = str(company_name).strip().lower()
        
        # Search data starting from row 2 (row 1 is header)
        for row in range(2, sheet.max_row + 1):
            company_cell = sheet.cell(row=row, column=company_col_idx).value
            quotation_cell = sheet.cell(row=row, column=quotation_col_idx).value
            
            if company_cell and quotation_cell:
                row_company = str(company_cell).strip()
                row_company_clean = row_company.lower()
                quotation_no = str(quotation_cell).strip()
                
                # Try exact match first
                if row_company_clean == company_name_clean:
                    return quotation_no
        
        # If exact match not found, try partial match
        for row in range(2, sheet.max_row + 1):
            company_cell = sheet.cell(row=row, column=company_col_idx).value
            quotation_cell = sheet.cell(row=row, column=quotation_col_idx).value
            
            if company_cell and quotation_cell:
                row_company_clean = str(company_cell).strip().lower()
                quotation_no = str(quotation_cell).strip()
                
                # Partial match - check if company name is in row or vice versa
                if (company_name_clean in row_company_clean) or (row_company_clean in company_name_clean):
                    return quotation_no
            
        return "-"  # Default fallback
        
    except Exception as e:
        print(f"Error getting quotation number: {str(e)}")
        return "-"  # Default fallback

def get_effluent_warranty_data(workbook, selected_warranty_type):
    """
    Get effluent warranty parameters based on selected warranty type
    
    Parameters:
    - workbook: openpyxl workbook object
    - selected_warranty_type: warranty type from DIP_Project Information.B83
    
    Returns:
    - dict: effluent warranty data with parameters and remarks
    """
    try:
        if 'Effluent Warranty' not in workbook.sheetnames:
            return {}
        
        sheet = workbook['Effluent Warranty']
        
        # Read headers from first row
        headers = {}
        for col in range(1, sheet.max_column + 1):
            cell_value = sheet.cell(row=1, column=col).value
            if cell_value:
                headers[str(cell_value).strip()] = col
        
        # Check required columns
        required_cols = ['Warranty_Type', 'Parameter_Name', 'Value', 'Unit', 'Row_Order']
        for req_col in required_cols:
            if req_col not in headers:
                return {}
        
        # Get column indices
        type_col = headers['Warranty_Type']
        param_col = headers['Parameter_Name']
        value_col = headers['Value']
        unit_col = headers['Unit']
        order_col = headers['Row_Order']
        
        # Find matching warranty type parameters
        warranty_params = []
        
        for row in range(2, sheet.max_row + 1):
            warranty_type = sheet.cell(row=row, column=type_col).value
            
            if warranty_type and str(warranty_type).strip() == str(selected_warranty_type).strip():
                param_name = sheet.cell(row=row, column=param_col).value
                param_value = sheet.cell(row=row, column=value_col).value
                param_unit = sheet.cell(row=row, column=unit_col).value
                row_order = sheet.cell(row=row, column=order_col).value
                
                # Skip if essential data is missing
                if not param_name:
                    continue
                
                # Handle Value: convert 0 to string "0", keep other values as string
                if param_value is None or param_value == "":
                    param_value_str = ""
                elif param_value == 0 or str(param_value).strip() == "0":
                    param_value_str = "0"  # Explicitly set to "0" string
                else:
                    param_value_str = str(param_value).strip()
                
                warranty_params.append({
                    'name': str(param_name).strip() if param_name else "",
                    'value': param_value_str,
                    'unit': str(param_unit).strip() if param_unit else "",
                    'order': int(row_order) if row_order and str(row_order).isdigit() else 999
                })
        
        # Sort by row order
        warranty_params.sort(key=lambda x: x['order'])
        
        # Set remarks to the selected warranty type itself
        warranty_remarks = selected_warranty_type
        
        # Create effluent data dictionary
        effluent_data = {}
        
        # Add parameters up to 22 (as specified in template)
        for i in range(1, 23):  # PARAM_1 to PARAM_22
            if i <= len(warranty_params):
                param = warranty_params[i-1]
                effluent_data[f'PARAM_{i}_NAME'] = param['name']
                effluent_data[f'PARAM_{i}_VALUE'] = param['value']
                effluent_data[f'PARAM_{i}_UNIT'] = param['unit']
            else:
                # Empty parameters for unused slots
                effluent_data[f'PARAM_{i}_NAME'] = ""
                effluent_data[f'PARAM_{i}_VALUE'] = ""
                effluent_data[f'PARAM_{i}_UNIT'] = ""
        
        # Add remarks from warranty type
        effluent_data['REMARKS'] = warranty_remarks
        
        # Store the actual number of parameters for table row management
        effluent_data['_PARAM_COUNT'] = len(warranty_params)
        
        return effluent_data
        
    except Exception as e:
        print(f"Error getting effluent warranty data: {str(e)}")
        import traceback
        traceback.print_exc()
        return {}

def remove_empty_effluent_table_rows(table, param_count):
    """
    Remove unused rows from effluent warranty table based on actual parameter count
    
    Parameters:
    - table: Word table object containing effluent parameters
    - param_count: actual number of parameters (rows to keep)
    
    Returns:
    - bool: True if successful
    """
    try:
        # Template has 22 parameter rows (PARAM_1 to PARAM_22)
        MAX_TEMPLATE_PARAMS = 22
        
        # Calculate how many rows need to be removed
        rows_to_remove_count = MAX_TEMPLATE_PARAMS - param_count
        
        if rows_to_remove_count <= 0:
            return True
    
        # Find all rows that contain parameter data
        effluent_rows_with_placeholders = []
        
        for i, row in enumerate(table.rows):
            row_has_effluent = False
            param_numbers = []
            
            # Check all cells in the row for EFFLUENT placeholders
            for cell in row.cells:
                for para in cell.paragraphs:
                    cell_text = para.text
                    
                    # Look for EFFLUENT patterns (with or without {{}})
                    if "EFFLUENT.PARAM_" in cell_text:
                        row_has_effluent = True
                        # Extract all parameter numbers from this row
                        import re
                        param_matches = re.findall(r'EFFLUENT\.PARAM_(\d+)_', cell_text)
                        for match in param_matches:
                            param_numbers.append(int(match))
            
            if row_has_effluent and param_numbers:
                # Use the minimum parameter number found in this row
                min_param_num = min(param_numbers)
                effluent_rows_with_placeholders.append((i, min_param_num, row))
        
        # Strategy 2: If no placeholders found (already replaced), use heuristic approach
        if not effluent_rows_with_placeholders:
            # Look for table structure pattern: find the "WWTP Effluent Warranty" header
            # and then count parameter rows from there
            header_row_idx = -1
            
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "WWTP Effluent Warranty" in para.text or "Effluent Warranty" in para.text:
                            header_row_idx = i
                            break
                    if header_row_idx >= 0:
                        break
                if header_row_idx >= 0:
                    break
            
            if header_row_idx >= 0:
                # Assume parameter rows start after header (usually next row)
                # and we need to remove the last (22 - param_count) rows
                total_rows = len(table.rows)
                
                # Calculate which rows to remove
                # Remove from the end backwards
                rows_to_remove = []
                start_remove_from = total_rows - rows_to_remove_count
                
                for i in range(start_remove_from, total_rows):
                    if i < len(table.rows):  # Safety check
                        rows_to_remove.append((i, 999, table.rows[i]))  # 999 as dummy param number
                
                effluent_rows_with_placeholders = rows_to_remove
        
        # Strategy 3: Simple approach - remove last N rows if nothing else works
        if not effluent_rows_with_placeholders:
            total_rows = len(table.rows)
            rows_to_remove = []
            
            # Remove the last rows_to_remove_count rows
            start_remove_from = max(1, total_rows - rows_to_remove_count)  # Keep at least header row
            
            for i in range(start_remove_from, total_rows):
                if i < len(table.rows):
                    rows_to_remove.append((i, 999, table.rows[i]))
            
            effluent_rows_with_placeholders = rows_to_remove
        
        # Sort by parameter number (or row index) to ensure correct order
        effluent_rows_with_placeholders.sort(key=lambda x: x[1] if x[1] != 999 else x[0])
        
        # Identify rows to remove based on parameter count
        rows_to_remove = []
        for row_idx, param_num, row_obj in effluent_rows_with_placeholders:
            if param_num == 999:  # Heuristic approach
                rows_to_remove.append((row_idx, param_num, row_obj))
            elif param_num > param_count:  # Placeholder approach
                rows_to_remove.append((row_idx, param_num, row_obj))
        
        # Sort by row index in reverse order to avoid index shifting when removing
        rows_to_remove.sort(key=lambda x: x[0], reverse=True)
        
        # Remove the rows from the table
        removed_count = 0
        for row_idx, param_num, row_obj in rows_to_remove:
            try:
                # Remove the row from table
                table._tbl.remove(row_obj._tr)
                removed_count += 1
            except Exception as e:
                print(f"Error removing row {row_idx}: {str(e)}")
        
        return removed_count > 0
        
    except Exception as e:
        print(f"Error removing empty effluent table rows: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def find_and_process_effluent_table(doc, effluent_data):
    """
    Find effluent warranty table and process it (remove unused rows)
    
    Parameters:
    - doc: Word document object
    - effluent_data: effluent data dictionary with _PARAM_COUNT
    
    Returns:
    - bool: True if table found and processed
    """
    try:
        param_count = effluent_data.get('_PARAM_COUNT', 0)
        
        # Find the table containing effluent warranty parameters
        effluent_table = None
        table_index = -1
        
        for i, table in enumerate(doc.tables):
            # Check if this table contains EFFLUENT placeholders
            table_has_effluent = False
            
            # Check all cells in all rows for EFFLUENT placeholders
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = ""
                    for para in cell.paragraphs:
                        cell_text += para.text + " "
                    
                    # Look for any EFFLUENT placeholder (not just PARAM_)
                    if "{{EFFLUENT." in cell_text or "EFFLUENT." in cell_text:
                        table_has_effluent = True
                        break
                if table_has_effluent:
                    break
            
            if table_has_effluent:
                effluent_table = table
                table_index = i
                break
        
        if effluent_table:
            # Remove unused rows
            success = remove_empty_effluent_table_rows(effluent_table, param_count)
            return success
        else:
            return False
            
    except Exception as e:
        print(f"Error processing effluent table: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def get_proposal_data_for_filename(workbook):
    """
    Get data required for generating proposal filename
    
    Parameters:
    - workbook: openpyxl workbook object
    
    Returns:
    - dict: data for filename
    """
    data = {
        'project_type': '',
        'capacity': '',
        'company_name': '',
        'user_code': ''
    }
    
    try:
        # 1. Project Type from DIP_Project_Information.B3
        if 'DIP_Project Information' in workbook.sheetnames:
            sheet = workbook['DIP_Project Information']
            try:
                project_type = sheet['B3'].value
                if project_type:
                    data['project_type'] = str(project_type).strip()
            except Exception:
                pass
        
        # 2. Capacity from DIP_Project_Information.B60
        if 'DIP_Project Information' in workbook.sheetnames:
            sheet = workbook['DIP_Project Information']
            try:
                capacity = sheet['B60'].value
                if capacity:
                    data['capacity'] = str(capacity).strip()
            except Exception:
                pass
        
        # 3. Company Name from DIP_Customer_Information.B4
        if 'DIP_Customer Information' in workbook.sheetnames:
            sheet = workbook['DIP_Customer Information']
            try:
                company_name = sheet['B4'].value
                if company_name:
                    data['company_name'] = str(company_name).strip()
            except Exception:
                pass
        
        # 4. User Code from DATA_TEMP.B1
        if 'DATA_TEMP' in workbook.sheetnames:
            sheet = workbook['DATA_TEMP']
            try:
                user_code = sheet['B1'].value
                if user_code:
                    data['user_code'] = str(user_code).strip()
            except Exception:
                pass
        
        return data
        
    except Exception as e:
        print(f"Error getting data for filename: {str(e)}")
        return data

def generate_dynamic_filename(workbook, fallback_customer_name=None, version="01"):
    """
    Generate dynamic proposal filename
    
    Format: Commercial and Technical_Project Type Capacity CMD_Company Name_Ver.01_BDE/PSE Code
    
    Parameters:
    - workbook: openpyxl workbook object
    - fallback_customer_name: fallback customer name if not in Excel
    - version: proposal version (default "01")
    
    Returns:
    - string: cleaned filename
    """
    try:
        # Get data from Excel
        data = get_proposal_data_for_filename(workbook)
        
        # Component 1: Fixed prefix
        prefix = "Commercial and Technical"
        
        # Component 2: Project Type
        project_type = data['project_type'] if data['project_type'] else "WWTP Project"
        
        # Component 3: Capacity with CMD
        capacity = data['capacity'] if data['capacity'] else "100"
        capacity_part = f"{capacity} CMD"
        
        # Component 4: Company Name
        company_name = data['company_name'] if data['company_name'] else fallback_customer_name if fallback_customer_name else "Unknown Company"
        
        # Component 5: Version
        version_part = f"Ver.{version}"
        
        # Component 6: User Code
        user_code = data['user_code'] if data['user_code'] else "000"
        
        # Combine all components
        filename_parts = [
            prefix,
            f"{project_type} {capacity_part}",
            company_name,
            version_part,
            user_code
        ]
        
        # Join with underscore separator
        filename = "_".join(filename_parts)
        
        # Add extension
        filename += ".docx"
        
        # Clean filename
        clean_name = clean_filename(filename)
        
        return clean_name
        
    except Exception as e:
        print(f"Error generating dynamic filename: {str(e)}")
        # Fallback to simple name
        fallback_name = f"WWTP_Quotation_{fallback_customer_name or 'Result'}_{version}.docx"
        return clean_filename(fallback_name)

def get_user_data_by_code(workbook, user_code):
    """
    Get user data based on user code from 'User Code' sheet
    
    Parameters:
    - workbook: openpyxl workbook object
    - user_code: string user code selected
    
    Returns:
    - dict: user data or empty dict if not found
    """
    try:
        if 'User Code' not in workbook.sheetnames:
            return {}
        
        # Access sheet using openpyxl
        sheet = workbook['User Code']
        
        # Read header from first row to determine columns
        headers = []
        for col in range(1, sheet.max_column + 1):
            cell_value = sheet.cell(row=1, column=col).value
            if cell_value:
                headers.append(str(cell_value).strip())
            else:
                headers.append(f"Col_{col}")
        
        # Find required column indices
        code_col_idx = None
        
        # Column mapping for user data
        column_mapping = {}
        
        for idx, header in enumerate(headers):
            col_idx = idx + 1  # openpyxl uses 1-based indexing
            
            if 'Code' in header:
                code_col_idx = col_idx
                column_mapping['Code'] = col_idx
            elif 'Name' in header:
                column_mapping['Name'] = col_idx
            elif 'Position' in header:
                column_mapping['Position'] = col_idx
            elif 'Email' in header:
                column_mapping['Email'] = col_idx
            elif 'Mobile' in header or 'Phone' in header:
                column_mapping['Mobile'] = col_idx
            else:
                # Add other columns that might exist
                column_mapping[header] = col_idx
        
        if code_col_idx is None:
            return {}
        
        # Clean user code for search
        user_code_clean = str(user_code).strip()
        
        # Search data starting from row 2 (row 1 is header)
        for row in range(2, sheet.max_row + 1):
            code_cell = sheet.cell(row=row, column=code_col_idx).value
            
            if code_cell:
                row_code = str(code_cell).strip()
                
                # Try exact match
                if row_code == user_code_clean:
                    # Get all data for this user
                    user_data = {}
                    for column_name, col_idx in column_mapping.items():
                        cell_value = sheet.cell(row=row, column=col_idx).value
                        if cell_value is not None:
                            user_data[column_name] = str(cell_value).strip()
                        else:
                            user_data[column_name] = ""
                    
                    return user_data
            
        return {}
        
    except Exception as e:
        print(f"Error getting user data: {str(e)}")
        return {}

def get_selected_user_code_from_excel(workbook):
    """
    Get selected user code from DATA_TEMP.B1
    """
    try:
        # Location for storing selected user code: DATA_TEMP.B1
        if 'DATA_TEMP' in workbook.sheetnames:
            sheet = workbook['DATA_TEMP']
            try:
                cell_value = sheet['B1'].value
                if cell_value and str(cell_value).strip():
                    return str(cell_value).strip()
                else:
                    return None
            except Exception:
                return None
        else:
            return None
        
    except Exception:
        return None

def get_selected_effluent_warranty_type(workbook):
    """
    Get selected effluent warranty type from the appropriate Excel location
    where the fd_Effluent Warranty field is saved by save_sheet_data
    
    Parameters:
    - workbook: openpyxl workbook object
    
    Returns:
    - string: selected warranty type or None if not found
    """
    try:
        # First, try to find the value in any DIP sheet where the fd_Effluent Warranty field might be
        for sheet_name in workbook.sheetnames:
            if sheet_name.startswith('DIP_'):
                try:
                    sheet = workbook[sheet_name]
                    # Scan the sheet to find fd_Effluent Warranty field
                    for row in range(1, sheet.max_row + 1):
                        cell_a = sheet.cell(row=row, column=1).value
                        if cell_a and str(cell_a).strip() == 'fd_Effluent Warranty':
                            # Found the field, get the value from column B
                            warranty_value = sheet.cell(row=row, column=2).value
                            if warranty_value and str(warranty_value).strip():
                                return str(warranty_value).strip()
                            break
                except Exception as e:
                    print(f"Error reading sheet {sheet_name}: {str(e)}")
                    continue
        
        # Fallback: check DIP_Project Information.B83 as secondary option
        if 'DIP_Project Information' in workbook.sheetnames:
            sheet = workbook['DIP_Project Information']
            try:
                warranty_type = sheet['B83'].value
                if warranty_type and str(warranty_type).strip():
                    return str(warranty_type).strip()
            except Exception as e:
                print(f"Error reading B83: {str(e)}")
        
        return None
        
    except Exception as e:
        print(f"Error getting effluent warranty type: {str(e)}")
        return None

def excel_to_word_by_cell(excel_path, template_path, output_path, selected_user_code=None):
    """
    Fill Word template with data from Excel file based on cell references.
    Maintains font formatting when replacing placeholders.
    Removes td_ and tdi_ prefixes from output.
    Changes ü character to Wingdings font.
    Converts all footer text to uppercase.
    Supports date placeholders.
    Supports mathematical placeholders $P1$ - $P4$ with fixed values.
    Supports USER_CODE placeholder for contact person data.
    Supports QUOTATION_NO placeholder for automatic quotation number.
    Supports EFFLUENT placeholder for effluent warranty parameters.
    
    Parameters:
    - excel_path: Path to Excel file containing data
    - template_path: Path to Word template
    - output_path: Path to save output Word file
    - selected_user_code: Selected user code (optional)
    
    Returns:
    - bool: True if successful, False if failed
    """
    
    # Prepare date data for placeholders
    now = datetime.datetime.now()
    
    # Date format DD MM YYYY
    day = now.strftime("%d")      # 2 digit format: 01, 02, ..., 31
    month = now.strftime("%m")     # 2 digit format: 01, 02, ..., 12
    year = now.strftime("%Y")     # 4 digit format: 2025
    full_date = f"{day} {month} {year}"
    
    # Create date data for placeholders
    date_data = {
        "NOW": full_date,
        "DAY": day,
        "MONTH": month,
        "YEAR": year
    }
    
    # Prepare data for mathematical placeholders
    math_placeholders = {
        "$P1$": "30",
        "$P2$": "50",
        "$P3$": "15",
        "$P4$": "5"
    }
    
    # Open Excel workbook using openpyxl for direct cell access
    try:
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception as e:
        print(f"Error opening Excel file: {e}")
        return False
    
    # Get selected user code if not provided as parameter
    if not selected_user_code:
        selected_user_code = get_selected_user_code_from_excel(workbook)
    
    # Get user data based on selected user code
    user_data = {}
    if selected_user_code:
        user_data = get_user_data_by_code(workbook, selected_user_code)
    
    # Prepare USER_CODE data for placeholders
    user_code_data = {
        "NAME": user_data.get('Name', 'Tia Amelia'),  # Default fallback
        "POSITION": user_data.get('Position', 'Product Strategist Engineer'),
        "EMAIL": user_data.get('Email', 'tiamalia@grinvirobiotekno.com'),
        "MOBILE": user_data.get('Mobile', '+62 856-5504-9457')
    }
    
    # Get company name to find quotation number
    company_name = ""
    try:
        if 'DIP_Customer Information' in workbook.sheetnames:
            sheet = workbook['DIP_Customer Information']
            company_cell = sheet['B4'].value
            if company_cell:
                company_name = str(company_cell).strip()
    except Exception:
        pass
    
    # Get quotation number based on company name
    quotation_number = get_quotation_number_by_company(workbook, company_name)
    
    # Prepare QUOTATION_NO data for placeholders
    quotation_data = {
        "NO": quotation_number
    }
    
    # Get effluent warranty data
    selected_warranty_type = get_selected_effluent_warranty_type(workbook)
    effluent_data = {}
    
    if selected_warranty_type:
        effluent_data = get_effluent_warranty_data(workbook, selected_warranty_type)
    else:
        # Create empty effluent data
        for i in range(1, 23):
            effluent_data[f'PARAM_{i}_NAME'] = ""
            effluent_data[f'PARAM_{i}_VALUE'] = ""
            effluent_data[f'PARAM_{i}_UNIT'] = ""
        effluent_data['REMARKS'] = ""
        effluent_data['_PARAM_COUNT'] = 0
    
    # Read Word template
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error opening Word template: {e}")
        return False
    
    # Class to handle placeholder replacement
    class Replacer:
        def __init__(self, effluent_data):
            self.replacement_count = 0
            self.math_replacement_count = 0
            self.user_code_replacement_count = 0
            self.quotation_no_replacement_count = 0
            self.effluent_replacement_count = 0
            self.effluent_data = effluent_data
            
        def get_cell_value(self, sheet_name, cell_ref):
            # Check if this is USER_CODE placeholder
            if sheet_name == "USER_CODE" and cell_ref in user_code_data:
                value = user_code_data[cell_ref]
                self.user_code_replacement_count += 1
                return value
            
            # Check if this is QUOTATION_NO placeholder
            if sheet_name == "QUOTATION_NO" and cell_ref in quotation_data:
                value = quotation_data[cell_ref]
                self.quotation_no_replacement_count += 1
                return value
            
            # Check if this is EFFLUENT placeholder
            if sheet_name == "EFFLUENT" and cell_ref in effluent_data:
                value = effluent_data[cell_ref]
                self.effluent_replacement_count += 1
                return value
            
            # Check if this is special date placeholder
            if sheet_name == "DATE" and cell_ref in date_data:
                value = date_data[cell_ref]
                return value
                
            # Existing code for getting values from Excel
            if sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                try:
                    cell_value = sheet[cell_ref].value
                    
                    # Process string values with td_ or tdi_ prefixes
                    if isinstance(cell_value, str):
                        # Remove td_ and tdi_ prefixes
                        if cell_value.startswith("td_"):
                            cell_value = cell_value[3:]  # Remove first 3 characters (td_)
                        elif cell_value.startswith("tdi_"):
                            cell_value = cell_value[4:]  # Remove first 4 characters (tdi_)
                        
                    return cell_value
                except Exception:
                    return f"ERROR: Invalid cell reference {cell_ref}"
            else:
                return f"ERROR: Sheet {sheet_name} not found"
        
        def set_wingdings_font(self, run):
            run.font.name = "Wingdings"
            # Add properties directly to XML to ensure font change
            run_properties = run._element.get_or_add_rPr()
            fonts = run_properties.xpath('./w:rFonts')
            if fonts:
                for font in fonts:
                    font.set(qn('w:ascii'), "Wingdings")
                    font.set(qn('w:hAnsi'), "Wingdings")
                    font.set(qn('w:cs'), "Wingdings")
            else:
                font_element = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Wingdings" w:hAnsi="Wingdings" w:cs="Wingdings"/>')
                run_properties.append(font_element)
        
        def replace_math_placeholders(self, text):
            """Replace mathematical placeholders $P1$-$P4$ with fixed values"""
            replaced = False
            
            # Find all mathematical placeholders in text
            for placeholder, value in math_placeholders.items():
                if placeholder in text:
                    text = text.replace(placeholder, value)
                    self.math_replacement_count += 1
                    replaced = True
            
            return text, replaced
        
        def replace_in_paragraph_runs(self, paragraph, is_footer=False):
            """Replace placeholders in paragraphs while maintaining formatting"""
            # Pattern to detect Excel placeholders, e.g.: {{Sheet1.A1}}, {{DATE.NOW}}, {{USER_CODE.NAME}}, {{QUOTATION_NO.NO}}, or {{EFFLUENT.PARAM_1_NAME}}
            pattern = r'\{\{([^}]+)\.([A-Z0-9_]+)\}\}'
            
            # Before processing Excel placeholders, check if there are mathematical placeholders
            for i, run in enumerate(paragraph.runs):
                # Check and replace mathematical placeholders
                if any(math_ph in run.text for math_ph in math_placeholders.keys()):
                    new_text, replaced = self.replace_math_placeholders(run.text)
                    if replaced:
                        run.text = new_text
            
            # We need to track changes in run structure, as this can change when we modify
            orig_runs = list(paragraph.runs)
            placeholder_runs = {}  # Store info about placeholders in which run
            
            # Identify which runs contain parts of placeholders
            for i, run in enumerate(orig_runs):
                if '{{' in run.text:
                    # Possible start of placeholder
                    full_placeholder = run.text
                    run_index = i
                    
                    # If placeholder is split across multiple runs, we need to combine them
                    while '}}' not in full_placeholder and run_index < len(orig_runs) - 1:
                        run_index += 1
                        full_placeholder += orig_runs[run_index].text
                    
                    # Extract placeholder parts using regex
                    matches = list(re.finditer(pattern, full_placeholder))
                    for match in matches:
                        # Store information about this placeholder
                        placeholder_text = match.group(0)  # {{Sheet1.A1}}, {{USER_CODE.NAME}}, {{QUOTATION_NO.NO}}, {{EFFLUENT.PARAM_1_NAME}}, etc
                        sheet_name = match.group(1)        # Sheet1, USER_CODE, DATE, QUOTATION_NO, EFFLUENT
                        cell_ref = match.group(2)          # A1, NAME, NOW, NO, PARAM_1_NAME
                        
                        # Get value from Excel, USER_CODE, QUOTATION_NO, EFFLUENT, or date placeholder
                        value = self.get_cell_value(sheet_name, cell_ref)
                        if value is not None:
                            value = str(value)
                            # If in footer, convert to uppercase
                            if is_footer:
                                value = value.upper()
                            
                            # Check if value contains mathematical placeholders and replace if necessary
                            for math_ph, math_val in math_placeholders.items():
                                if math_ph in value:
                                    value = value.replace(math_ph, math_val)
                                    self.math_replacement_count += 1
                        else:
                            value = ""
                        
                        # Store information for later processing
                        placeholder_runs[i] = {
                            'start_run': i,
                            'end_run': run_index,
                            'placeholder': placeholder_text,
                            'value': value,
                            'contains_umlaut': 'ü' in value
                        }
                        self.replacement_count += 1
            
            # Process replacements starting from last run to prevent index shifting
            for start_run_idx in sorted(placeholder_runs.keys(), reverse=True):
                info = placeholder_runs[start_run_idx]
                
                # Simple case: placeholder exists in one run
                if info['start_run'] == info['end_run']:
                    run = orig_runs[info['start_run']]
                    
                    # Check if value contains ü character
                    if 'ü' in info['value']:
                        # Split text based on ü character
                        parts = info['value'].split('ü')
                        
                        # Remove placeholder in original run and replace with first part
                        run.text = run.text.replace(info['placeholder'], parts[0])
                        
                        # For each ü and text after it
                        for i in range(len(parts) - 1):
                            # Create new run for ü character with Wingdings font
                            u_run = paragraph.add_run('ü')
                            self.set_wingdings_font(u_run)
                            
                            # Create new run for text after ü (if any)
                            if parts[i+1]:
                                paragraph.add_run(parts[i+1])
                    else:
                        # No ü, just replace placeholder
                        run.text = run.text.replace(info['placeholder'], info['value'])
                
                # Complex case: placeholder split across multiple runs
                else:
                    # Get first and last runs
                    first_run = orig_runs[info['start_run']]
                    last_run = orig_runs[info['end_run']]
                    
                    # Get text before placeholder in first run
                    before_text = first_run.text.split('{{')[0]
                    
                    # Get text after placeholder in last run
                    after_text = last_run.text.split('}}')[1] if '}}' in last_run.text else ''
                    
                    # Check if value contains ü character
                    if 'ü' in info['value']:
                        # Split text based on ü character
                        parts = info['value'].split('ü')
                        
                        # Set first part to first run
                        first_run.text = before_text + parts[0]
                        
                        # Clear or empty other runs that are part of placeholder
                        for i in range(info['start_run'] + 1, info['end_run'] + 1):
                            orig_runs[i].text = ''
                        
                        # For each ü and text after it
                        for i in range(len(parts) - 1):
                            # Create new run for ü character with Wingdings font
                            u_run = paragraph.add_run('ü')
                            self.set_wingdings_font(u_run)
                            
                            # If this is the last part and there's after_text
                            if i == len(parts) - 2 and after_text:
                                paragraph.add_run(parts[i+1] + after_text)
                            else:
                                # Create new run for text after ü (if any)
                                if parts[i+1]:
                                    paragraph.add_run(parts[i+1])
                    else:
                        # No ü, normal process
                        first_run.text = before_text + info['value']
                        
                        # Clear or empty other runs that are part of placeholder
                        for i in range(info['start_run'] + 1, info['end_run'] + 1):
                            orig_runs[i].text = ''
                        
                        # Add after text to last run if any
                        if after_text:
                            last_run.text = after_text
        
        def replace_in_tables(self, tables, is_footer=False):
            """Process tables"""
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.replace_in_paragraph_runs(para, is_footer)
        
        def replace_in_section_headers_footers(self, doc):
            """Process headers and footers"""
            for section in doc.sections:
                # Header
                header = section.header
                for para in header.paragraphs:
                    self.replace_in_paragraph_runs(para)
                self.replace_in_tables(header.tables)
                
                # Footer - with is_footer=True parameter for uppercase conversion
                footer = section.footer
                for para in footer.paragraphs:
                    # Check if paragraph contains page number field
                    contains_page_field = False
                    for run in para.runs:
                        # Page number field is usually represented with special characters or patterns
                        if run._element.xpath('.//w:fldChar') or "PAGE" in run.text or run._element.xpath('.//w:instrText'):
                            contains_page_field = True
                            break
                    
                    # Only process paragraphs that don't contain page number fields
                    if not contains_page_field:
                        self.replace_in_paragraph_runs(para, is_footer=True)
                    else:
                        # For paragraphs with page number fields, only convert normal text to uppercase
                        # without disturbing page number fields
                        for run in para.runs:
                            if not (run._element.xpath('.//w:fldChar') or run._element.xpath('.//w:instrText')):
                                # Convert to uppercase only if not part of field
                                run.text = run.text.upper()
                
                self.replace_in_tables(footer.tables, is_footer=True)
        
        def process_document(self, doc):
            """Enhanced process_document method"""
            # IMPORTANT: Process effluent table BEFORE text replacement
            # This ensures we can still find the placeholder patterns
            if self.effluent_data and '_PARAM_COUNT' in self.effluent_data:
                param_count = self.effluent_data.get('_PARAM_COUNT', 0)
                
                # Find and mark effluent table before text replacement
                effluent_table = self.find_effluent_table_before_replacement(doc)
                if effluent_table:
                    # Store reference for later processing
                    self.marked_effluent_table = effluent_table
                else:
                    self.marked_effluent_table = None
            
            # Process all text replacements
            # Process main paragraphs
            for para in doc.paragraphs:
                self.replace_in_paragraph_runs(para)
            
            # Process tables (including effluent table text replacement)
            self.replace_in_tables(doc.tables)
            
            # Process headers and footers
            self.replace_in_section_headers_footers(doc)
            
            # After text replacements, remove unused rows from the marked effluent table
            if hasattr(self, 'marked_effluent_table') and self.marked_effluent_table:
                param_count = self.effluent_data.get('_PARAM_COUNT', 0)
                success = remove_empty_effluent_table_rows(self.marked_effluent_table, param_count)
        
        def find_effluent_table_before_replacement(self, doc):
            """
            Find effluent table before text replacement occurs
            """
            try:
                for i, table in enumerate(doc.tables):
                    # Check if this table contains EFFLUENT placeholders
                    table_has_effluent = False
                    
                    # Check all cells in all rows for EFFLUENT placeholders
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = ""
                            for para in cell.paragraphs:
                                cell_text += para.text + " "
                            
                            # Look for EFFLUENT placeholders (should still have {{ }} format)
                            if "{{EFFLUENT." in cell_text:
                                table_has_effluent = True
                                break
                        if table_has_effluent:
                            break
                    
                    if table_has_effluent:
                        return table
                
                return None
                
            except Exception as e:
                print(f"Error in pre-scan: {str(e)}")
                return None
    
    # Use enhanced Replacer class to process document
    replacer = Replacer(effluent_data)
    replacer.process_document(doc)
    
    # Save result to new file
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            
        doc.save(output_path)
        success = True
    except Exception as e:
        print(f"Error saving Word document: {e}")
        success = False
    
    # Close Excel workbook
    workbook.close()
    
    return success

# Main function that can be called from other applications
def generate_proposal(excel_path, template_path, output_dir, selected_user_code=None, customer_name=None, version="01"):
    """
    Main function to be called from other applications with dynamic filename.
    
    Parameters:
    - excel_path: Path to Excel file containing data
    - template_path: Path to Word template
    - output_dir: Directory to save output Word file
    - selected_user_code: User code selected from dropdown
    - customer_name: Customer name for fallback filename
    - version: Proposal version (default "01")
    
    Returns:
    - tuple: (bool success, str output_path) - True if successful and result file path
    """
    # Validate input paths
    if not os.path.exists(excel_path):
        print(f"Error: Excel file not found: {excel_path}")
        return False, ""
        
    if not os.path.exists(template_path):
        print(f"Error: Word template not found: {template_path}")
        return False, ""
    
    # Ensure output directory exists
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception as e:
            print(f"Error creating output directory: {e}")
            return False, ""
    
    try:
        # Open workbook to generate filename
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        
        # Generate dynamic filename
        dynamic_filename = generate_dynamic_filename(
            workbook, 
            fallback_customer_name=customer_name, 
            version=version
        )
        
        # Close workbook temporarily
        workbook.close()
        
        # Complete output path
        output_path = os.path.join(output_dir, dynamic_filename)
        
    except Exception as e:
        print(f"Error generating dynamic filename: {str(e)}")
        # Fallback to simple filename
        fallback_filename = f"WWTP_Quotation_{customer_name or 'Result'}_{version}.docx"
        output_path = os.path.join(output_dir, clean_filename(fallback_filename))
            
    # Run function to create document
    success = excel_to_word_by_cell(excel_path, template_path, output_path, selected_user_code)
    
    # If successful in generating proposal, update cell A1 in DATA_PROPOSAL sheet
    if success:
        try:
            # Open the same Excel workbook
            wb = openpyxl.load_workbook(excel_path)
            
            # Check if DATA_PROPOSAL sheet exists
            if 'DATA_PROPOSAL' in wb.sheetnames:
                sheet = wb['DATA_PROPOSAL']
                
                # Get relative path from output_path
                # If output_path is in customer folder, create relative path to customer folder
                
                # Base path is project root
                root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                
                # If output_path is subpath of root_dir, create relative path
                if os.path.commonpath([root_dir]) == os.path.commonpath([root_dir, output_path]):
                    relative_path = os.path.relpath(output_path, root_dir)
                else:
                    # If different drive or no common path, use absolute path
                    relative_path = output_path
                
                # Convert to Windows path format with backslash
                relative_path = relative_path.replace('/', '\\')
                
                # Update cell A1 with relative path
                sheet['A1'] = relative_path
                
                # Save workbook
                wb.save(excel_path)
                wb.close()
            else:
                wb.close()
        except Exception as e:
            print(f"Error updating cell A1 in DATA_PROPOSAL sheet: {e}")
    
    return success, output_path if success else ""

# If this script is run directly (not imported)
if __name__ == "__main__":
    """If this file is run directly (not imported), 
    this code allows using customer path if provided as argument.
    
    Usage:
    python generate_proposal.py [customer_name] [user_code] [version]
    
    If customer_name is provided, script will process SET_BDU.xlsx file
    in that customer's folder and save output in the same folder.
    
    If user_code is provided, script will use that user's data.
    If version is provided, it will be used for file naming.
    """
    # Import argparse to handle command line arguments
    import argparse
    
    # Create argument parser
    parser = argparse.ArgumentParser(description='Generate proposal from SET_BDU.xlsx with dynamic filename and quotation number lookup')
    parser.add_argument('customer_name', nargs='?', help='Optional customer name')
    parser.add_argument('user_code', nargs='?', help='Optional user code')
    parser.add_argument('version', nargs='?', default='01', help='Optional version number (default: 01)')
    
    # Parse arguments
    args = parser.parse_args()
    
    # Determine relative path to project root
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    if args.customer_name:
        # If customer_name is provided, use customer-specific path
        from modules.fix_customer_system import clean_folder_name
        
        # Create customer folder path
        customer_folder = os.path.join(root_dir, "data", "customers", clean_folder_name(args.customer_name))
        
        # Customer file paths
        excel_file = os.path.join(customer_folder, "SET_BDU.xlsx")
        output_dir = customer_folder
        
        # Template path always refers to main data folder
        template_file = os.path.join(root_dir, "data", "Trial WWTP ANP Quotation Template.docx")
        
    else:
        # Use default path if no customer_name
        data_folder = os.path.join(root_dir, "data")
        
        # File names
        excel_filename = "SET_BDU.xlsx"
        template_filename = "Trial WWTP ANP Quotation Template.docx"
        
        # Complete paths
        excel_file = os.path.join(data_folder, excel_filename)
        template_file = os.path.join(data_folder, template_filename)
        output_dir = data_folder
    
    # Ensure data folder exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # Check if required files exist
    if not os.path.exists(excel_file):
        print(f"ERROR: Excel file '{excel_file}' not found!")
    elif not os.path.exists(template_file):
        print(f"ERROR: Word template file '{template_file}' not found!")
    else:
        # Call generate_proposal with complete parameters
        success, output_file = generate_proposal(
            excel_file, 
            template_file, 
            output_dir, 
            selected_user_code=args.user_code,
            customer_name=args.customer_name,
            version=args.version
        )
        
        if success:
            print(f"✅ Proposal successfully generated: {output_file}")
        else:
            print("❌ Failed to generate proposal")